"""
Operasi tingkat-rendah pada dokumen Word (python-docx) yang dipakai berulang
oleh berbagai generator surat: mengisi baris tabel dari master, memaksa page
break eksplisit, membersihkan paragraf placeholder kosong sisa render
docxtpl, dan menggabungkan beberapa file .docx jadi satu dokumen multi-halaman.
"""
import copy
import os
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.core.app_logging import safe_log


def ensure_numbered_slots(doc, prefix, needed_count):
    """Pastikan jumlah slot placeholder '{{prefix_N}}' (mis. pihak_terkait_1,
    pihak_terkait_2, dst -- termasuk paragraf List Paragraph berpenomoran
    otomatis di form 'Pihak Terkait') di TEMPLATE minimal `needed_count`.

    BUG YANG DIPERBAIKI: form UI ('+ Tambah Pihak Terkait') tidak membatasi
    jumlah entri yang bisa diketik pengguna, tapi master template .docx
    cuma punya sejumlah slot placeholder yang di-hardcode (mis. cuma sampai
    pihak_terkait_5). Kalau pengguna mengisi lebih dari itu, entri
    ke-(N+1) dst DIAM-DIAM HILANG dari dokumen tanpa pesan error apapun,
    karena docxtpl hanya mengisi placeholder yang benar-benar ada.

    Solusinya: sebelum render, cari slot placeholder bernomor TERBESAR yang
    ada, lalu kloning paragraf itu (termasuk numPr auto-numbering & rPr-nya
    persis) sebanyak selisih yang kurang, sisipkan tepat sesudahnya. Nomor
    tampilan (1, 2, 3, ...) TIDAK perlu ditulis manual -- itu dihitung
    ulang otomatis oleh Word dari urutan paragraf List Paragraph yang
    memakai numId yang sama, jadi hasilnya selalu berurutan benar berapa
    pun jumlah slot yang ditambahkan. Slot yang TIDAK terpakai (kosong)
    tetap dibuang seperti biasa oleh cleanup_skenario_paripurna setelah
    render, jadi fungsi ini aman dipanggil dengan `needed_count` sebesar
    apa pun, termasuk 0 (tidak melakukan apa-apa)."""
    if needed_count <= 0:
        return
    pattern = re.compile(r'^\{\{\s*' + re.escape(prefix) + r'_(\d+)\s*\}\}$')
    slots = []
    for p in doc.paragraphs:
        merged_text = ''.join(r.text for r in p.runs).strip()
        m = pattern.match(merged_text)
        if m:
            slots.append((int(m.group(1)), p))
    if not slots:
        safe_log(f"ensure_numbered_slots: tidak menemukan slot '{prefix}_N' di template, dilewati.")
        return
    slots.sort(key=lambda t: t[0])
    max_num, last_para = slots[-1]
    if needed_count <= max_num:
        return

    anchor_el = last_para._element
    template_rPr = None
    if last_para.runs and last_para.runs[0]._element.find(qn('w:rPr')) is not None:
        template_rPr = last_para.runs[0]._element.find(qn('w:rPr'))

    for n in range(max_num + 1, needed_count + 1):
        new_el = copy.deepcopy(anchor_el)
        # Kosongkan semua run lama di klon, ganti dengan SATU run baru berisi
        # placeholder bernomor baru -- menghindari masalah placeholder yang
        # aslinya terpecah jadi beberapa run (mis. '{{' + 'pihak_terkait_5' + '}}').
        for r_el in new_el.findall(qn('w:r')):
            new_el.remove(r_el)
        new_r = OxmlElement('w:r')
        if template_rPr is not None:
            new_r.append(copy.deepcopy(template_rPr))
        new_t = OxmlElement('w:t')
        new_t.text = f"{{{{{prefix}_{n}}}}}"
        new_r.append(new_t)
        new_el.append(new_r)
        anchor_el.addnext(new_el)
        anchor_el = new_el


def _fill_table_rows_from_master(doc, header_keywords, rows_data, max_tables=None):
    def _find_tables_by_header(doc, header_keywords):
        """Mengembalikan SEMUA tabel yang cocok dengan header_keywords, bukan
        cuma yang pertama ditemukan. Penting untuk dokumen seperti Daftar
        Hadir yang punya lebih dari satu tabel tanda tangan (satu per
        halaman/blok) -- kalau cuma tabel pertama yang diisi, tabel
        berikutnya akan tertinggal dengan baris kosong bawaan template dan
        mendorong halaman jadi kosong."""
        found = []
        for tbl in doc.tables:
            if len(tbl.rows) < 2:
                continue
            header_texts = [c.text.strip().lower() for c in tbl.rows[0].cells]
            if all(any(kw.lower() in h for h in header_texts) for kw in header_keywords):
                found.append(tbl)
        return found

    def _set_cell_text_preserve_style(cell, lines):
        para = cell.paragraphs[0]
        font_name, font_size, font_bold = "Arial", Pt(11), False
        if para.runs:
            r0 = para.runs[0]
            if r0.font.name: font_name = r0.font.name
            if r0.font.size: font_size = r0.font.size
            font_bold = bool(r0.font.bold)

        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        for p in list(cell.paragraphs[1:]):
            p._element.getparent().remove(p._element)

        for idx, line in enumerate(lines):
            if idx > 0:
                run = para.add_run()
                run._element.append(OxmlElement('w:br'))
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = font_size
            run.bold = font_bold

    if not rows_data:
        return
    tables = _find_tables_by_header(doc, header_keywords)
    if max_tables is not None:
        # Hanya isi N tabel pertama yang cocok (mis. tabel "keberangkatan").
        # Tabel sisanya (mis. tabel "tempat tugas/tujuan") SENGAJA dibiarkan
        # apa adanya sesuai master template -- harus tetap kosong karena
        # akan diisi tanda tangan manual oleh pihak tempat tugas, BUKAN
        # diisi otomatis oleh program.
        tables = tables[:max_tables]
    for tbl in tables:
        template_tr = tbl.rows[1]._tr
        parent = template_tr.getparent()
        new_trs = []
        for data_row in rows_data:
            new_tr = copy.deepcopy(template_tr)
            new_trs.append(new_tr)

        ref = template_tr
        for new_tr in new_trs:
            ref.addnext(new_tr)
            ref = new_tr
        parent.remove(template_tr)

        for new_tr, data_row in zip(new_trs, rows_data):
            cells_tc = new_tr.findall(qn('w:tc'))
            from docx.table import _Cell
            for tc, val in zip(cells_tc, data_row):
                cell = _Cell(tc, tbl)
                lines = val.split('\n') if isinstance(val, str) else [str(val)]
                _set_cell_text_preserve_style(cell, lines)

def _force_daftar_hadir_page_break(doc):
    """Memastikan SETIAP periode/tujuan daftar hadir selalu terdiri dari
    TEPAT 2 halaman: (1) format keberangkatan, (2) format tempat tugas.

    Master template (DAFTAR_HADIR_DPRD.docx) mengandalkan belasan paragraf
    kosong sebagai "pengganjal" supaya format ke-2 jatuh ke halaman
    berikutnya -- ini RAPUH karena tinggi baris tabel/jumlah pelaksana bisa
    berubah-ubah, sehingga kadang menyisakan halaman kosong, kadang malah
    format ke-2 ikut nempel di halaman pertama. Fungsi ini menggantinya
    dengan page-break EKSPLISIT yang hasilnya selalu pasti, lalu membuang
    semua paragraf kosong "pengganjal" yang sudah tidak diperlukan lagi
    (baik yang ada di antara format 1 dan format 2, maupun sisa di akhir
    dokumen) supaya tidak ada halaman kosong tambahan yang nyelip."""
    body = doc.element.body
    children = list(body)

    def _ptext(el):
        return ''.join(t.text or '' for t in el.findall('.//' + qn('w:t'))).strip()

    title_positions = [i for i, el in enumerate(children)
                        if el.tag == qn('w:p') and _ptext(el) == "DAFTAR HADIR"]

    if len(title_positions) >= 2:
        first_title_idx, second_title_idx = title_positions[0], title_positions[1]
        # Buang semua paragraf KOSONG di antara format 1 dan format 2 --
        # sebelumnya dipakai sebagai pengganjal halaman, sekarang tidak
        # diperlukan lagi karena page break sudah eksplisit.
        for el in children[first_title_idx + 1:second_title_idx]:
            if el.tag == qn('w:p') and not _ptext(el):
                el.getparent().remove(el)

        # Sisipkan page break eksplisit tepat sebelum paragraf judul format 2,
        # supaya format 2 SELALU mulai di halaman baru, terlepas dari berapa
        # banyak baris yang ada di tabel format 1.
        second_title_el = children[second_title_idx]
        pb_p = OxmlElement('w:p')
        pb_r = OxmlElement('w:r')
        pb_br = OxmlElement('w:br')
        pb_br.set(qn('w:type'), 'page')
        pb_r.append(pb_br)
        pb_p.append(pb_r)
        second_title_el.addprevious(pb_p)

    # Buang sisa paragraf kosong di akhir dokumen (setelah tabel terakhir)
    # supaya tidak menambah halaman kosong tambahan saat beberapa periode
    # digabungkan jadi satu dokumen.
    for el in reversed(list(body)):
        if el.tag == qn('w:sectPr'):
            continue
        if el.tag == qn('w:p') and not _ptext(el):
            el.getparent().remove(el)
        else:
            break

def cleanup_surat_tugas_biasa(doc, template_type):
    paragraphs = list(doc.paragraphs)
    to_remove = set()

    if template_type == 'dprd':
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            stripped = p.text.replace('\t', ' ').strip()
            if re.match(r'^Nama\s*:\s*$', stripped):
                to_remove.add(id(p))
                if i + 1 < len(paragraphs):
                    next_p = paragraphs[i + 1]
                    next_stripped = next_p.text.replace('\t', ' ').strip()
                    if re.match(r'^Jabatan\s*:\s*$', next_stripped):
                        to_remove.add(id(next_p))
                        i += 2
                        continue
            i += 1
    else:
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            stripped = p.text.replace('\t', ' ').strip()
            # Baris "Nama" di template ASN diawali nomor urut, contoh:
            # "1. Nama :", "2. Nama :", "3. Nama :" (bukan "Nama :" polos seperti di DPRD)
            if re.match(r'^\d*\.?\s*Nama\s*:\s*$', stripped):
                to_remove.add(id(p))
                j = i + 1
                while j < len(paragraphs):
                    ns = paragraphs[j].text.replace('\t', ' ').strip()
                    # Label lengkap di template ASN: "Pangkat / Gol. Ruang :", "N I P :", "Jabatan :"
                    if re.match(r'^(Pangkat(\s*/\s*Gol\.?\s*Ruang)?|N\s*I\s*P|NIP|Jabatan)\s*:\s*$', ns):
                        to_remove.add(id(paragraphs[j]))
                        j += 1
                    else:
                        break
                i = j
                continue
            i += 1

    for p in paragraphs:
        if id(p) in to_remove:
            p._element.getparent().remove(p._element)

def cleanup_skenario_paripurna(doc):
    """
    Hapus paragraf skenario kosong di undangan paripurna.
    Logika identik dengan cleanup_surat_tugas_biasa: setelah DocxTemplate
    merender variabel kosong (''), paragraf ListParagraph (skenario 2-7)
    yang tidak terpakai menjadi baris kosong berpenomoran — hapus semuanya.
    Skenario_1 ada di baris yang sama dengan 'Acara :' sehingga tidak dihapus.
    """
    for p in list(doc.paragraphs):
        # Hanya ListParagraph dengan numPr (skenario 2-7 yang berpenomoran otomatis)
        if p.style and p.style.name == 'List Paragraph':
            if p.text.strip() == '':
                p._element.getparent().remove(p._element)


def _force_paripurna_page_break(doc):
    """Memastikan SETIAP blok halaman Undangan Paripurna (1 blok per tujuan,
    template rapat_paripurna.docx berisi 8 blok identik) SELALU mulai persis
    di halaman baru, tidak peduli berapa baris skenario acara atau seberapa
    panjang teks isi undangan yang diketik pengguna.

    BUG YANG DIPERBAIKI ("teks naik turun antar halaman"): master template
    memisahkan tiap blok HANYA dengan ~10 paragraf kosong "pengganjal" di
    akhir tiap blok (mengandalkan tinggi baris alami supaya blok berikutnya
    jatuh ke halaman baru) -- persis pola rapuh yang sama yang sudah
    diperbaiki utk Daftar Hadir lewat _force_daftar_hadir_page_break. Kalau
    isi_undangan_paripurna panjang / skenario acara terisi banyak baris,
    kontennya melebihi sisa ruang di halaman lalu MENDORONG paragraf
    pengganjal (dan judul "Bitung, ..." blok berikutnya) naik/turun antar
    halaman secara tidak terduga. Solusinya: pola yang sama dengan Daftar
    Hadir & Undangan Rapat Biasa -- pasang page-break EKSPLISIT tepat
    sebelum judul "Bitung, ..." tiap blok (kecuali blok pertama), lalu buang
    semua paragraf kosong pengganjal yang sudah tidak diperlukan lagi.
    """
    body = doc.element.body

    def _ptext(el):
        return ''.join(t.text or '' for t in el.findall('.//' + qn('w:t'))).strip()

    titles = [el for el in body if el.tag == qn('w:p') and _ptext(el).startswith('Bitung,')]
    if len(titles) < 2:
        return

    for prev_el, next_el in zip(titles, titles[1:]):
        # Buang semua paragraf kosong "pengganjal" di antara akhir blok
        # sebelumnya dan judul blok berikutnya -- tidak diperlukan lagi
        # karena page break sekarang eksplisit.
        el = prev_el.getnext()
        while el is not None and el is not next_el:
            nxt = el.getnext()
            if el.tag == qn('w:p') and not _ptext(el):
                el.getparent().remove(el)
            el = nxt

        # Sisipkan page break eksplisit tepat sebelum judul blok berikutnya,
        # supaya blok itu SELALU mulai di halaman baru terlepas dari
        # panjang konten blok sebelumnya.
        pb_p = OxmlElement('w:p')
        pb_r = OxmlElement('w:r')
        pb_br = OxmlElement('w:br')
        pb_br.set(qn('w:type'), 'page')
        pb_r.append(pb_br)
        pb_p.append(pb_r)
        next_el.addprevious(pb_p)


def _remove_blank_pages(doc):
    """Pemeriksaan akhir generik: pindai seluruh dokumen yang sudah
    digabung, cari "halaman" (segmen di antara dua page-break) yang sama
    sekali tidak punya teks ataupun tabel, lalu hilangkan page-break
    pemicunya supaya halaman kosong itu lenyap. Ini jaring pengaman
    terakhir di atas page-break eksplisit yang sudah dipasang di
    _force_daftar_hadir_page_break / _combine_word_pages, untuk berjaga-jaga
    seandainya ada page-break ganda yang tidak sengaja menyisakan halaman
    kosong."""
    body = doc.element.body

    def _is_page_break_para(el):
        if el.tag != qn('w:p'):
            return False
        breaks = el.findall('.//' + qn('w:br'))
        has_page_br = any(b.get(qn('w:type')) == 'page' for b in breaks)
        text = ''.join(t.text or '' for t in el.findall('.//' + qn('w:t'))).strip()
        return has_page_br and not text

    children = [el for el in body if el.tag != qn('w:sectPr')]
    # Bagi dokumen jadi segmen-segmen "halaman" berdasarkan paragraf
    # page-break murni (tanpa teks lain di paragraf yang sama).
    segments, current = [], []
    for el in children:
        if _is_page_break_para(el):
            segments.append(current)
            current = []
        else:
            current.append(el)
    segments.append(current)

    def _segment_is_empty(seg):
        for el in seg:
            if el.tag == qn('w:tbl'):
                return False
            if el.tag == qn('w:p'):
                text = ''.join(t.text or '' for t in el.findall('.//' + qn('w:t'))).strip()
                if text:
                    return False
        return True

    # Halaman pertama (segments[0]) tidak punya page-break pemicu sendiri,
    # jadi dilewati. Untuk halaman ke-2 dst, kalau isinya kosong, hapus
    # page-break paragraf yang memicunya supaya halaman itu hilang.
    pb_paragraphs = [el for el in children if _is_page_break_para(el)]
    for seg, pb_el in zip(segments[1:], pb_paragraphs):
        if _segment_is_empty(seg):
            pb_el.getparent().remove(pb_el)

def build_halaman_tujuan_lain_template(source_path, cache_path, nomor_placeholder="nomor_surat_rapat_4"):
    """Membuat template satu-halaman dari 'master halaman ke-4' surat undangan
    rapat biasa (blok Yth. {{tujuan_surat_lainnya}}), untuk dipakai berulang
    kali membuat halaman tujuan-lain tambahan yang JUMLAHNYA BEBAS.

    Caranya: salin utuh dokumen sumber (supaya header/footer/logo/style/
    numbering ikut terbawa), lalu hapus semua elemen body SELAIN blok
    halaman-4 tsb (dicari otomatis lewat tabel yang memuat placeholder
    nomor_placeholder) dan sectPr penutup. Hasilnya di-cache ke cache_path
    supaya tidak perlu diekstrak ulang tiap kali generate.
    """
    if os.path.exists(cache_path) and os.path.getmtime(cache_path) >= os.path.getmtime(source_path):
        return cache_path

    doc = Document(source_path)
    body = doc.element.body
    children = list(body)

    idx_nomor_tbl = None
    for i, el in enumerate(children):
        if el.tag == qn('w:tbl') and nomor_placeholder in ''.join(el.itertext()):
            idx_nomor_tbl = i
            break
    if idx_nomor_tbl is None:
        raise ValueError(f"Tidak menemukan tabel nomor dengan placeholder {{{{{nomor_placeholder}}}}} di template.")

    sect_pr = children[-1] if children[-1].tag == qn('w:sectPr') else None
    keep_start = idx_nomor_tbl - 3  # baris tanggal + 2 paragraf kosong sebelum tabel nomor
    keep_end = len(children) - (2 if sect_pr is not None else 1)  # sampai sebelum sectPr

    for i, el in enumerate(children):
        if sect_pr is not None and el is sect_pr:
            continue
        if not (keep_start <= i <= keep_end):
            body.remove(el)

    os.makedirs(os.path.dirname(cache_path), exist_ok=True)
    doc.save(cache_path)
    return cache_path


def build_tujuan_richtext(names):
    """Membuat RichText multi-baris untuk field 'Yth. {{tujuan_surat_lainnya}}'
    dengan penomoran per tujuan (1. , 2. , dst) dan format font YANG SAMA
    PERSIS dengan placeholder aslinya di template (Arial, 12pt/sz=24, bold),
    supaya ukuran font tiap baris konsisten baik itu tujuan pertama di
    halaman ke-4 maupun tujuan-tujuan di halaman tambahan.
    """
    from docxtpl import RichText
    rt = RichText()
    for i, nm in enumerate(names):
        if i > 0:
            rt.xml += "<w:br/>"
        rt.add(f"{i + 1}. {nm.upper()}", bold=True, font="Arial", size=24)
    return rt


def _dedupe_shape_ids(body):
    """PENTING - perbaikan bug utama: setiap kali docxtpl me-render sebuah
    dokumen yang mengandung shape/text-box (mis. kotak tanda tangan
    "DEWAN PERWAKILAN RAKYAT DAERAH..." di template rapat_biasa.docx),
    ID internal shape tsb (<wp:docPr id="...">) DIRESET ulang oleh docxtpl
    mulai dari angka yang sama (mis. 1001) UNTUK SETIAP FILE yang dirender
    terpisah. Selama tiap halaman disimpan sebagai file .docx sendiri-
    sendiri itu tidak masalah -- tapi begitu beberapa file digabung jadi
    SATU dokumen oleh _combine_word_pages (persis kasus "Tambah Halaman
    Baru" pada Undangan Rapat Biasa, yang me-render halaman ke-4 + setiap
    halaman tambahan sebagai file docxtpl TERPISAH lalu menggabungkannya),
    beberapa shape di halaman berbeda berakhir dengan id YANG SAMA PERSIS
    dalam satu dokumen akhir.

    ID shape (wp:docPr id) wajib unik di seluruh dokumen menurut skema
    OOXML. Kalau ada duplikat, Microsoft Word (bukan LibreOffice -- itu
    kenapa bug ini baru kelihatan lewat jalur konversi Word/COM) menganggap
    dokumen itu "punya masalah konten" dan mencoba menampilkan dialog
    perbaikan otomatis. Karena Word dijalankan headless (Visible=False)
    lewat automation COM untuk live preview, dialog itu TIDAK PERNAH bisa
    diklik siapapun -- panggilan Documents.Open() jadi menggantung selamanya.
    Itulah sebabnya setiap kali preview dicoba lagi (mis. saat mengetik di
    halaman tambahan), instance Word BARU dibuka (karena instance lama
    dianggap mati/tidak merespons) sementara instance lama yang menggantung
    tadi tidak pernah benar-benar tertutup -- menumpuk terus di Task Manager
    sampai akhirnya Word bahkan gagal dibuka manual oleh pengguna.

    Perbaikan di sini: setelah semua halaman digabung jadi satu <w:body>,
    beri ULANG setiap <wp:docPr id="..."> nomor yang dijamin unik di
    seluruh dokumen akhir (juga <v:shape id="..."> dan o:spid="..." milik
    fallback VML shape yang sama, supaya kedua representoriginal tetap
    konsisten satu sama lain)."""
    docpr_elements = body.findall('.//' + qn('wp:docPr'))
    used_ids = set()
    next_id = 9000  # rentang tinggi supaya tidak bentrok dgn id asli manapun
    for el in docpr_elements:
        current = el.get('id')
        if current in used_ids or current is None:
            while str(next_id) in used_ids:
                next_id += 1
            el.set('id', str(next_id))
            used_ids.add(str(next_id))
            next_id += 1
        else:
            used_ids.add(current)

    # Fallback VML (mc:Fallback) shape id & o:spid -- dibuat unik juga
    # dengan pola sama, terpisah dari namespace id DrawingML di atas.
    vml_ns = 'urn:schemas-microsoft-com:vml'
    vml_shapes = body.findall('.//{%s}shape' % vml_ns)
    used_spids = set()
    next_spid = 9000
    for shp in vml_shapes:
        spid = shp.get('{urn:schemas-microsoft-com:office:office}spid')
        if spid and spid in used_spids:
            new_spid = f"_x0000_s{next_spid}"
            while new_spid in used_spids:
                next_spid += 1
                new_spid = f"_x0000_s{next_spid}"
            shp.set('{urn:schemas-microsoft-com:office:office}spid', new_spid)
            shp.set('id', new_spid)
            used_spids.add(new_spid)
            next_spid += 1
        elif spid:
            used_spids.add(spid)


def _combine_word_pages(files_list, out_path):
    if not files_list: return
    try:
        master = Document(files_list[0])
        body = master.element.body

        # PENTING: python-docx menyimpan properti halaman dokumen sebagai
        # <w:sectPr> yang WAJIB jadi elemen TERAKHIR di <w:body>. Memanggil
        # master.add_page_break() (API tingkat tinggi) menyisipkan paragraf
        # barunya tepat SEBELUM sectPr ini, sedangkan elemen dari subdoc
        # lain yang ditempel manual lewat body.append(...) selalu masuk ke
        # posisi PALING AKHIR (yaitu SETELAH sectPr). Kalau dua cara ini
        # dicampur, sectPr jadi nyangkut di TENGAH dokumen (struktur OOXML
        # tidak valid) dan page-break antar-periode bisa salah tempat atau
        # malah hilang -- inilah salah satu sumber halaman kosong/hilang
        # yang dilaporkan. Solusinya: lepas dulu sectPr aslinya, gabungkan
        # SEMUA elemen (termasuk page break) lewat append() biasa secara
        # berurutan, baru pasang lagi sectPr di posisi terakhir di akhir.
        master_sectpr = body.find(qn('w:sectPr'))
        if master_sectpr is not None:
            body.remove(master_sectpr)

        for f in files_list[1:]:
            pb_p = OxmlElement('w:p')
            pb_r = OxmlElement('w:r')
            pb_br = OxmlElement('w:br')
            pb_br.set(qn('w:type'), 'page')
            pb_r.append(pb_br)
            pb_p.append(pb_r)
            body.append(pb_p)

            subdoc = Document(f)
            for element in list(subdoc.element.body):
                if element.tag != qn('w:sectPr'):
                    body.append(element)

        if master_sectpr is not None:
            body.append(master_sectpr)

        if master.paragraphs:
            last_p = master.paragraphs[-1]
            if not last_p.text.strip():
                last_p._element.getparent().remove(last_p._element)
        _dedupe_shape_ids(body)
        _remove_blank_pages(master)
        master.save(out_path)
    except Exception as e:
        # PENTING: JANGAN pakai print() di sini (lihat app/core/app_logging.py
        # untuk alasannya -- print() bisa melempar exception baru & membungkam
        # error asli di build .exe mode windowed). Dulu exception di sini
        # hanya di-print lalu FUNGSI INI KEMBALI DIAM-DIAM TANPA MENULIS
        # out_path SAMA SEKALI. Akibatnya kode pemanggil (preview & cetak
        # dokumen) melanjutkan proses seolah berhasil padahal file belum
        # dibuat/masih file lama -- baru gagal belakangan di tahap konversi
        # PDF dengan pesan yang menyesatkan ("Gagal konversi ke PDF") karena
        # sumber masalah sebenarnya (gagal menggabungkan halaman) sudah
        # tertutup. Sekarang: catat ke log lalu lempar ulang exception-nya
        # supaya pemanggil tahu proses ini benar-benar gagal dan bisa
        # menampilkan pesan error yang sebenarnya ke pengguna.
        safe_log(f"Gagal menggabungkan halaman ({out_path}): {e}")
        raise

