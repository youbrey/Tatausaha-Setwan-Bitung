"""
Window utama aplikasi (SIPSApp): seluruh tata letak panel (sidebar, form
input, checklist personel, live preview) serta orkestrasi pemanggilan
generator dokumen di app.core dan repository data di app.data.

NB: Kelas ini sengaja masih besar karena widget Tkinter/CustomTkinter saling
terikat erat lewat banyak atribut `self.*` (form input, state checklist,
preview). Pemisahan lebih lanjut per "view" (perjalanan dinas / undangan
paripurna / preview panel) ke file masing-masing adalah langkah refactor
lanjutan yang direkomendasikan -- lihat README.md bagian "Roadmap Refactor".
"""
import json
import os
import re
import shutil
import subprocess
import queue
import sys
import tempfile
import threading
import uuid
from datetime import datetime, timedelta

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, Twips
from docxtpl import DocxTemplate, RichText

try:
    from docx2pdf import convert as convert_to_pdf_word
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

try:
    import fitz  # PyMuPDF
    from PIL import Image
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    from tkcalendar import DateEntry
    HAS_TKCALENDAR = True
except ImportError:
    HAS_TKCALENDAR = False

from app.config.constants import (
    PREVIEW_TEMPLATES,
    AKD_LAINNYA_DISPLAY_NAMES,
    KATEGORI_AKD_INTI,
    KATEGORI_AKD_LAINNYA,
    KATEGORI_FRAKSI,
    KATEGORI_TENAGA_AHLI,
    KATEGORI_STRUKTUR,
    PELAKSANA_RAPAT_CUSTOM,
    PELAKSANA_RAPAT_OPTIONS,
    JENIS_RAPAT_OPTIONS,
    TUJUAN_SURAT_DPRD_MAP,
)
from app.config.theme import (
    CARD_CORNER_RADIUS,
    COLOR_BODY_BG,
    COLOR_BORDER,
    COLOR_CARD_BG,
    COLOR_DANGER,
    COLOR_DANGER_HOVER,
    COLOR_GREY_DARK,
    COLOR_GREY_DARK_HOVER,
    COLOR_INFO,
    COLOR_INFO_HOVER,
    COLOR_PRIMARY,
    COLOR_PRIMARY_HOVER,
    COLOR_SECONDARY,
    COLOR_SECONDARY_HOVER,
    COLOR_SIDEBAR_BG,
    COLOR_SOFT_BG,
    COLOR_SUCCESS,
    COLOR_SUCCESS_HOVER,
    COLOR_TEXT_BODY,
    COLOR_TEXT_DARK,
    COLOR_TEXT_MUTED,
    COLOR_WARNING,
    COLOR_WARNING_HOVER,
    SIDEBAR_ITEM_PADY,
    font as themed_font,
)
from app.config.settings import (
    APP_GEOMETRY,
    APP_ICON_PATH,
    APP_MIN_SIZE,
    APP_TITLE,
    APP_VERSION,
    TEMPLATE_DAFTAR_HADIR,
    TEMPLATE_NASKAH_DINAS_RAPAT,
    TEMPLATE_DAFTAR_HADIR_RAPAT_MAP,
    TEMPLATE_DAFTAR_HADIR_PARIPURNA,
    TEMPLATE_DAFTAR_HADIR_PIHAK_TERKAIT,
    TEMPLATE_DAFTAR_HADIR_SEKRETARIAT,
    TEMPLATE_DAFTAR_HADIR_TAF,
    TEMPLATE_PARIPURNA,
    TEMPLATE_RAPAT_BIASA,
    TEMPLATE_PEMBERITAHUAN,
    TEMPLATE_SPD_BELAKANG,
    TEMPLATE_SPD_DEPAN,
    TEMPLATE_ST_ASN_BIASA,
    TEMPLATE_ST_ASN_TABEL,
    TEMPLATE_ST_DPRD_BIASA,
    TEMPLATE_ST_DPRD_TABEL,
)
from app.core.docx_utils import (
    _combine_word_pages,
    _fill_table_rows_from_master,
    _force_daftar_hadir_page_break,
    _force_paripurna_page_break,
    build_halaman_tujuan_lain_template,
    build_tujuan_richtext,
    cleanup_skenario_paripurna,
    ensure_numbered_slots,
)
from app.core.document_generators import (
    buat_surat_pemberitahuan_multi,
    buat_surat_tugas_asn,
    buat_surat_tugas_dprd,
)
from app.core.sppd_generators import _build_person_sppd_context, buat_sppd_asn, buat_sppd_dprd
from app.core.app_logging import safe_log
from app.core.naskah_dinas_generator import generate_naskah_dinas
from app.core.daftar_hadir_rapat_generator import generate_daftar_hadir_rapat
from app.core.text_utils import (
    detect_zona_waktu,
    extract_city_name,
    format_jabatan_penandatanganan,
    format_rentang_tanggal,
    generate_periods,
    increment_nomor,
    increment_nomor_paripurna,
    is_in_sulawesi_utara,
    is_plain_region_name,
    slugify_filename,
    strip_jenis_perjalanan_prefix,
)
from app.data.database_repository import (
    load_database as repo_load_database,
    read_dprd_asn_from_file,
    save_database as repo_save_database,
    normalize_keys as repo_normalize_keys,
)
from app.data.history_repository import (
    load_history as repo_load_history,
    save_history as repo_save_history,
    cari_judul_serupa,
    nomor_surat_sudah_dipakai,
)
from app.ui.components.toast import ToastManager
from app.ui.components.step_lock import StepLockManager
from app.ui.riwayat_window import RiwayatSuratView
from app.ui.dashboard_window import DashboardView
from app.ui.judul_perjadin_window import JudulPerjadinView


class SIPSApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(APP_TITLE)
        self.geometry(APP_GEOMETRY)
        self.minsize(*APP_MIN_SIZE)
        self.configure(fg_color=COLOR_BODY_BG)
        try:
            self.iconbitmap(APP_ICON_PATH)
        except Exception:
            pass

        self.db_dprd = []
        self.db_asn = []
        self.dprd_vars = {}
        self.asn_vars = {}
        self.pelaksana_vars = {}
        self.pendamping_vars = {}
        self.history_data = {"perjalanan_dinas": {}, "undangan_rapat": {}}
        self._editing_perjalanan_key = None
        self.mode = "dprd"
        self.current_view = "perjalanan_dinas"
        self.active_categories = {}

        self.database_tujuan = [
            "Kota Manado", "Kota Bitung", "Kota Tomohon", "Kota Kotamobagu",
            "Kabupaten Minahasa", "Kabupaten Minahasa Utara", "Kabupaten Minahasa Selatan",
            "Kabupaten Minahasa Tenggara", "Kabupaten Bolaang Mongondow",
            "Kabupaten Bolaang Mongondow Utara", "Kabupaten Bolaang Mongondow Selatan",
            "Kabupaten Bolaang Mongondow Timur", "Kabupaten Kepulauan Sangihe",
            "Kabupaten Kepulauan Talaud", "Kabupaten Kepulauan Sitaro",
            "DKI Jakarta", "Kota Surabaya", "Kota Bandung", "Kota Medan",
            "Kota Semarang", "Kota Makassar", "Kota Palembang", "Kota Tangerang",
            "Kota Tangerang Selatan", "Kota Bekasi", "Kota Depok", "Kota Yogyakarta",
            "Kota Surakarta (Solo)", "Kota Balikpapan", "Kota Samarinda",
            "Kota Banjarmasin", "Kota Pontianak", "Kota Denpasar", "Kota Mataram",
            "Kota Kupang", "Kota Ambon", "Kota Jayapura", "Kota Sorong",
            "Kota Palu", "Kota Kendari", "Kota Gorontalo", "Kota Palangkaraya",
            "Kota Tarakan", "Kota Banda Aceh", "Kota Padang", "Kota Pekanbaru",
            "Kota Jambi", "Kota Bengkulu", "Kota Bandar Lampung", "Kota Pangkalpinang",
            "Kota Tanjungpinang"
        ]
        self.tujuan_terpilih = []

        # Toast notifikasi non-blocking (dipakai a.l. oleh StepLockManager)
        # dan pengunci urutan pengisian formulir khusus Perjalanan Dinas:
        # pelaksana wajib dipilih dulu, lalu bagian-bagian formulir lain
        # wajib diisi berurutan dari atas ke bawah tanpa boleh dilompati.
        self.toast = ToastManager(self)
        self.step_lock = StepLockManager(self.toast.show)

        self.preview_dir = tempfile.mkdtemp(prefix="sips_preview_")
        self._preview_after_id = None
        self._judul_dupe_after_id = None
        self._preview_busy = False
        self._preview_pending = False
        self._preview_lock = threading.Lock()
        self._preview_ctk_image = None
        # Instance Word COM yang dibiarkan tetap terbuka (bukan dibuka-tutup
        # setiap refresh) supaya konversi docx->pdf untuk Live Preview jauh
        # lebih cepat -- lihat _get_persistent_word_app().
        #
        # PENTING (perbaikan bug "proses Microsoft Word menumpuk di Task
        # Manager"): versi sebelumnya membuat THREAD PYTHON BARU di setiap
        # panggilan konversi (satu per refresh preview), dan tiap thread
        # baru itu mem-panggil pythoncom.CoInitialize() sendiri -- artinya
        # tiap thread punya APARTEMEN COM SENDIRI. Word COM (STA) TIDAK
        # BOLEH dipanggil dari apartemen/thread yang berbeda dari yang
        # membuatnya. Begitu thread ke-2 mencoba "ping" (word.Visible) pada
        # instance Word yang dibuat oleh thread ke-1, panggilan itu gagal
        # dengan error lintas-apartemen -- kode lama menganggap itu berarti
        # "instance lama mati", lalu membuat instance Word BARU lewat
        # DispatchEx dari thread ke-2. Instance LAMA (milik thread ke-1)
        # tidak pernah di-Quit() karena referensinya sudah ditimpa -- jadi
        # jadi proses winword.exe "hantu" yang tidak pernah tertutup. Kalau
        # ini terjadi berkali-kali (preview refresh terpicu berkali-kali
        # per menit saat mengetik), puluhan proses Word menumpuk persis
        # seperti yang dilaporkan.
        #
        # Perbaikannya: SATU worker thread yang hidup selama aplikasi
        # berjalan (dibuat sekali, lihat _get_word_worker_thread), semua
        # operasi Word COM SELALU dikerjakan di thread itu saja lewat
        # antrian tugas -- jadi hanya ada SATU apartemen COM, konsisten.
        self._word_app = None
        self._word_pid = None
        self._word_lock = threading.Lock()
        self._word_worker_thread = None
        self._word_task_queue = None

        self.load_database()
        self.load_history()
        self.setup_ui()
        self.calculate_duration()

        self.protocol("WM_DELETE_WINDOW", self.on_close)
        self.after(600, lambda: self.schedule_preview_refresh(immediate=True))

    def on_close(self):
        try:
            shutil.rmtree(self.preview_dir, ignore_errors=True)
        except Exception:
            pass
        self._quit_persistent_word_app()
        if self._word_task_queue is not None:
            try: self._word_task_queue.put(None)  # sinyal worker thread berhenti
            except Exception: pass
        self.destroy()

    def terbilang(self, n):
        satuan = ["", "Satu", "Dua", "Tiga", "Empat", "Lima", "Enam", "Tujuh", "Delapan", "Sembilan", "Sepuluh", "Sebelas"]
        if n < 12: return satuan[n]
        elif n < 20: return self.terbilang(n - 10) + " Belas"
        elif n < 100: return self.terbilang(n // 10) + " Puluh " + (satuan[n % 10] if n % 10 != 0 else "")
        return str(n)

    def format_indonesian_date(self, date_obj):
        if not date_obj: return ""
        months = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
                  "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
        return f"{date_obj.day} {months[date_obj.month]} {date_obj.year}"

    # ------------------------------------------------------------------
    # PENAMAAN FILE OTOMATIS BERDASARKAN KATEGORI/TUJUAN/TANGGAL
    # ------------------------------------------------------------------
    _BULAN_SLUG = ["", "januari", "februari", "maret", "april", "mei", "juni",
                   "juli", "agustus", "september", "oktober", "november", "desember"]
    _HARI_SLUG = ["senin", "selasa", "rabu", "kamis", "jumat", "sabtu", "minggu"]

    def _hari_tanggal_slug(self, date_obj):
        """'rabu-10-juli' dari objek date -- TANPA tahun, sesuai contoh
        yang diminta ('undangan-paripurna-rabu-10-juli')."""
        if not date_obj:
            return "tanggal"
        hari = self._HARI_SLUG[date_obj.weekday()]
        bulan = self._BULAN_SLUG[date_obj.month]
        return f"{hari}-{date_obj.day}-{bulan}"

    def _slug_kategori_pelaksana(self):
        """Slug kategori pelaksana tugas perjalanan dinas yang aktif
        (mis. 'Komisi III' -> 'komisi-III'), dipakai untuk penamaan file
        Surat Tugas / Daftar Hadir / SPD dst."""
        cats = [c for c, v in getattr(self, "active_categories", {}).items()
                if v and c not in ("Pendamping ASN", "Pelaksana ASN")]
        if not cats:
            cats = [c for c, v in getattr(self, "active_categories", {}).items() if v] or ["umum"]
        return "-".join(slugify_filename(c) for c in cats)

    def _slug_tujuan_pertama(self):
        if not getattr(self, "tujuan_terpilih", None):
            return "tujuan"
        return slugify_filename(self.tujuan_terpilih[0])

    def _nama_file_perjalanan(self, jenis_label):
        """mis. 'surat-tugas-dprd-komisi-III-bandung'."""
        mode_tag = "dprd" if self.mode == "dprd" else "setwan"
        return f"{jenis_label}-{mode_tag}-{self._slug_kategori_pelaksana()}-{self._slug_tujuan_pertama()}"

    def _nama_file_undangan_paripurna(self):
        """mis. 'undangan-paripurna-rabu-10-juli'."""
        date_obj = None
        try:
            if HAS_TKCALENDAR and "tanggal_surat" in self.undangan_inputs:
                date_obj = self.undangan_inputs["tanggal_surat"].get_date()
        except Exception:
            date_obj = None
        return f"undangan-paripurna-{self._hari_tanggal_slug(date_obj)}"

    def _nama_file_undangan_biasa(self, ctx):
        """mis. 'rapat-kerja-komisi-III-rabu-10-juli'."""
        jenis_rapat = self.undangan_inputs.get("jenis_rapat")
        jenis_slug = slugify_filename(jenis_rapat.get()) if jenis_rapat else "rapat"
        pelaksana_rapat = self.undangan_inputs.get("pelaksana_rapat")
        kategori_slug = slugify_filename(pelaksana_rapat.get()) if pelaksana_rapat else ""
        date_obj = None
        try:
            if HAS_TKCALENDAR and "tanggal_rapat" in self.undangan_inputs:
                date_obj = self.undangan_inputs["tanggal_rapat"].get_date()
        except Exception:
            date_obj = None
        parts = [p for p in [jenis_slug, kategori_slug, self._hari_tanggal_slug(date_obj)] if p]
        return "-".join(parts)


    def normalize_keys(self, data_list):
        return repo_normalize_keys(data_list)

    # ------------------------------------------------------------------
    # DATABASE & UI SETUP
    # (Logika baca/tulis file sesungguhnya ada di app.data.database_repository
    # dan app.data.history_repository -- method di sini hanya menyimpan hasil
    # ke state widget/aplikasi.)
    # ------------------------------------------------------------------
    def load_database(self):
        self.db_dprd, self.db_asn = repo_load_database()

    def save_database(self):
        repo_save_database(self.db_dprd, self.db_asn)

    def show_network_settings_dialog(self):
        """Dialog kecil untuk mengatur folder BERSAMA (network share) supaya
        SIPS di komputer ini terhubung dengan komputer lain. Pengaturan ini
        LOKAL untuk komputer ini saja (lihat app/config/network_config.py)
        -- harus diisi satu per satu di tiap komputer dengan path yang
        sama-sama menunjuk ke satu folder jaringan.

        Sengaja tidak mengubah DATA_FILE/HISTORY_FILE/ACCOUNTS_FILE yang
        sudah dipakai (kalkulasi itu terjadi sekali saat aplikasi start,
        lihat app/config/settings.py) supaya tidak ada file yang setengah
        pindah folder saat aplikasi sedang berjalan -- karena itu, setelah
        menyimpan pengaturan di sini pengguna diminta RESTART aplikasi.
        """
        from app.config.network_config import get_shared_dir, set_shared_dir, is_shared_dir_reachable
        from app.config.settings import IS_NETWORK_MODE

        win = ctk.CTkToplevel(self)
        win.title("Pengaturan Jaringan (Multi-Komputer)")
        win.geometry("560x360")
        win.transient(self)
        win.grab_set()

        ctk.CTkLabel(
            win, text="🌐 Pengaturan Jaringan (Multi-Komputer)",
            font=themed_font(16, "bold"),
        ).pack(anchor="w", padx=20, pady=(18, 6))

        status_saat_ini = (
            f"Status saat ini: {'✅ TERSAMBUNG ke folder bersama' if IS_NETWORK_MODE else '⚪ Mode LOKAL (belum tersambung)'}"
        )
        ctk.CTkLabel(win, text=status_saat_ini, font=themed_font(12), text_color=COLOR_TEXT_BODY).pack(
            anchor="w", padx=20, pady=(0, 12)
        )

        info = (
            "Supaya data surat, riwayat, database pelaksana, dan akun login "
            "SALING TERHUBUNG di beberapa komputer (tanpa internet), arahkan "
            "SEMUA komputer ke SATU folder yang sama di jaringan lokal "
            "(mis. folder yang di-share lewat Windows File Sharing).\n\n"
            "Contoh isian:  \\\\KOMPUTER-TU\\SIPS_DATA   atau   Z:\\SIPS_DATA"
        )
        ctk.CTkLabel(
            win, text=info, font=themed_font(11), text_color=COLOR_TEXT_BODY,
            wraplength=510, justify="left",
        ).pack(anchor="w", padx=20, pady=(0, 12))

        ent_path = ctk.CTkEntry(win, placeholder_text="Path folder bersama...", height=36)
        ent_path.pack(fill="x", padx=20, pady=(0, 6))
        existing = get_shared_dir()
        if existing:
            ent_path.insert(0, existing)

        lbl_hasil = ctk.CTkLabel(win, text="", font=themed_font(11, "bold"))
        lbl_hasil.pack(anchor="w", padx=20, pady=(4, 0))

        def _cek_folder():
            path = ent_path.get().strip()
            if not path:
                lbl_hasil.configure(text="Isi dulu path folder-nya.", text_color=COLOR_WARNING)
                return
            if is_shared_dir_reachable(path):
                lbl_hasil.configure(text="✅ Folder bisa diakses (baca & tulis).", text_color=COLOR_SUCCESS)
            else:
                lbl_hasil.configure(
                    text="⛔ Folder TIDAK bisa diakses. Pastikan sudah di-share & komputer ini terhubung ke jaringannya.",
                    text_color="#DC2626",
                )

        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(fill="x", padx=20, pady=(14, 6))
        ctk.CTkButton(btn_row, text="Cek Folder", command=_cek_folder, width=120).pack(side="left")

        def _simpan():
            path = ent_path.get().strip()
            set_shared_dir(path, enabled=True)
            messagebox.showinfo(
                "Tersimpan",
                "Pengaturan jaringan disimpan untuk KOMPUTER INI.\n\n"
                "Tutup dan buka ulang aplikasi supaya perubahan berlaku "
                "(data surat/riwayat/database akan mulai dibaca-tulis dari "
                "folder bersama tsb).",
                parent=win,
            )
            win.destroy()

        def _lepas():
            set_shared_dir("", enabled=False)
            messagebox.showinfo(
                "Tersimpan",
                "Komputer ini dikembalikan ke mode LOKAL (tidak tersambung "
                "ke komputer lain). Restart aplikasi untuk menerapkan.",
                parent=win,
            )
            win.destroy()

        ctk.CTkButton(
            btn_row, text="💾 Simpan & Sambungkan", command=_simpan,
            fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER, width=180,
        ).pack(side="right")
        ctk.CTkButton(
            btn_row, text="Lepas Sambungan (Lokal)", command=_lepas,
            fg_color=COLOR_GREY_DARK, hover_color=COLOR_GREY_DARK_HOVER, width=180,
        ).pack(side="right", padx=(0, 8))

    def import_excel_database(self):
        file_path = filedialog.askopenfilename(
            title="Pilih File Database Excel/CSV",
            filetypes=[("Excel & CSV files", "*.xlsx *.xls *.csv")]
        )
        if not file_path: return

        progress = self._show_progress_dialog("Import Database", "Membaca file database...")

        def worker():
            try:
                new_dprd, new_asn = read_dprd_asn_from_file(file_path)
            except Exception as e:
                self.after(0, lambda: self._finish_import_database(progress, error=str(e)))
                return
            self.after(0, lambda: self._finish_import_database(progress, new_dprd=new_dprd, new_asn=new_asn))

        threading.Thread(target=worker, daemon=True).start()

    def _finish_import_database(self, dlg, new_dprd=None, new_asn=None, error=None):
        try:
            dlg.grab_release()
            dlg.destroy()
        except Exception:
            pass

        if error:
            messagebox.showerror("Import Gagal", f"Terjadi kesalahan:\n{error}")
            return

        try:
            if new_dprd: self.db_dprd = self.normalize_keys(new_dprd)
            if new_asn: self.db_asn = self.normalize_keys(new_asn)
            self.save_database()
            self.dprd_vars = {}
            self.asn_vars = {}
            self.refresh_personnel_list()
            self.refresh_signer_dropdowns()
            self.schedule_preview_refresh(immediate=True)
            messagebox.showinfo("Import Berhasil", "Database berhasil diperbarui.")
        except Exception as e:
            messagebox.showerror("Import Gagal", f"Terjadi kesalahan:\n{e}")

    def load_history(self):
        self.history_data = repo_load_history()

    def save_history(self):
        repo_save_history(self.history_data)

    def _schedule_cek_judul_duplikat(self, delay_ms=700):
        """Cek apakah Materi/Agenda Kegiatan yang sedang diketik SAMA PERSIS
        dengan yang sudah pernah dibuat sebelumnya (termasuk yang dibuat di
        komputer lain, karena riwayat dibaca dari file bersama), lalu
        tampilkan toast peringatan di TENGAH LAYAR kalau memang duplikat.

        Sengaja di-debounce (ditunda `delay_ms`, dibatalkan lagi kalau
        pengguna masih mengetik) supaya TIDAK mengecek riwayat pada SETIAP
        ketukan tombol -- kalau folder riwayat ada di jaringan (mode
        multi-komputer), membaca file di tiap keystroke bisa terasa berat/
        lambat, apalagi di komputer spesifikasi rendah."""
        if getattr(self, "_judul_dupe_after_id", None) is not None:
            try:
                self.after_cancel(self._judul_dupe_after_id)
            except Exception:
                pass
        self._judul_dupe_after_id = self.after(delay_ms, self._cek_judul_duplikat_now)

    def _cek_judul_duplikat_now(self):
        self._judul_dupe_after_id = None
        try:
            materi = self.txt_materi_st.get("1.0", tk.END).strip()
        except Exception:
            return
        # Kecualikan nomor surat yang SEDANG terisi di formulir -- kalau
        # pengguna sedang mengedit/merevisi surat lama, judulnya sendiri
        # tidak akan ditandai sebagai duplikat. Untuk surat baru (nomor
        # belum pernah ada di riwayat), pengecualian ini tidak berpengaruh.
        # Dikecualikan dari pengecekan duplikat: KEY internal record yang
        # sedang diedit (BUKAN teks nomor surat -- lihat catatan panjang di
        # build_context soal kenapa key riwayat sekarang UUID, bukan nomor
        # surat apa adanya). Untuk surat baru ini bernilai None, jadi tidak
        # mengecualikan apa pun.
        nomor_induk_sedang_diedit = getattr(self, "_editing_perjalanan_key", None)
        try:
            duplikat = cari_judul_serupa(materi, kecuali_nomor_induk=nomor_induk_sedang_diedit)
        except Exception:
            return
        if duplikat:
            pembuat = duplikat.get("dibuat_oleh", "-")
            tanggal = duplikat.get("tanggal_surat", "-")
            self.toast.show_center(
                f"Judul/Materi ini SUDAH PERNAH dibuat sebelumnya "
                f"(surat tanggal {tanggal}, oleh {pembuat}). Periksa kembali "
                f"sebelum melanjutkan.",
                kind="warning",
            )

    def _refresh_riwayat_window_if_open(self):
        """Nama method dipertahankan (dipanggil dari banyak tempat setelah
        surat baru disimpan) walau sekarang Riwayat Surat & Dashboard bukan
        lagi jendela terpisah -- cukup hitung ulang datanya di tempat,
        murah karena cuma agregasi dict, tidak masalah dipanggil tiap kali
        surat disimpan meskipun sedang tidak terlihat."""
        if hasattr(self, "riwayat_view"):
            self.riwayat_view.refresh_data()
        if hasattr(self, "dashboard_view"):
            self.dashboard_view.refresh_data()
        if hasattr(self, "judul_perjadin_view"):
            self.judul_perjadin_view.refresh_data()

    def _set_nav_button_states(self, active):
        """Nonaktifkan tombol sidebar utk panel yang sedang tampil, aktifkan
        sisanya -- pola yang sama dipakai utk Undangan Paripurna/Biasa."""
        mapping = {
            "undangan_paripurna": self.btn_undangan_paripurna,
            "undangan_biasa": self.btn_undangan_biasa,
            "dashboard": self.btn_dashboard,
            "riwayat_surat": self.btn_riwayat_surat,
            "judul_perjadin": self.btn_judul_perjadin,
        }
        for name, btn in mapping.items():
            btn.configure(state="disabled" if name == active else "normal")

    def show_judul_perjadin(self):
        """Tampilkan panel Judul Perjadin TEPAT DI SAMPING SIDEBAR, bukan
        jendela terpisah -- pola sama seperti show_dashboard/show_riwayat_surat."""
        self.current_view = "judul_perjadin"
        self.btn_back_to_perjalanan.grid()
        self._set_nav_button_states("judul_perjadin")
        self._switch_to_fullwidth_layout()
        self.dashboard_view.grid_remove()
        self.riwayat_view.grid_remove()
        self.judul_perjadin_view.grid(row=0, column=0, sticky="nsew")
        self.judul_perjadin_view.refresh_data()

    def show_dashboard(self):
        """Tampilkan panel Dashboard TEPAT DI SAMPING SIDEBAR (menggantikan
        area form + preview), bukan jendela terpisah."""
        self.current_view = "dashboard"
        self.btn_back_to_perjalanan.grid()
        self._set_nav_button_states("dashboard")
        self._switch_to_fullwidth_layout()
        self.riwayat_view.grid_remove()
        self.judul_perjadin_view.grid_remove()
        self.dashboard_view.grid(row=0, column=0, sticky="nsew")
        self.dashboard_view.refresh_data()

    def show_riwayat_surat(self):
        """Tampilkan panel Riwayat Surat TEPAT DI SAMPING SIDEBAR, bukan
        jendela terpisah."""
        self.current_view = "riwayat_surat"
        self.btn_back_to_perjalanan.grid()
        self._set_nav_button_states("riwayat_surat")
        self._switch_to_fullwidth_layout()
        self.dashboard_view.grid_remove()
        self.judul_perjadin_view.grid_remove()
        self.riwayat_view.grid(row=0, column=0, sticky="nsew")
        self.riwayat_view.refresh_data()

    def load_riwayat_perjalanan(self, nomor_surat):
        """Muat ulang 1 riwayat Surat Perjalanan Dinas ke formulir, dipanggil
        dari tombol Edit di jendela Riwayat Surat. `nomor_surat` di sini
        sebenarnya KUNCI internal riwayat (lihat catatan di build_context
        tentang kenapa kunci riwayat TIDAK BOLEH memakai nomor surat
        apa adanya)."""
        data = self.history_data.get("perjalanan_dinas", {}).get(nomor_surat)
        if data:
            if self.current_view != "perjalanan_dinas":
                self.show_perjalanan_dinas()
            if not hasattr(self, 'inputs') or not self.inputs:
                return
            # Tandai formulir sedang MENGEDIT record riwayat ini (bukan
            # membuat surat baru) supaya build_context() menimpa record yang
            # SAMA saat disimpan ulang, bukan membuat entri duplikat baru.
            self._editing_perjalanan_key = nomor_surat
            try:
                self.inputs["nomor_surat"].delete(0, tk.END)
                self.inputs["nomor_surat"].insert(0, data.get("nomor_surat", ""))
            except: pass
            try:
                self.inputs["nomor_surat_asn"].delete(0, tk.END)
                self.inputs["nomor_surat_asn"].insert(0, data.get("nomor_surat_asn", ""))
            except: pass
            try:
                self.inputs["nomor_pemberitahuan_dprd"].delete(0, tk.END)
                self.inputs["nomor_pemberitahuan_dprd"].insert(0, data.get("nomor_pemberitahuan_dprd", ""))
            except: pass
            try:
                self.inputs["nomor_pemberitahuan_asn"].delete(0, tk.END)
                self.inputs["nomor_pemberitahuan_asn"].insert(0, data.get("nomor_pemberitahuan_asn", ""))
            except: pass
            try:
                self.inputs["nomor_spd_dprd"].delete(0, tk.END)
                self.inputs["nomor_spd_dprd"].insert(0, data.get("nomor_spd_dprd", ""))
            except: pass
            try:
                self.inputs["nomor_spd_asn"].delete(0, tk.END)
                self.inputs["nomor_spd_asn"].insert(0, data.get("nomor_spd_asn", ""))
            except: pass
            try:
                self.txt_materi_st.delete("1.0", tk.END)
                self.txt_materi_st.insert("1.0", data.get("materi_tugas", ""))
            except: pass
            try:
                self.txt_materi_pb.delete("1.0", tk.END)
                self.txt_materi_pb.insert("1.0", data.get("materi_tugas_pb", ""))
            except: pass
            try:
                self.txt_dasar_dprd.delete("1.0", tk.END)
                self.txt_dasar_dprd.insert("1.0", data.get("dasar_surat_dprd", ""))
            except: pass
            try:
                self.txt_dasar_asn.delete("1.0", tk.END)
                self.txt_dasar_asn.insert("1.0", data.get("dasar_surat_asn", ""))
            except: pass
            try:
                self.combo_jenis.set(data.get("jenis_perjalanan", ""))
            except: pass

            tujuan_saved = data.get("tujuan_bertugas_list", [])
            if tujuan_saved:
                self.tujuan_terpilih = tujuan_saved
            else:
                old = data.get("tujuan_bertugas", "")
                self.tujuan_terpilih = [old] if old else []
            self.refresh_tujuan_list_ui()

            for var in self.dprd_vars.values(): var.set(False)
            for var in self.asn_vars.values(): var.set(False)
            for n in data.get("dprd_terpilih", []):
                if "||" in n:
                    nama, kategori = n.split("||", 1)
                    key = (nama, kategori)
                    if key in self.dprd_vars:
                        self.dprd_vars[key].set(True)
                else:
                    # Kompatibilitas mundur dgn riwayat lama (tersimpan nama saja,
                    # sebelum diperbaiki jadi kunci nama+kategori) -- centang semua
                    # baris kategori milik nama tsb yang sedang dirender.
                    for (var_nama, _var_kat), var in self.dprd_vars.items():
                        if var_nama == n:
                            var.set(True)
            for n in data.get("asn_terpilih", []):
                if n in self.asn_vars: self.asn_vars[n].set(True)

            self.schedule_preview_refresh(immediate=True)
            messagebox.showinfo("Riwayat Dimuat", "Formulir telah diisi dengan data surat sebelumnya.")
        else:
            messagebox.showwarning("Gagal", "Riwayat tidak ditemukan.")

    def load_riwayat_undangan(self, key):
        """Muat ulang 1 riwayat Surat Undangan Rapat (paripurna/biasa) ke
        formulir undangan yang sesuai, dipanggil dari tombol Edit di jendela
        Riwayat Surat."""
        data = self.history_data.get("undangan_rapat", {}).get(key)
        if not data:
            messagebox.showwarning("Gagal", "Riwayat tidak ditemukan.")
            return

        tipe = data.get("tipe", "biasa")
        if tipe == "paripurna":
            self.show_undangan_paripurna()
        else:
            self.show_undangan_biasa()

        if not hasattr(self, "undangan_inputs") or not self.undangan_inputs:
            return
        inp = self.undangan_inputs

        def _set_date(widget_key, iso_value):
            widget = inp.get(widget_key)
            if widget is None or not iso_value:
                return
            try:
                y, m, d = [int(x) for x in iso_value.split("-")]
                if HAS_TKCALENDAR:
                    widget.set_date(datetime(y, m, d).date())
                else:
                    widget.delete(0, tk.END)
                    widget.insert(0, iso_value)
            except Exception:
                pass

        try:
            inp["nomor_undangan"].delete(0, tk.END)
            inp["nomor_undangan"].insert(0, data.get("nomor_surat", ""))
        except Exception: pass
        _set_date("tanggal_surat", data.get("tanggal_surat_iso", ""))
        _set_date("tanggal_rapat", data.get("tanggal_rapat_iso", ""))
        self.update_hari_rapat()
        try:
            inp["jam_pelaksanaan"].delete(0, tk.END)
            inp["jam_pelaksanaan"].insert(0, data.get("jam_pelaksanaan", ""))
        except Exception: pass
        try:
            inp["isi_surat"].delete("1.0", tk.END)
            inp["isi_surat"].insert("1.0", data.get("isi_surat", ""))
        except Exception: pass
        try:
            if data.get("penandatanganan"):
                inp["penandatanganan"].set(data["penandatanganan"])
        except Exception: pass

        if tipe == "paripurna":
            try:
                if data.get("pakaian"):
                    inp["pakaian"].set(data["pakaian"])
            except Exception: pass
            try:
                for ent in list(inp["skenario"]):
                    ent.master.destroy()
                inp["skenario"] = []
                for teks in (data.get("skenario_list") or [""]):
                    self._add_skenario_row()
                    inp["skenario"][-1].insert(0, teks)
            except Exception: pass
        else:
            try:
                if data.get("pelaksana_rapat"):
                    inp["pelaksana_rapat"].set(data["pelaksana_rapat"])
                if data.get("jenis_rapat"):
                    inp["jenis_rapat"].set(data["jenis_rapat"])
            except Exception: pass
            try:
                for ent in list(inp["pihak_terkait"]):
                    ent.master.destroy()
                inp["pihak_terkait"] = []
                for teks in (data.get("pihak_terkait_list") or [""]):
                    self._add_pihak_terkait_row()
                    inp["pihak_terkait"][-1].insert(0, teks)
            except Exception: pass
            try:
                for g in list(inp["tujuan_halaman_groups"]):
                    g["frame"].destroy()
                inp["tujuan_halaman_groups"] = []
                for names in (data.get("tujuan_halaman_groups") or [[]]):
                    self._add_tujuan_halaman_group()
                    group = inp["tujuan_halaman_groups"][-1]
                    for ent in list(group["entries"]):
                        ent.master.destroy()
                    group["entries"] = []
                    for nama in (names or [""]):
                        self._add_tujuan_entry(group)
                        group["entries"][-1].insert(0, nama)
            except Exception: pass

        self.schedule_preview_refresh(immediate=True)
        messagebox.showinfo("Riwayat Dimuat", "Formulir undangan telah diisi dengan data surat sebelumnya.")

    def _record_riwayat_undangan(self, tipe):
        """Catat 1 entri riwayat surat undangan (paripurna/biasa) setiap kali
        surat berhasil dicetak. Sebelumnya surat undangan sama sekali tidak
        tercatat di riwayat -- ini menambahkannya, setara dengan yang sudah
        ada untuk Surat Perjalanan Dinas."""
        try:
            inp = self.undangan_inputs
            nomor = inp["nomor_undangan"].get().strip() or "-"

            if HAS_TKCALENDAR:
                tgl_surat_disp = self.format_indonesian_date(inp["tanggal_surat"].get_date())
                tgl_surat_iso = inp["tanggal_surat"].get_date().isoformat()
                tgl_rapat_disp = self.format_indonesian_date(inp["tanggal_rapat"].get_date())
                tgl_rapat_iso = inp["tanggal_rapat"].get_date().isoformat()
            else:
                tgl_surat_disp = tgl_surat_iso = inp["tanggal_surat"].get()
                tgl_rapat_disp = tgl_rapat_iso = inp["tanggal_rapat"].get()

            waktu_dibuat = datetime.now().strftime("%d-%m-%Y %H:%M")
            record = {
                "tipe": tipe,
                "nomor_surat": nomor,
                "tanggal_surat": tgl_surat_disp,
                "tanggal_surat_iso": tgl_surat_iso,
                "tanggal_rapat": tgl_rapat_disp,
                "tanggal_rapat_iso": tgl_rapat_iso,
                "hari_rapat": inp["hari_rapat"].get(),
                "jam_pelaksanaan": inp["jam_pelaksanaan"].get(),
                "isi_surat": inp["isi_surat"].get("1.0", tk.END).strip(),
                "penandatanganan": inp["penandatanganan"].get(),
                "dibuat_oleh": getattr(self, "_current_user", "-"),
                "tanggal_dibuat": waktu_dibuat,
            }
            if tipe == "paripurna":
                record["pakaian"] = inp["pakaian"].get()
                record["skenario_list"] = [e.get().strip() for e in inp["skenario"]]
            else:
                record["pelaksana_rapat"] = inp["pelaksana_rapat"].get()
                record["jenis_rapat"] = inp["jenis_rapat"].get()
                record["pihak_terkait_list"] = [e.get().strip() for e in inp["pihak_terkait"]]
                record["tujuan_halaman_groups"] = [
                    [e.get().strip() for e in g["entries"]] for g in inp["tujuan_halaman_groups"]
                ]

            key = f"{tipe}::{nomor}::{waktu_dibuat}"
            self.history_data.setdefault("undangan_rapat", {})[key] = record
            self.save_history()
            self._refresh_riwayat_window_if_open()
        except Exception:
            pass

    def setup_ui(self):
        self.grid_columnconfigure(0, weight=0, minsize=240)
        self.grid_columnconfigure(1, weight=1, minsize=420)
        self.grid_columnconfigure(2, weight=1, minsize=380)
        self.grid_columnconfigure(3, weight=2, minsize=480)
        self.grid_rowconfigure(0, weight=1)

        # SIDEBAR (gaya Adminator: latar putih, border kanan tipis, bukan panel gelap)
        self.sidebar_frame = ctk.CTkFrame(
            self, width=240, corner_radius=0,
            fg_color=COLOR_SIDEBAR_BG, border_width=0,
        )
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        # Baris 11 adalah baris KOSONG khusus untuk spacer yang melar mengisi
        # sisa ruang vertikal -- BUKAN baris yang juga berisi tombol, supaya
        # tombol "Undangan Biasa" tidak ikut tertarik melebar ke bawah dan
        # menyisakan jarak besar sebelum footer (bug tampilan sebelumnya).
        self.sidebar_frame.grid_rowconfigure(11, weight=1)

        self.logo_label = ctk.CTkLabel(
            self.sidebar_frame, text="Sekretariat DPRD Kota BITUNG",
            font=themed_font(16, "bold"), text_color=COLOR_TEXT_DARK,
            wraplength=200, justify="center",
        )
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 12))
        row1_frame = ctk.CTkFrame(self.sidebar_frame, fg_color="transparent")
        row1_frame.grid(row=1, column=0, padx=20, pady=(0, 10), sticky="ew")
        row1_frame.grid_columnconfigure(0, weight=1)
        row1_frame.grid_columnconfigure(1, weight=0)
        self.btn_import_db = ctk.CTkButton(
            row1_frame, text="📥 Import Database", command=self.import_excel_database,
            fg_color=COLOR_PRIMARY, hover_color=COLOR_PRIMARY_HOVER,
            font=themed_font(12), corner_radius=6, height=34,
        )
        self.btn_import_db.grid(row=0, column=0, sticky="ew")
        self.btn_network_settings = ctk.CTkButton(
            row1_frame, text="🌐", width=34, height=34, corner_radius=6,
            fg_color=COLOR_GREY_DARK, hover_color=COLOR_GREY_DARK_HOVER,
            font=themed_font(12), command=self.show_network_settings_dialog,
        )
        self.btn_network_settings.grid(row=0, column=1, padx=(6, 0))

        self.mode_selector = ctk.CTkSegmentedButton(
            self.sidebar_frame, values=["DPRD", "Setwan"], command=self.change_mode,
            fg_color=COLOR_SOFT_BG, selected_color=COLOR_PRIMARY,
            selected_hover_color=COLOR_PRIMARY_HOVER, unselected_color=COLOR_SOFT_BG,
            text_color=COLOR_TEXT_DARK, font=themed_font(12),
        )
        self.mode_selector.grid(row=2, column=0, padx=20, pady=(0, 10), sticky="ew")
        self.mode_selector.set("DPRD")

        self.btn_dashboard = ctk.CTkButton(
            self.sidebar_frame, text="📊 Dashboard", command=self.show_dashboard,
            fg_color=COLOR_INFO, hover_color=COLOR_INFO_HOVER,
            font=themed_font(12, "bold"), corner_radius=6, height=36,
        )
        self.btn_dashboard.grid(row=3, column=0, padx=20, pady=(8, 8), sticky="ew")

        self.btn_riwayat_surat = ctk.CTkButton(
            self.sidebar_frame, text="📜 Riwayat Surat", command=self.show_riwayat_surat,
            fg_color=COLOR_WARNING, hover_color=COLOR_WARNING_HOVER,
            text_color=COLOR_TEXT_DARK, font=themed_font(12, "bold"), corner_radius=6, height=36,
        )
        self.btn_riwayat_surat.grid(row=4, column=0, padx=20, pady=(0, 12), sticky="ew")

        self.btn_judul_perjadin = ctk.CTkButton(
            self.sidebar_frame, text="🗂️ Judul Perjadin", command=self.show_judul_perjadin,
            fg_color=COLOR_SECONDARY, hover_color=COLOR_SECONDARY_HOVER,
            font=themed_font(12, "bold"), corner_radius=6, height=36,
        )
        self.btn_judul_perjadin.grid(row=5, column=0, padx=20, pady=(0, 12), sticky="ew")

        self.btn_generate_main = ctk.CTkButton(
            self.sidebar_frame, text="⚡ CETAK SURAT & SPD", command=self.generate_documents_action,
            fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER, font=themed_font(14, "bold"),
            corner_radius=6, height=40,
        )
        self.btn_generate_main.grid(row=6, column=0, padx=20, pady=(0, 14), sticky="ew")

        lbl_undangan_title = ctk.CTkLabel(self.sidebar_frame, text="Kategori Surat Undangan:", font=themed_font(11, "bold"), text_color=COLOR_TEXT_BODY)
        lbl_undangan_title.grid(row=7, column=0, padx=20, pady=(0, 6), sticky="w")
        self.btn_undangan_paripurna = ctk.CTkButton(
            self.sidebar_frame, text="📨 Undangan Paripurna", command=self.show_undangan_paripurna,
            fg_color=COLOR_PRIMARY, hover_color=COLOR_PRIMARY_HOVER, font=themed_font(12),
            corner_radius=6, height=34,
        )
        self.btn_undangan_paripurna.grid(row=8, column=0, padx=20, pady=(0, SIDEBAR_ITEM_PADY), sticky="ew")
        self.btn_undangan_biasa = ctk.CTkButton(
            self.sidebar_frame, text="📋 Undangan Biasa", command=self.show_undangan_biasa,
            fg_color=COLOR_SECONDARY, hover_color=COLOR_SECONDARY_HOVER, font=themed_font(12),
            corner_radius=6, height=34,
        )
        self.btn_undangan_biasa.grid(row=9, column=0, padx=20, pady=(0, SIDEBAR_ITEM_PADY), sticky="ew")

        self.btn_back_to_perjalanan = ctk.CTkButton(
            self.sidebar_frame, text="← Kembali ke Perjalanan Dinas", command=self.show_perjalanan_dinas,
            fg_color=COLOR_GREY_DARK, hover_color=COLOR_GREY_DARK_HOVER, font=themed_font(12),
            corner_radius=6, height=32,
        )
        self.btn_back_to_perjalanan.grid(row=10, column=0, padx=20, pady=(2, 0), sticky="ew")
        self.btn_back_to_perjalanan.grid_remove()

        # row=11 sengaja dibiarkan kosong -- inilah baris spacer (weight=1 di
        # atas) yang menyerap sisa ruang, sehingga jarak antar tombol di atas
        # tetap rapat dan footer (versi, user info, Kelola Akun, Logout) tetap
        # menempel rapat satu sama lain di bagian bawah sidebar.

        self.lbl_credit = ctk.CTkLabel(
            self.sidebar_frame, text=f"{APP_VERSION} © DPRD Kota Bitung",
            font=themed_font(9), text_color=COLOR_TEXT_MUTED,
        )
        self.lbl_credit.grid(row=12, column=0, padx=20, pady=(0, 6), sticky="s")

        # PANEL TENGAH (kartu putih bergaya Adminator: corner radius + border tipis)
        self.middle_frame = ctk.CTkScrollableFrame(
            self, label_text="Data Perjalanan Dinas",
            fg_color=COLOR_CARD_BG, corner_radius=CARD_CORNER_RADIUS,
            border_width=1, border_color=COLOR_BORDER,
            label_font=themed_font(13, "bold"), label_text_color=COLOR_TEXT_DARK,
        )
        self.middle_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        self.inputs = {}
        self.mode_specific_widgets = {}
        self.undangan_inputs = {}

        # PANEL KANAN
        self.right_frame = ctk.CTkFrame(
            self, fg_color=COLOR_CARD_BG, corner_radius=CARD_CORNER_RADIUS,
            border_width=1, border_color=COLOR_BORDER,
        )
        self.right_frame.grid(row=0, column=2, padx=10, pady=10, sticky="nsew")
        self.right_frame.grid_columnconfigure(0, weight=1)
        self.right_frame.grid_rowconfigure(4, weight=1)

        lbl_cat_title = ctk.CTkLabel(self.right_frame, text="1. Filter Kategori Calon Pelaksana", font=themed_font(13, "bold"), text_color=COLOR_TEXT_DARK, anchor="w")
        lbl_cat_title.grid(row=0, column=0, padx=15, pady=(10, 5), sticky="w")
        self.category_check_frame = ctk.CTkFrame(self.right_frame, fg_color=COLOR_SOFT_BG, corner_radius=6, border_width=1, border_color=COLOR_BORDER)
        self.category_check_frame.grid(row=1, column=0, padx=15, pady=5, sticky="ew")
        
        lbl_person_title = ctk.CTkLabel(self.right_frame, text="2. Checklist Personel Pelaksana", font=themed_font(13, "bold"), text_color=COLOR_TEXT_DARK, anchor="w")
        lbl_person_title.grid(row=2, column=0, padx=15, pady=(15, 2), sticky="w")
        
        self.btn_action_frame = ctk.CTkFrame(self.right_frame, fg_color="transparent")
        self.btn_action_frame.grid(row=3, column=0, padx=15, pady=2, sticky="ew")
        self.btn_sel_all = ctk.CTkButton(self.btn_action_frame, text="Centang Semua Tampil", command=self.select_all_visible, height=24, fg_color=COLOR_PRIMARY, hover_color=COLOR_PRIMARY_HOVER, font=themed_font(11), corner_radius=5)
        self.btn_sel_all.pack(side="left", padx=2)
        self.btn_clear_all = ctk.CTkButton(self.btn_action_frame, text="Bersihkan", fg_color=COLOR_GREY_DARK, hover_color=COLOR_GREY_DARK_HOVER, command=self.clear_all_visible, height=24, font=themed_font(11), corner_radius=5)
        self.btn_clear_all.pack(side="left", padx=2)

        self.scroll_personnel = ctk.CTkScrollableFrame(self.right_frame, fg_color="transparent")
        self.scroll_personnel.grid(row=4, column=0, padx=15, pady=(5, 15), sticky="nsew")
        
        # PANEL PREVIEW
        self.preview_frame = ctk.CTkFrame(
            self, fg_color=COLOR_CARD_BG, corner_radius=CARD_CORNER_RADIUS,
            border_width=1, border_color=COLOR_BORDER,
        )
        self.preview_frame.grid(row=0, column=3, padx=10, pady=10, sticky="nsew")
        self.preview_frame.grid_columnconfigure(0, weight=1)
        self.preview_frame.grid_rowconfigure(3, weight=1)

        lbl_preview_title = ctk.CTkLabel(self.preview_frame, text="👁 Live Preview", font=themed_font(13, "bold"), text_color=COLOR_TEXT_DARK, anchor="w")
        lbl_preview_title.grid(row=0, column=0, padx=15, pady=(10, 5), sticky="w")
        
        self.preview_toolbar = ctk.CTkFrame(self.preview_frame, fg_color="transparent")
        self.preview_toolbar.grid(row=1, column=0, padx=15, pady=(0, 5), sticky="ew")
        self.preview_toolbar.grid_columnconfigure(0, weight=1)

        preview_choices = [item[0] for item in PREVIEW_TEMPLATES]
        self.combo_preview_jenis = ctk.CTkComboBox(self.preview_toolbar, values=preview_choices, command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_preview_jenis.set(preview_choices[0])
        self.combo_preview_jenis.grid(row=0, column=0, sticky="ew")
        
        self.btn_refresh_preview = ctk.CTkButton(self.preview_toolbar, text="🔄 Refresh", width=90, command=lambda: self.schedule_preview_refresh(immediate=True))
        self.btn_refresh_preview.grid(row=0, column=1, padx=(8, 0))

        self.preview_status_lbl = ctk.CTkLabel(self.preview_frame, text="Menyiapkan preview...", font=themed_font(11), text_color="gray", wraplength=420, justify="left")
        self.preview_status_lbl.grid(row=2, column=0, padx=15, pady=(0, 5), sticky="w")
        self.preview_canvas_frame = ctk.CTkScrollableFrame(self.preview_frame, fg_color="#E5E7EB")
        self.preview_canvas_frame.grid(row=3, column=0, padx=15, pady=(0, 15), sticky="nsew")
        self.preview_image_label = ctk.CTkLabel(self.preview_canvas_frame, text="")
        self.preview_image_label.pack(expand=True, pady=10)

        # PANEL LEBAR PENUH (menggantikan middle+right+preview sekaligus):
        # dipakai oleh Dashboard & Riwayat Surat, supaya keduanya tampil
        # TEPAT DI SAMPING SIDEBAR seperti mode lain -- bukan jendela
        # terpisah. Lihat _switch_to_fullwidth_layout().
        self.fullwidth_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.fullwidth_frame.grid_columnconfigure(0, weight=1)
        self.fullwidth_frame.grid_rowconfigure(0, weight=1)
        self.fullwidth_frame.grid_remove()  # tersembunyi sampai show_dashboard()/show_riwayat_surat()

        self.dashboard_view = DashboardView(self.fullwidth_frame, self)
        self.dashboard_view.grid(row=0, column=0, sticky="nsew")
        self.dashboard_view.grid_remove()

        self.riwayat_view = RiwayatSuratView(self.fullwidth_frame, self)
        self.judul_perjadin_view = JudulPerjadinView(self.fullwidth_frame, self)
        self.riwayat_view.grid(row=0, column=0, sticky="nsew")
        self.riwayat_view.grid_remove()

        self.show_perjalanan_dinas()

    def change_mode(self, mode):
        self.mode = "setwan" if mode == "Setwan" else "dprd"
        if self.current_view != "perjalanan_dinas":
            self.setup_category_checkboxes()
            self.refresh_personnel_list()
            return
        self.show_perjalanan_dinas()

    def setup_category_checkboxes(self):
        for widget in self.category_check_frame.winfo_children():
            widget.destroy()
        self.cat_chk_widgets = {}
        if self.mode == "dprd":
            default_active = {
                "Pimpinan DPRD": True, "Komisi I": True, "Komisi II": False, "Komisi III": False,
                "Pendamping ASN": True,
            }
            row = [0]  # pakai list supaya bisa diubah dalam closure

            def add_header(text, color="#1E3A8A", pady_top=8):
                lbl = ctk.CTkLabel(self.category_check_frame, text=text, font=themed_font(12, "bold"), text_color=color, anchor="w")
                lbl.grid(row=row[0], column=0, columnspan=2, padx=10, pady=(pady_top, 2), sticky="w")
                row[0] += 1

            def add_checkbox_row(cats, indent=10):
                for i in range(0, len(cats), 2):
                    pair = cats[i:i + 2]
                    for j, cat in enumerate(pair):
                        val = tk.BooleanVar(value=default_active.get(cat, False))
                        # 'cat' tetap dipakai APA ADANYA sebagai key pencocokan
                        # kategori di database Excel (mis. "Banggar"), hanya
                        # TEKS yang ditampilkan ke pengguna yang memakai nama
                        # lengkap (mis. "BADAN ANGGARAN").
                        display_text = AKD_LAINNYA_DISPLAY_NAMES.get(cat, cat).upper()
                        chk = ctk.CTkCheckBox(self.category_check_frame, text=display_text, variable=val, font=themed_font(11, "bold"), command=self.on_category_changed)
                        chk.grid(row=row[0], column=j, padx=(indent, 10), pady=6, sticky="w")
                        self.cat_chk_widgets[cat] = val
                    row[0] += 1

            # --- AKD ---
            add_header("AKD")
            add_checkbox_row(KATEGORI_AKD_INTI)

            # --- AKD Lainnya (sub-grup dari AKD) ---
            add_header("AKD Lainnya", color="#374151", pady_top=6)
            add_checkbox_row(KATEGORI_AKD_LAINNYA)

            # --- Pansus: nama bisa diketik ulang oleh pengguna ---
            pansus_frame = ctk.CTkFrame(self.category_check_frame, fg_color="transparent")
            pansus_frame.grid(row=row[0], column=0, columnspan=2, padx=(10, 10), pady=6, sticky="w")
            row[0] += 1
            if not hasattr(self, "pansus_name_var"):
                self.pansus_name_var = tk.StringVar(value="Pansus")
            self.pansus_active_var = tk.BooleanVar(value=False)
            chk_pansus = ctk.CTkCheckBox(pansus_frame, text="", variable=self.pansus_active_var, width=20, command=self.on_category_changed)
            chk_pansus.pack(side="left")
            ent_pansus = ctk.CTkEntry(pansus_frame, textvariable=self.pansus_name_var, width=180, font=themed_font(11, "bold"), placeholder_text="Nama Pansus (mis. Pansus RTRW)")
            ent_pansus.pack(side="left", padx=(4, 0))
            ent_pansus.bind("<KeyRelease>", lambda e: self.on_category_changed())
            self.cat_chk_widgets["Pansus"] = self.pansus_active_var

            # --- Fraksi ---
            add_header("Fraksi", pady_top=10)
            add_checkbox_row(KATEGORI_FRAKSI)

            # --- Tenaga Ahli ---
            add_header("Tenaga Ahli", pady_top=10)
            add_checkbox_row(KATEGORI_TENAGA_AHLI)

            # --- Pendamping ASN (tidak berubah) ---
            val = tk.BooleanVar(value=default_active.get("Pendamping ASN", True))
            chk = ctk.CTkCheckBox(self.category_check_frame, text="PENDAMPING ASN", variable=val, font=themed_font(11, "bold"), command=self.on_category_changed)
            chk.grid(row=row[0], column=0, columnspan=2, padx=10, pady=(10, 10), sticky="w")
            self.cat_chk_widgets["Pendamping ASN"] = val
        else:
            categories = ["Pelaksana ASN", "Pendamping ASN"]
            default_active = {"Pelaksana ASN": True, "Pendamping ASN": True}
            for idx, cat in enumerate(categories):
                val = tk.BooleanVar(value=default_active.get(cat, False))
                chk = ctk.CTkCheckBox(self.category_check_frame, text=cat, variable=val, font=themed_font(11, "bold"), command=self.on_category_changed)
                chk.grid(row=idx // 3, column=idx % 3, padx=10, pady=10, sticky="w")
                self.cat_chk_widgets[cat] = val
        self.active_categories = {cat: var.get() for cat, var in self.cat_chk_widgets.items()}

    def on_category_changed(self):
        for cat, var in self.cat_chk_widgets.items():
            self.active_categories[cat] = var.get()
        self.refresh_personnel_list()

    def refresh_personnel_list(self):
        for widget in self.scroll_personnel.winfo_children():
            widget.destroy()
        self.rendered_dprd_widgets, self.rendered_asn_widgets = {}, {}
        self.rendered_pelaksana_widgets, self.rendered_pendamping_widgets = {}, {}

        if self.mode == "dprd":
            pansus_name = getattr(self, "pansus_name_var", None)
            pansus_name = (pansus_name.get().strip() if pansus_name else "Pansus") or "Pansus"
            for p in self.db_dprd:
                cat = str(p.get("kategori") or "").strip()
                cat_lower = cat.lower()
                if cat in KATEGORI_AKD_INTI or cat in KATEGORI_AKD_LAINNYA or cat in KATEGORI_FRAKSI or cat in KATEGORI_TENAGA_AHLI:
                    group_cat = cat
                elif cat_lower.startswith("pansus") or (cat_lower and cat_lower == pansus_name.lower()):
                    group_cat = "Pansus"
                else:
                    # Kategori kosong/tidak dikenal (tidak cocok AKD/Fraksi/Pansus
                    # manapun) -> tidak ditampilkan di checklist manapun supaya
                    # data personel tetap rapi sesuai struktur kategori resmi.
                    group_cat = None
                if group_cat and self.active_categories.get(group_cat, False):
                    nama, jab = p.get('nama', ''), p.get('jabatan', '')
                    # Kunci per (nama, kategori) -- BUKAN nama saja. Satu orang bisa
                    # rangkap jabatan di banyak kategori (baris terpisah di database),
                    # jadi centang di satu kategori TIDAK BOLEH ikut mencentang baris
                    # kategori lain milik orang yang sama.
                    var_key = (nama, cat)
                    if var_key not in self.dprd_vars: self.dprd_vars[var_key] = tk.BooleanVar(value=False)
                    chk = ctk.CTkCheckBox(self.scroll_personnel, text=f"{nama} ({jab}) [{cat}]", variable=self.dprd_vars[var_key], command=lambda: self.schedule_preview_refresh(immediate=True))
                    chk.pack(fill="x", padx=10, pady=4, anchor="w")
                    self.rendered_dprd_widgets[var_key] = p

            if self.active_categories.get("Pendamping ASN", False):
                if self.rendered_dprd_widgets:
                    ctk.CTkFrame(self.scroll_personnel, height=2, fg_color="gray").pack(fill="x", padx=10, pady=10)
                for p in self.db_asn:
                    nama, nip, jab = p.get('nama', ''), p.get('nip', '-'), p.get('jabatan', '-')
                    if nama not in self.asn_vars: self.asn_vars[nama] = tk.BooleanVar(value=False)
                    chk = ctk.CTkCheckBox(self.scroll_personnel, text=f"{nama}\nNIP: {nip} | Jabatan: {jab}", variable=self.asn_vars[nama], command=lambda: self.schedule_preview_refresh(immediate=True))
                    chk.pack(fill="x", padx=10, pady=5, anchor="w")
                    self.rendered_asn_widgets[nama] = p
        else:
            if self.active_categories.get("Pelaksana ASN", False):
                lbl_pelaksana = ctk.CTkLabel(self.scroll_personnel, text="Pelaksana ASN:", font=themed_font(12, "bold"), text_color="#1E3A8A")
                lbl_pelaksana.pack(fill="x", padx=10, pady=(5, 2), anchor="w")
                for p in self.db_asn:
                    nama, nip, jab = p.get('nama', ''), p.get('nip', '-'), p.get('jabatan', '-')
                    key = f"pelaksana_{nama}"
                    if key not in self.pelaksana_vars: self.pelaksana_vars[key] = tk.BooleanVar(value=False)
                    chk = ctk.CTkCheckBox(self.scroll_personnel, text=f"{nama}\nNIP: {nip} | Jabatan: {jab}", variable=self.pelaksana_vars[key], command=lambda: self.schedule_preview_refresh(immediate=True))
                    chk.pack(fill="x", padx=10, pady=5, anchor="w")
                    self.rendered_pelaksana_widgets[nama] = p

            if self.active_categories.get("Pendamping ASN", False):
                if self.rendered_pelaksana_widgets:
                    ctk.CTkFrame(self.scroll_personnel, height=2, fg_color="gray").pack(fill="x", padx=10, pady=10)
                lbl_pendamping = ctk.CTkLabel(self.scroll_personnel, text="Pendamping ASN:", font=themed_font(12, "bold"), text_color="#059669")
                lbl_pendamping.pack(fill="x", padx=10, pady=(5, 2), anchor="w")
                for p in self.db_asn:
                    nama, nip, jab = p.get('nama', ''), p.get('nip', '-'), p.get('jabatan', '-')
                    key = f"pendamping_{nama}"
                    if key not in self.pendamping_vars: self.pendamping_vars[key] = tk.BooleanVar(value=False)
                    chk = ctk.CTkCheckBox(self.scroll_personnel, text=f"{nama}\nNIP: {nip} | Jabatan: {jab}", variable=self.pendamping_vars[key], command=lambda: self.schedule_preview_refresh(immediate=True))
                    chk.pack(fill="x", padx=10, pady=5, anchor="w")
                    self.rendered_pendamping_widgets[nama] = p

    def _has_pelaksana_selected(self):
        """True kalau minimal satu personel (DPRD/ASN/Pelaksana/Pendamping)
        sudah dicentang di panel kanan -- ini gerbang PERTAMA yang wajib
        dipenuhi sebelum bagian formulir lain (nomor surat, tanggal, dst)
        di panel tengah boleh diisi, sesuai permintaan: pelaksana perjalanan
        dinas harus ditentukan terlebih dahulu."""
        for var_dict in (self.dprd_vars, self.asn_vars, self.pelaksana_vars, self.pendamping_vars):
            for v in var_dict.values():
                try:
                    if v.get():
                        return True
                except Exception:
                    pass
        return False

    def select_all_visible(self):
        for n in self.rendered_dprd_widgets: self.dprd_vars[n].set(True)
        for n in self.rendered_asn_widgets: self.asn_vars[n].set(True)
        for n in self.rendered_pelaksana_widgets:
            if f"pelaksana_{n}" in self.pelaksana_vars: self.pelaksana_vars[f"pelaksana_{n}"].set(True)
        for n in self.rendered_pendamping_widgets:
            if f"pendamping_{n}" in self.pendamping_vars: self.pendamping_vars[f"pendamping_{n}"].set(True)
        self.schedule_preview_refresh(immediate=True)

    def clear_all_visible(self):
        for n in self.rendered_dprd_widgets: self.dprd_vars[n].set(False)
        for n in self.rendered_asn_widgets: self.asn_vars[n].set(False)
        for n in self.rendered_pelaksana_widgets:
            if f"pelaksana_{n}" in self.pelaksana_vars: self.pelaksana_vars[f"pelaksana_{n}"].set(False)
        for n in self.rendered_pendamping_widgets:
            if f"pendamping_{n}" in self.pendamping_vars: self.pendamping_vars[f"pendamping_{n}"].set(False)
        self.schedule_preview_refresh(immediate=True)

    def tambah_tujuan(self):
        val = self.ent_tujuan.get().strip()
        if not val: return
        if val not in self.tujuan_terpilih:
            self.tujuan_terpilih.append(val)
            self.refresh_tujuan_list_ui()
        self.ent_tujuan.delete(0, tk.END)
        self.hide_tujuan_suggestions()
        self.schedule_preview_refresh(immediate=True)

    def hapus_tujuan(self, kota):
        if kota in self.tujuan_terpilih:
            self.tujuan_terpilih.remove(kota)
            self.refresh_tujuan_list_ui()
            self.schedule_preview_refresh(immediate=True)

    def refresh_tujuan_list_ui(self):
        if not hasattr(self, 'tujuan_list_frame'): return
        for w in self.tujuan_list_frame.winfo_children(): w.destroy()
        if not self.tujuan_terpilih:
            lbl = ctk.CTkLabel(self.tujuan_list_frame, text="(Belum ada tujuan)", text_color="gray", font=themed_font(11))
            lbl.pack(padx=8, pady=4, anchor="w")
        for kota in self.tujuan_terpilih:
            row_frame = ctk.CTkFrame(self.tujuan_list_frame, fg_color="transparent")
            row_frame.pack(fill="x", padx=4, pady=2)
            lbl = ctk.CTkLabel(row_frame, text=f"📍 {kota}", anchor="w", font=themed_font(11), text_color="#1E3A8A")
            lbl.pack(side="left", padx=(4, 8))
            btn_del = ctk.CTkButton(row_frame, text="✕", width=28, height=22, fg_color="#EF4444", hover_color="#DC2626", font=themed_font(10, "bold"), command=lambda k=kota: self.hapus_tujuan(k))
            btn_del.pack(side="right", padx=2)

    def on_tujuan_key_release(self, event):
        val = self.ent_tujuan.get().strip().lower()
        if len(val) >= 2:
            matches = [item for item in self.database_tujuan if val in item.lower()]
            if matches: self.show_tujuan_suggestions(matches)
            else: self.hide_tujuan_suggestions()
        else: self.hide_tujuan_suggestions()
        if event.keysym == "Return": self.tambah_tujuan()

    def show_tujuan_suggestions(self, matches):
        for widget in self.suggestion_frame.winfo_children(): widget.destroy()
        for match in matches[:6]:
            btn = ctk.CTkButton(self.suggestion_frame, text=match, anchor="w", fg_color="transparent", text_color="black", hover_color="#E5E7EB", command=lambda m=match: self.select_tujuan_suggestion(m))
            btn.pack(fill="x", padx=5, pady=1)
        self.suggestion_frame.pack(fill="x", padx=10, pady=2, before=self.tujuan_list_frame)

    def hide_tujuan_suggestions(self):
        if hasattr(self, 'suggestion_frame'):
            self.suggestion_frame.pack_forget()

    def select_tujuan_suggestion(self, val):
        self.ent_tujuan.delete(0, tk.END)
        self.ent_tujuan.insert(0, val)
        self.hide_tujuan_suggestions()
        self.tambah_tujuan()

    def refresh_signer_dropdowns(self):
        if not hasattr(self, 'combo_ttd_dprd'): return
        ttd_dprd_values = [f"{p.get('jabatan', '')} - {p.get('nama', '')}" for p in self.db_dprd if str(p.get('kategori', '')).lower() == 'pimpinan dprd']
        self.combo_ttd_dprd.configure(values=ttd_dprd_values if ttd_dprd_values else ["-"])
        self.combo_ttd_dprd.set(ttd_dprd_values[0] if ttd_dprd_values else "-")

        ttd_asn_values = [f"{p.get('jabatan', '')} - {p.get('nama', '')}" for p in self.db_asn if 'sekretaris dprd' in str(p.get('jabatan', '')).lower()]
        if not ttd_asn_values: ttd_asn_values = [f"{p.get('jabatan', '')} - {p.get('nama', '')}" for p in self.db_asn]
        self.combo_ttd_asn.configure(values=ttd_asn_values if ttd_asn_values else ["-"])
        self.combo_ttd_asn.set(ttd_asn_values[0] if ttd_asn_values else "-")

    def calculate_duration(self, *args):
        if not hasattr(self, 'dp_mulai'): return
        try:
            if HAS_TKCALENDAR:
                s_date, e_date = self.dp_mulai.get_date(), self.dp_akhir.get_date()
            else:
                s_date = datetime.strptime(self.dp_mulai.get(), "%d/%m/%Y")
                e_date = datetime.strptime(self.dp_akhir.get(), "%d/%m/%Y")
            delta = (e_date - s_date).days + 1
            if delta < 1: delta = 1
            self.ent_lama_hari.configure(state="normal")
            self.ent_lama_hari.delete(0, tk.END)
            self.ent_lama_hari.insert(0, str(delta))
            self.ent_lama_hari.configure(state="readonly")
        except:
            pass
        self.schedule_preview_refresh()

    def _switch_to_undangan_layout(self):
        self.fullwidth_frame.grid_remove()
        self.middle_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        self.right_frame.grid_remove()
        self.preview_frame.grid(row=0, column=2, padx=10, pady=10, sticky="nsew")
        self.grid_columnconfigure(2, weight=2, minsize=480)
        self.grid_columnconfigure(3, weight=0, minsize=0)
        self.update_idletasks()  # settle geometry SEBELUM preview dirender, lihat _apply_preview_image

    def _switch_to_perjalanan_layout(self):
        self.fullwidth_frame.grid_remove()
        self.middle_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        self.right_frame.grid(row=0, column=2, padx=10, pady=10, sticky="nsew")
        self.preview_frame.grid(row=0, column=3, padx=10, pady=10, sticky="nsew")
        self.grid_columnconfigure(2, weight=1, minsize=380)
        self.grid_columnconfigure(3, weight=2, minsize=480)
        self.update_idletasks()  # settle geometry SEBELUM preview dirender, lihat _apply_preview_image

    def _switch_to_fullwidth_layout(self):
        """Dipakai oleh Dashboard & Riwayat Surat: sembunyikan panel form
        (tengah/kanan/preview), lalu tampilkan satu panel lebar penuh yang
        membentang di kolom 1-3 -- tepat di samping sidebar, persis seperti
        mode form lainnya, bukan jendela terpisah."""
        self.middle_frame.grid_remove()
        self.right_frame.grid_remove()
        self.preview_frame.grid_remove()
        self.fullwidth_frame.grid(row=0, column=1, columnspan=3, padx=10, pady=10, sticky="nsew")
        self.update_idletasks()

    def clear_preview(self):
        self.preview_image_label.configure(image=None, text="")

    # ------------------------------------------------------------------
    # SURAT UNDANGAN PARIPURNA (UPDATE)
    # ------------------------------------------------------------------
    def show_undangan_paripurna(self):
        self.current_view = "undangan_paripurna"
        self.middle_frame.configure(label_text="Surat Undangan Paripurna")
        self.btn_generate_main.configure(text="⚡ CETAK SURAT UNDANGAN", command=self.generate_undangan_paripurna)
        self.btn_back_to_perjalanan.grid()
        self._set_nav_button_states("undangan_paripurna")
        self._switch_to_undangan_layout()
        self.setup_undangan_paripurna_form()
        self.combo_preview_jenis.set("Undangan Paripurna")
        # Render preview SEKALI setelah delay singkat supaya layout grid
        # sempat settle dulu (kalau dirender langsung, ukuran panel kadang
        # masih terbaca 0 sehingga preview sempat kosong). Sebelumnya di
        # sini ada 2x render (langsung + setelah delay) yang masing-masing
        # menjalankan generate-docx + convert-PDF penuh -- salah satu
        # penyebab perpindahan menu terasa berat. Sekarang cukup 1x.
        self.after(350, lambda: self.schedule_preview_refresh(immediate=True))

    def show_undangan_biasa(self):
        self.current_view = "undangan_biasa"
        self.middle_frame.configure(label_text="Surat Undangan Rapat Biasa")
        self.btn_generate_main.configure(text="⚡ CETAK SURAT UNDANGAN", command=self.generate_undangan_biasa)
        self.btn_back_to_perjalanan.grid()
        self._set_nav_button_states("undangan_biasa")
        self._switch_to_undangan_layout()
        self.setup_undangan_biasa_form()
        self.combo_preview_jenis.set("Undangan Rapat Biasa")
        self.after(350, lambda: self.schedule_preview_refresh(immediate=True))

    def setup_undangan_paripurna_form(self):
        for widget in self.middle_frame.winfo_children(): widget.destroy()
        self.undangan_inputs = {}

        lbl_tgl_surat = ctk.CTkLabel(self.middle_frame, text="1. Tanggal Surat", anchor="w", font=themed_font(12, "bold"))
        lbl_tgl_surat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.undangan_inputs["tanggal_surat"] = DateEntry(self.middle_frame, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.undangan_inputs["tanggal_surat"].pack(fill="x", padx=10, pady=(0, 6))
            self.undangan_inputs["tanggal_surat"].bind("<<DateEntrySelected>>", lambda e: self.schedule_preview_refresh())
        else:
            self.undangan_inputs["tanggal_surat"] = ctk.CTkEntry(self.middle_frame)
            self.undangan_inputs["tanggal_surat"].pack(fill="x", padx=10, pady=(0, 6))

        lbl_nomor = ctk.CTkLabel(self.middle_frame, text="2. Nomor Undangan", anchor="w", font=themed_font(12, "bold"))
        lbl_nomor.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["nomor_undangan"] = ctk.CTkEntry(self.middle_frame, placeholder_text="Contoh: 01/UNDP/X/2026")
        self.undangan_inputs["nomor_undangan"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["nomor_undangan"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_isi = ctk.CTkLabel(self.middle_frame, text="3. Isi Surat / Agenda Rapat", anchor="w", font=themed_font(12, "bold"))
        lbl_isi.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["isi_surat"] = ctk.CTkTextbox(self.middle_frame, height=80, wrap="word")
        self.undangan_inputs["isi_surat"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["isi_surat"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_tgl_rapat = ctk.CTkLabel(self.middle_frame, text="4. Tanggal Pelaksanaan Rapat", anchor="w", font=themed_font(12, "bold"))
        lbl_tgl_rapat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.undangan_inputs["tanggal_rapat"] = DateEntry(self.middle_frame, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.undangan_inputs["tanggal_rapat"].pack(fill="x", padx=10, pady=(0, 6))
            self.undangan_inputs["tanggal_rapat"].bind("<<DateEntrySelected>>", lambda e: self.update_hari_rapat())
        else:
            self.undangan_inputs["tanggal_rapat"] = ctk.CTkEntry(self.middle_frame)
            self.undangan_inputs["tanggal_rapat"].pack(fill="x", padx=10, pady=(0, 6))

        hari_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        hari_frame.pack(fill="x", padx=10, pady=(0, 6))
        lbl_hari = ctk.CTkLabel(hari_frame, text="Hari:", anchor="w", font=themed_font(11))
        lbl_hari.pack(side="left", padx=(0, 10))
        self.undangan_inputs["hari_rapat"] = ctk.CTkEntry(hari_frame, width=150, state="readonly")
        self.undangan_inputs["hari_rapat"].pack(side="left")

        lbl_jam = ctk.CTkLabel(self.middle_frame, text="5. Jam Pelaksanaan", anchor="w", font=themed_font(12, "bold"))
        lbl_jam.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["jam_pelaksanaan"] = ctk.CTkEntry(self.middle_frame, placeholder_text="Contoh: 09.00 WITA s.d. selesai")
        self.undangan_inputs["jam_pelaksanaan"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["jam_pelaksanaan"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_skenario = ctk.CTkLabel(self.middle_frame, text="6. Skenario Rapat", anchor="w", font=themed_font(12, "bold"))
        lbl_skenario.pack(fill="x", padx=10, pady=(8, 2))
        self.skenario_container = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.skenario_container.pack(fill="x", padx=10, pady=(0, 2))
        self.undangan_inputs["skenario"] = []
        self._add_skenario_row()
        btn_add_skenario = ctk.CTkButton(self.middle_frame, text="+ Tambah Skenario", command=self._add_skenario_row, fg_color="#6366F1", hover_color="#4F46E5", height=28)
        btn_add_skenario.pack(fill="x", padx=10, pady=(2, 6))

        lbl_pakaian = ctk.CTkLabel(self.middle_frame, text="7. Pakaian", anchor="w", font=themed_font(12, "bold"))
        lbl_pakaian.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["pakaian"] = ctk.CTkComboBox(self.middle_frame, values=["PSH", "PSR", "PSL"], command=lambda choice: self.schedule_preview_refresh())
        self.undangan_inputs["pakaian"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["pakaian"].set("PSH")

        lbl_ttd = ctk.CTkLabel(self.middle_frame, text="8. Penandatanganan", anchor="w", font=themed_font(12, "bold"))
        lbl_ttd.pack(fill="x", padx=10, pady=(8, 2))
        ttd_values = [f"{p.get('jabatan', '')} - {p.get('nama', '')}" for p in self.db_dprd if "Pimpinan" in p.get("kategori", "")]
        self.undangan_inputs["penandatanganan"] = ctk.CTkComboBox(self.middle_frame, values=ttd_values if ttd_values else ["-"])
        self.undangan_inputs["penandatanganan"].pack(fill="x", padx=10, pady=(0, 6))
        if ttd_values: self.undangan_inputs["penandatanganan"].set(ttd_values[0])

        self._build_dokumen_pendukung_rapat_section(is_paripurna=True)

        self.update_hari_rapat()

    def _add_skenario_row(self):
        idx = len(self.undangan_inputs["skenario"]) + 1
        row_frame = ctk.CTkFrame(self.skenario_container, fg_color="transparent")
        row_frame.pack(fill="x", pady=2)
        row_frame.grid_columnconfigure(0, weight=1)
        ent = ctk.CTkEntry(row_frame, placeholder_text=f"Skenario {idx}...")
        ent.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        def remove_row(f=row_frame, e=ent):
            self.undangan_inputs["skenario"].remove(e)
            f.destroy()
            self.schedule_preview_refresh()
        btn_del = ctk.CTkButton(row_frame, text="✕", width=30, height=28, fg_color="#EF4444", hover_color="#DC2626", command=remove_row)
        btn_del.grid(row=0, column=1, sticky="e")
        self.undangan_inputs["skenario"].append(ent)

    def update_hari_rapat(self):
        if HAS_TKCALENDAR and "tanggal_rapat" in self.undangan_inputs:
            tanggal = self.undangan_inputs["tanggal_rapat"].get_date()
            hari_nama = ["Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu"]
            hari = hari_nama[tanggal.weekday()]
            self.undangan_inputs["hari_rapat"].configure(state="normal")
            self.undangan_inputs["hari_rapat"].delete(0, tk.END)
            self.undangan_inputs["hari_rapat"].insert(0, hari)
            self.undangan_inputs["hari_rapat"].configure(state="readonly")
            self.schedule_preview_refresh()

    # ------------------------------------------------------------------
    # NASKAH DINAS & DAFTAR HADIR (Undangan Biasa & Paripurna)
    # ------------------------------------------------------------------
    def _build_dokumen_pendukung_rapat_section(self, is_paripurna):
        """Tombol cetak Naskah Dinas & Daftar Hadir. Dipasang di kedua
        formulir (Undangan Paripurna & Undangan Biasa) lewat fungsi yang
        sama supaya tidak duplikat kode -- keduanya memakai template Word
        RESMI kantor (nama anggota sudah tertulis tetap di templatenya,
        BUKAN diambil dari database), jadi generator di sini murni
        mengisi 3-4 variabel header (perihal, hari, tanggal, jam)."""
        lbl = ctk.CTkLabel(
            self.middle_frame, text="📎 Dokumen Pendukung Rapat (opsional)",
            anchor="w", font=themed_font(12, "bold"),
        )
        lbl.pack(fill="x", padx=10, pady=(14, 2))

        if not is_paripurna:
            # Sekretariat & Pihak Terkait dipakai sbg lembar TAMBAHAN yang
            # digabung jadi 1 file bersama Daftar Hadir pelaksana utama
            # (lihat generate_daftar_hadir_rapat_action). Default Pihak
            # Terkait AKTIF karena Undangan Biasa selalu punya kolom Pihak
            # Terkait yang sudah diisi pengguna di form nomor 6.
            self.undangan_inputs["sertakan_pihak_terkait_dh"] = ctk.CTkCheckBox(
                self.middle_frame, text="Sertakan lembar Daftar Hadir Pihak Terkait",
                font=themed_font(11),
            )
            self.undangan_inputs["sertakan_pihak_terkait_dh"].pack(fill="x", padx=10, pady=(0, 4))
            self.undangan_inputs["sertakan_pihak_terkait_dh"].select()

            self.undangan_inputs["sertakan_sekretariat_dh"] = ctk.CTkCheckBox(
                self.middle_frame, text="Sertakan lembar Daftar Hadir Sekretariat",
                font=themed_font(11),
            )
            self.undangan_inputs["sertakan_sekretariat_dh"].pack(fill="x", padx=10, pady=(0, 8))

        btn_row = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        btn_row.pack(fill="x", padx=10, pady=(0, 10))
        btn_row.grid_columnconfigure((0, 1), weight=1)
        btn_naskah = ctk.CTkButton(
            btn_row, text="📋 Cetak Naskah Dinas", height=34,
            fg_color="#0EA5E9", hover_color="#0284C7",
            command=lambda: self.generate_naskah_dinas_action(is_paripurna),
        )
        btn_naskah.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        btn_daftar_hadir = ctk.CTkButton(
            btn_row, text="📝 Cetak Daftar Hadir", height=34,
            fg_color="#16A34A", hover_color="#15803D",
            command=lambda: self.generate_daftar_hadir_rapat_action(is_paripurna),
        )
        btn_daftar_hadir.grid(row=0, column=1, sticky="ew", padx=(4, 0))

    def _ambil_tanggal_jam_rapat_untuk_dokumen_pendukung(self):
        """Baca Hari/Tanggal/Jam Pelaksanaan & Perihal langsung dari
        formulir Undangan yang SEDANG diisi -- Daftar Hadir & Naskah Dinas
        HARUS selalu sinkron dengan isi surat Undangan-nya sendiri (sesuai
        permintaan awal), jadi TIDAK ada input terpisah/ketik ulang."""
        try:
            tanggal_rapat = (
                self.format_indonesian_date(self.undangan_inputs["tanggal_rapat"].get_date())
                if HAS_TKCALENDAR else self.undangan_inputs["tanggal_rapat"].get()
            )
        except Exception:
            tanggal_rapat = ""
        hari = self.undangan_inputs["hari_rapat"].get()
        jam = self.undangan_inputs["jam_pelaksanaan"].get().strip()
        perihal = self.undangan_inputs["isi_surat"].get("1.0", tk.END).strip()
        return hari, tanggal_rapat, jam, perihal

    def generate_naskah_dinas_action(self, is_paripurna):
        if not hasattr(self, "undangan_inputs") or not self.undangan_inputs:
            return
        out_dir = filedialog.askdirectory(title="Pilih Folder untuk Menyimpan Naskah Dinas")
        if not out_dir:
            return
        try:
            nomor = self.undangan_inputs["nomor_undangan"].get().strip()
            try:
                tanggal_surat = (
                    self.format_indonesian_date(self.undangan_inputs["tanggal_surat"].get_date())
                    if HAS_TKCALENDAR else self.undangan_inputs["tanggal_surat"].get()
                )
            except Exception:
                tanggal_surat = ""
            perihal = self.undangan_inputs["isi_surat"].get("1.0", tk.END).strip()

            date_obj = self.undangan_inputs["tanggal_rapat"].get_date() if HAS_TKCALENDAR else None
            prefix = "naskah-dinas-paripurna" if is_paripurna else "naskah-dinas-rapat"
            nama_file = f"{prefix}-{self._hari_tanggal_slug(date_obj)}.docx"
            out_path = os.path.join(out_dir, nama_file)

            generate_naskah_dinas(TEMPLATE_NASKAH_DINAS_RAPAT, out_path, nomor, tanggal_surat, perihal)
            messagebox.showinfo("Berhasil", f"Naskah Dinas berhasil dibuat:\n{out_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Gagal membuat Naskah Dinas:\n{str(e)}")

    def generate_daftar_hadir_rapat_action(self, is_paripurna):
        if not hasattr(self, "undangan_inputs") or not self.undangan_inputs:
            return
        try:
            if is_paripurna:
                template_utama = TEMPLATE_DAFTAR_HADIR_PARIPURNA
            else:
                pelaksana_pilihan = self.undangan_inputs["pelaksana_rapat"].get()
                if pelaksana_pilihan == PELAKSANA_RAPAT_CUSTOM or not pelaksana_pilihan.strip():
                    messagebox.showwarning(
                        "Pelaksana Rapat Belum Dipilih",
                        "Pilih salah satu Pelaksana Rapat dari daftar dulu (bukan opsi "
                        "'Masukan Nama Pelaksana') sebelum mencetak Daftar Hadir -- belum "
                        "ada template Daftar Hadir untuk pelaksana rapat custom.",
                    )
                    return
                template_utama = TEMPLATE_DAFTAR_HADIR_RAPAT_MAP.get(pelaksana_pilihan)
                if not template_utama:
                    # CATATAN: sampai saat ini belum ada template Daftar Hadir
                    # utk "Pimpinan dan Anggota Badan Kehormatan" (tidak ikut
                    # diupload). Daripada mencetak dokumen yang salah/kosong,
                    # aplikasi berhenti & memberi tahu dgn jelas.
                    messagebox.showerror(
                        "Template Belum Tersedia",
                        f"Belum ada template Daftar Hadir untuk '{pelaksana_pilihan}'.\n\n"
                        "Silakan siapkan file Word-nya (format sama seperti Daftar Hadir "
                        "Komisi/AKD lain) lalu tambahkan ke resources/templates/ dan pemetaannya "
                        "di app/config/settings.py (TEMPLATE_DAFTAR_HADIR_RAPAT_MAP).",
                    )
                    return

            hari, tanggal_rapat, jam, perihal = self._ambil_tanggal_jam_rapat_untuk_dokumen_pendukung()

            lembar_tambahan = []
            if not is_paripurna:
                if self.undangan_inputs.get("sertakan_pihak_terkait_dh") and \
                        self.undangan_inputs["sertakan_pihak_terkait_dh"].get():
                    lembar_tambahan.append(TEMPLATE_DAFTAR_HADIR_PIHAK_TERKAIT)
                if self.undangan_inputs.get("sertakan_sekretariat_dh") and \
                        self.undangan_inputs["sertakan_sekretariat_dh"].get():
                    lembar_tambahan.append(TEMPLATE_DAFTAR_HADIR_SEKRETARIAT)
                # TAF (Tenaga Ahli Fraksi) otomatis ikut disertakan HANYA
                # kalau memang ada di daftar Pihak Terkait yang diisi
                # pengguna -- konsisten dengan logika _is_tenaga_ahli_fraksi
                # yang sudah dipakai utk halaman 3 surat Undangan Biasa.
                pihak_raw = [ent.get().strip() for ent in self.undangan_inputs.get("pihak_terkait", []) if ent.get().strip()]
                ada_taf = any(
                    " ".join(p.strip().lower().split()).startswith("tenaga ahli fraksi")
                    for p in pihak_raw
                )
                if ada_taf:
                    lembar_tambahan.append(TEMPLATE_DAFTAR_HADIR_TAF)

            out_dir = filedialog.askdirectory(title="Pilih Folder untuk Menyimpan Daftar Hadir")
            if not out_dir:
                return

            date_obj = self.undangan_inputs["tanggal_rapat"].get_date() if HAS_TKCALENDAR else None
            if is_paripurna:
                nama_file = f"daftar-hadir-paripurna-{self._hari_tanggal_slug(date_obj)}.docx"
            else:
                pelaksana_slug = slugify_filename(self.undangan_inputs["pelaksana_rapat"].get())
                nama_file = f"daftar-hadir-{pelaksana_slug}-{self._hari_tanggal_slug(date_obj)}.docx"
            out_path = os.path.join(out_dir, nama_file)

            generate_daftar_hadir_rapat(
                template_utama, out_path, hari=hari, tanggal=tanggal_rapat, jam=jam,
                isi_perihal=perihal, is_paripurna=is_paripurna, lembar_tambahan=lembar_tambahan,
            )
            jumlah_lembar = 1 + len(lembar_tambahan)
            messagebox.showinfo(
                "Berhasil",
                f"Daftar Hadir berhasil dibuat ({jumlah_lembar} lembar):\n{out_path}",
            )
        except Exception as e:
            messagebox.showerror("Error", f"Gagal membuat Daftar Hadir:\n{str(e)}")

    def get_undangan_context(self):
        """Membuat context data untuk undangan paripurna dengan penomoran skenario."""
        try:
            if HAS_TKCALENDAR:
                tgl_surat = self.format_indonesian_date(self.undangan_inputs["tanggal_surat"].get_date())
                tgl_rapat = self.format_indonesian_date(self.undangan_inputs["tanggal_rapat"].get_date())
            else:
                tgl_surat = self.undangan_inputs["tanggal_surat"].get()
                tgl_rapat = self.undangan_inputs["tanggal_rapat"].get()
        except:
            tgl_surat = ""
            tgl_rapat = ""

        # Skenario: isi teks saja (tanpa nomor), nomor sudah otomatis dari ListParagraph di template
        skenario_raw = [ent.get().strip() for ent in self.undangan_inputs["skenario"] if ent.get().strip()]
        ctx_skenario = {}
        for i in range(7):
            key = f"skenario_paripurna_{i+1}"
            if i < len(skenario_raw):
                # Teks saja — penomoran otomatis dari Word (ListParagraph numId) di template
                ctx_skenario[key] = skenario_raw[i]
            else:
                # Kosong: akan dihapus oleh cleanup_skenario_paripurna setelah render
                ctx_skenario[key] = ""

        ttd_raw = self.undangan_inputs["penandatanganan"].get()
        jab_ttd, nama_ttd = ttd_raw.split(" - ", 1) if " - " in ttd_raw else ("", ttd_raw)
        jab_ttd = format_jabatan_penandatanganan(jab_ttd)

        # Template rapat_paripurna.docx berisi 8 blok halaman yang identik
        # (nomor_surat_paripurna_1 .. nomor_surat_paripurna_8). Nomor yang
        # diketik user dipakai sebagai nomor awal, lalu di-increment per
        # halaman memakai fungsi yang sama dengan surat pemberitahuan
        # (increment_nomor), supaya tiap halaman punya nomor urut berbeda.
        # Format baku kantor: '005/DPRD/XXX/VII/2026' -- '005/DPRD' tetap,
        # yang naik berurutan tiap halaman adalah posisi XXX (segmen ke-3).
        nomor_base = self.undangan_inputs["nomor_undangan"].get().strip()
        ctx_nomor = {
            f"nomor_surat_paripurna_{i+1}": increment_nomor_paripurna(nomor_base, i)
            for i in range(8)
        }

        ctx = {
            "tgl_surat_paripurna": tgl_surat,
            "isi_undangan_paripurna": self.undangan_inputs["isi_surat"].get("1.0", tk.END).strip(),
            "hari_paripurna": self.undangan_inputs["hari_rapat"].get(),
            "tanggal_paripurna": tgl_rapat,
            "jam_paripurna": self.undangan_inputs["jam_pelaksanaan"].get(),
            "pakaian_paripurna": self.undangan_inputs["pakaian"].get(),
            "jabatan_ttd_paripurna": jab_ttd,
            "nama_ttd_paripurna": nama_ttd,
        }
        ctx.update(ctx_skenario)
        ctx.update(ctx_nomor)
        return ctx

    def generate_undangan_paripurna(self):
        if not hasattr(self, 'undangan_inputs'): return
        out_dir = filedialog.askdirectory(title="Pilih Folder untuk Menyimpan Surat Undangan")
        if not out_dir: return

        template_original = TEMPLATE_PARIPURNA
        if not os.path.exists(template_original):
            messagebox.showerror("Error", f"Template tidak ditemukan:\n{template_original}")
            return

        try:
            ctx = self.get_undangan_context()
            
            # MERENDER TEMPLATE 1 KALI SAJA (Karena template sudah punya 8 halaman)
            out_path = os.path.join(out_dir, self._nama_file_undangan_paripurna() + ".docx")
            doc = DocxTemplate(template_original)
            doc.render(ctx)
            doc.save(out_path)
            # Hapus paragraf skenario yang kosong (slot tidak terpakai), lalu
            # pasang page-break eksplisit antar-blok supaya tiap tujuan
            # SELALU mulai di halamannya sendiri (lihat _force_paripurna_page_break).
            doc_clean = Document(out_path)
            cleanup_skenario_paripurna(doc_clean)
            _force_paripurna_page_break(doc_clean)
            doc_clean.save(out_path)

            self._record_riwayat_undangan("paripurna")

            messagebox.showinfo("Berhasil", f"Surat Undangan Paripurna berhasil dibuat:\n{out_path}\n\nBerisi pas 8 halaman tujuan sesuai template master.")
        except Exception as e:
            messagebox.showerror("Error", f"Gagal membuat surat undangan:\n{str(e)}")

    # ------------------------------------------------------------------
    # SURAT UNDANGAN RAPAT BIASA (BARU)
    # ------------------------------------------------------------------
    def _on_pelaksana_rapat_change(self, choice):
        """CTkComboBox pada form ini bisa diketik langsung (editable), jadi
        TIDAK perlu textbox terpisah untuk pelaksana custom. Begitu pengguna
        memilih opsi placeholder PELAKSANA_RAPAT_CUSTOM ('Masukan Nama
        Pelaksana') dari daftar dropdown, satu-satunya hal yang perlu
        dilakukan supaya terasa responsif: kosongkan dulu teksnya (kalau
        tidak, pengguna harus repot select-all + hapus manual teks
        placeholder itu dulu sebelum bisa mengetik nama sungguhan) lalu
        pindahkan fokus & kursor ke widget-nya supaya pengguna bisa langsung
        mengetik tanpa klik apa pun lagi."""
        combo = self.undangan_inputs.get("pelaksana_rapat")
        if choice == PELAKSANA_RAPAT_CUSTOM and combo is not None:
            combo.set("")
            combo.focus_set()
        self.schedule_preview_refresh(immediate=True)

    def setup_undangan_biasa_form(self):
        """Sama seperti setup_undangan_paripurna_form, TAPI TANPA input
        Skenario Rapat & Pakaian (tidak dipakai di template rapat_biasa.docx),
        dan DENGAN tambahan: combobox Pelaksana Rapat, combobox Jenis Rapat,
        dan daftar Pihak Terkait (multi baris, bebas ditambah/dihapus)."""
        for widget in self.middle_frame.winfo_children(): widget.destroy()
        self.undangan_inputs = {}

        lbl_tgl_surat = ctk.CTkLabel(self.middle_frame, text="1. Tanggal Surat", anchor="w", font=themed_font(12, "bold"))
        lbl_tgl_surat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.undangan_inputs["tanggal_surat"] = DateEntry(self.middle_frame, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.undangan_inputs["tanggal_surat"].pack(fill="x", padx=10, pady=(0, 6))
            self.undangan_inputs["tanggal_surat"].bind("<<DateEntrySelected>>", lambda e: self.schedule_preview_refresh())
        else:
            self.undangan_inputs["tanggal_surat"] = ctk.CTkEntry(self.middle_frame)
            self.undangan_inputs["tanggal_surat"].pack(fill="x", padx=10, pady=(0, 6))

        lbl_nomor = ctk.CTkLabel(self.middle_frame, text="2. Nomor Undangan", anchor="w", font=themed_font(12, "bold"))
        lbl_nomor.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["nomor_undangan"] = ctk.CTkEntry(self.middle_frame, placeholder_text="Contoh: 01/UND/X/2026")
        self.undangan_inputs["nomor_undangan"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["nomor_undangan"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_pelaksana = ctk.CTkLabel(self.middle_frame, text="3. Pelaksana Rapat", anchor="w", font=themed_font(12, "bold"))
        lbl_pelaksana.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["pelaksana_rapat"] = ctk.CTkComboBox(self.middle_frame, values=PELAKSANA_RAPAT_OPTIONS, command=self._on_pelaksana_rapat_change)
        self.undangan_inputs["pelaksana_rapat"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["pelaksana_rapat"].set(PELAKSANA_RAPAT_OPTIONS[0])
        # CTkComboBox ini editable (boleh diketik langsung, bukan cuma pilih
        # dari dropdown) -- itu jalur utama untuk pelaksana rapat custom.
        # Ketikan bebas langsung tertangkap oleh KeyRelease supaya preview
        # ikut update real-time, sama seperti field teks bebas lainnya.
        self.undangan_inputs["pelaksana_rapat"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_jenis = ctk.CTkLabel(self.middle_frame, text="4. Jenis Rapat", anchor="w", font=themed_font(12, "bold"))
        lbl_jenis.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["jenis_rapat"] = ctk.CTkComboBox(self.middle_frame, values=JENIS_RAPAT_OPTIONS, command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.undangan_inputs["jenis_rapat"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["jenis_rapat"].set(JENIS_RAPAT_OPTIONS[0])

        lbl_isi = ctk.CTkLabel(self.middle_frame, text="5. Isi Surat / Agenda Rapat", anchor="w", font=themed_font(12, "bold"))
        lbl_isi.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["isi_surat"] = ctk.CTkTextbox(self.middle_frame, height=80, wrap="word")
        self.undangan_inputs["isi_surat"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["isi_surat"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_pihak = ctk.CTkLabel(self.middle_frame, text="6. Pihak Terkait (bisa lebih dari 1)", anchor="w", font=themed_font(12, "bold"))
        lbl_pihak.pack(fill="x", padx=10, pady=(8, 2))
        self.pihak_terkait_container = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.pihak_terkait_container.pack(fill="x", padx=10, pady=(0, 2))
        self.undangan_inputs["pihak_terkait"] = []
        self._add_pihak_terkait_row()
        btn_add_pihak = ctk.CTkButton(self.middle_frame, text="+ Tambah Pihak Terkait", command=self._add_pihak_terkait_row, fg_color="#6366F1", hover_color="#4F46E5", height=28)
        btn_add_pihak.pack(fill="x", padx=10, pady=(2, 6))

        lbl_tujuan_lain = ctk.CTkLabel(self.middle_frame, text="7. Tujuan Surat Lainnya mulai Halaman ke-4 (tiap halaman bisa diisi lebih dari 1 tujuan, dan halaman bisa ditambah tanpa batas)", anchor="w", font=themed_font(12, "bold"))
        lbl_tujuan_lain.pack(fill="x", padx=10, pady=(8, 2))
        self.tujuan_halaman_container = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.tujuan_halaman_container.pack(fill="x", padx=10, pady=(0, 2))
        self.undangan_inputs["tujuan_halaman_groups"] = []
        self._add_tujuan_halaman_group()  # halaman 4 (dasar, tidak bisa dihapus)
        btn_add_halaman = ctk.CTkButton(self.middle_frame, text="+ Tambah Halaman Baru", command=self._add_tujuan_halaman_group, fg_color="#0EA5E9", hover_color="#0284C7", height=30)
        btn_add_halaman.pack(fill="x", padx=10, pady=(2, 6))

        lbl_tgl_rapat = ctk.CTkLabel(self.middle_frame, text="8. Tanggal Pelaksanaan Rapat", anchor="w", font=themed_font(12, "bold"))
        lbl_tgl_rapat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.undangan_inputs["tanggal_rapat"] = DateEntry(self.middle_frame, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.undangan_inputs["tanggal_rapat"].pack(fill="x", padx=10, pady=(0, 6))
            self.undangan_inputs["tanggal_rapat"].bind("<<DateEntrySelected>>", lambda e: self.update_hari_rapat())
        else:
            self.undangan_inputs["tanggal_rapat"] = ctk.CTkEntry(self.middle_frame)
            self.undangan_inputs["tanggal_rapat"].pack(fill="x", padx=10, pady=(0, 6))

        hari_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        hari_frame.pack(fill="x", padx=10, pady=(0, 6))
        lbl_hari = ctk.CTkLabel(hari_frame, text="Hari:", anchor="w", font=themed_font(11))
        lbl_hari.pack(side="left", padx=(0, 10))
        self.undangan_inputs["hari_rapat"] = ctk.CTkEntry(hari_frame, width=150, state="readonly")
        self.undangan_inputs["hari_rapat"].pack(side="left")

        lbl_jam = ctk.CTkLabel(self.middle_frame, text="9. Jam Pelaksanaan", anchor="w", font=themed_font(12, "bold"))
        lbl_jam.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["jam_pelaksanaan"] = ctk.CTkEntry(self.middle_frame, placeholder_text="Contoh: 09.00 WITA s.d. selesai")
        self.undangan_inputs["jam_pelaksanaan"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["jam_pelaksanaan"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_ttd = ctk.CTkLabel(self.middle_frame, text="10. Penandatanganan", anchor="w", font=themed_font(12, "bold"))
        lbl_ttd.pack(fill="x", padx=10, pady=(8, 2))
        ttd_values = [f"{p.get('jabatan', '')} - {p.get('nama', '')}" for p in self.db_dprd if "Pimpinan" in p.get("kategori", "")]
        self.undangan_inputs["penandatanganan"] = ctk.CTkComboBox(self.middle_frame, values=ttd_values if ttd_values else ["-"])
        self.undangan_inputs["penandatanganan"].pack(fill="x", padx=10, pady=(0, 6))
        if ttd_values: self.undangan_inputs["penandatanganan"].set(ttd_values[0])

        self._build_dokumen_pendukung_rapat_section(is_paripurna=False)

        self.update_hari_rapat()

    def _add_pihak_terkait_row(self):
        """Baris isian bebas untuk daftar pihak terkait yang diundang.
        Polanya identik dengan _add_skenario_row (paripurna): tambah/hapus
        baris bebas, bukan dropdown/checklist -- karena pihak terkait rapat
        biasa isinya sembarang instansi/nama, bukan personel dari database."""
        idx = len(self.undangan_inputs["pihak_terkait"]) + 1
        row_frame = ctk.CTkFrame(self.pihak_terkait_container, fg_color="transparent")
        row_frame.pack(fill="x", pady=2)
        row_frame.grid_columnconfigure(0, weight=1)
        ent = ctk.CTkEntry(row_frame, placeholder_text=f"Pihak Terkait {idx}... (contoh: Kepala Dinas PUPR Kota Bitung)")
        ent.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        def remove_row(f=row_frame, e=ent):
            self.undangan_inputs["pihak_terkait"].remove(e)
            f.destroy()
            self.schedule_preview_refresh()
        btn_del = ctk.CTkButton(row_frame, text="✕", width=30, height=28, fg_color="#EF4444", hover_color="#DC2626", command=remove_row)
        btn_del.grid(row=0, column=1, sticky="e")
        self.undangan_inputs["pihak_terkait"].append(ent)

    def _add_tujuan_halaman_group(self):
        """Menambah 1 'halaman' baru untuk Tujuan Surat Lainnya. Halaman
        pertama (index 0) mewakili halaman ke-4 template utama dan tidak
        bisa dihapus. Halaman ke-2 dst adalah halaman tambahan bebas.
        Setiap halaman punya daftar tujuan sendiri yang bisa ditambah tanpa
        batas -- persis pola form nomor 6 (Pihak Terkait)."""
        idx = len(self.undangan_inputs["tujuan_halaman_groups"])
        is_base = (idx == 0)

        group_frame = ctk.CTkFrame(self.tujuan_halaman_container, fg_color="#F1F5F9", corner_radius=8)
        group_frame.pack(fill="x", pady=(4, 4))

        header_frame = ctk.CTkFrame(group_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=8, pady=(6, 2))
        header_text = "Halaman ke-4 (utama)" if is_base else f"Halaman Tambahan #{idx + 1}"
        lbl_header = ctk.CTkLabel(header_frame, text=header_text, anchor="w", font=themed_font(11, "bold"), text_color="#334155")
        lbl_header.pack(side="left")

        entries_container = ctk.CTkFrame(group_frame, fg_color="transparent")
        entries_container.pack(fill="x", padx=8, pady=(0, 2))

        group = {"frame": group_frame, "entries_container": entries_container, "entries": []}

        if not is_base:
            def remove_group(g=group):
                self.undangan_inputs["tujuan_halaman_groups"].remove(g)
                g["frame"].destroy()
                self.schedule_preview_refresh()
            btn_del_group = ctk.CTkButton(header_frame, text="✕ Hapus Halaman Ini", width=120, height=24, fg_color="#EF4444", hover_color="#DC2626", command=remove_group)
            btn_del_group.pack(side="right")

        self.undangan_inputs["tujuan_halaman_groups"].append(group)
        self._add_tujuan_entry(group)  # baris tujuan pertama di halaman ini

        btn_add_entry = ctk.CTkButton(group_frame, text="+ Tambah Tujuan di Halaman Ini", command=lambda g=group: self._add_tujuan_entry(g), fg_color="#6366F1", hover_color="#4F46E5", height=26)
        btn_add_entry.pack(fill="x", padx=8, pady=(2, 8))

    def _add_tujuan_entry(self, group):
        """Menambah 1 baris isian tujuan di dalam sebuah halaman (group).
        Jumlah tujuan per halaman tidak dibatasi."""
        idx = len(group["entries"]) + 1
        row_frame = ctk.CTkFrame(group["entries_container"], fg_color="transparent")
        row_frame.pack(fill="x", pady=2)
        row_frame.grid_columnconfigure(0, weight=1)
        ent = ctk.CTkEntry(row_frame, placeholder_text=f"Tujuan {idx}... (contoh: CAMAT BITUNG UTARA)")
        ent.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        def remove_entry(f=row_frame, e=ent):
            group["entries"].remove(e)
            f.destroy()
            self.schedule_preview_refresh()
        btn_del = ctk.CTkButton(row_frame, text="✕", width=30, height=28, fg_color="#EF4444", hover_color="#DC2626", command=remove_entry)
        btn_del.grid(row=0, column=1, sticky="e")
        group["entries"].append(ent)

    def get_undangan_biasa_context(self):
        """Membuat context data untuk undangan rapat biasa. Mengikuti pola
        yang sama dengan get_undangan_context (paripurna): nomor per-halaman
        di-increment otomatis, slot pihak_terkait yang kosong dibersihkan
        setelah render (lihat cleanup_skenario_paripurna, dipakai ulang
        karena template rapat_biasa.docx juga memakai style 'List Paragraph'
        yang sama untuk slot dinamisnya)."""
        try:
            if HAS_TKCALENDAR:
                tgl_surat = self.format_indonesian_date(self.undangan_inputs["tanggal_surat"].get_date())
                tgl_rapat = self.format_indonesian_date(self.undangan_inputs["tanggal_rapat"].get_date())
            else:
                tgl_surat = self.undangan_inputs["tanggal_surat"].get()
                tgl_rapat = self.undangan_inputs["tanggal_rapat"].get()
        except Exception:
            tgl_surat = ""
            tgl_rapat = ""

        # Pihak terkait: sampai 5 slot (jumlah maksimum di template, lihat
        # halaman 1). Sisa slot kosong dihapus otomatis oleh cleanup.
        pihak_raw = [ent.get().strip() for ent in self.undangan_inputs["pihak_terkait"] if ent.get().strip()]
        # PENTING: JANGAN dibatasi ke 5 slot lagi (bug lama -- entri ke-6 dst
        # diam-diam hilang karena master template cuma punya 5 placeholder
        # pihak_terkait_N). Sekarang jumlah slot di template diperluas
        # dinamis sesuai jumlah entri sesungguhnya lewat ensure_numbered_slots
        # (dipanggil di _render_undangan_biasa_docx SEBELUM render), jadi di
        # sini semua entri pengguna diikutkan apa adanya, tanpa batas atas.
        ctx_pihak = {f"pihak_terkait_{i+1}": nama for i, nama in enumerate(pihak_raw)}

        # Halaman ke-2 (permohonan ke WALI KOTA BITUNG) memakai daftar pihak
        # yang SUDAH DISARING: entri "Tenaga Ahli Fraksi (DPRD Kota Bitung)"
        # tidak boleh ikut dimintakan lewat Wali Kota, karena pihak itu sudah
        # jadi tujuan surat tersendiri di halaman ke-3 (lihat placeholder
        # pihak_terkait_wk_1..N di template, dipisah dari pihak_terkait_1..N
        # supaya halaman 1 tetap menampilkan semua pihak apa adanya).
        def _is_tenaga_ahli_fraksi(nama):
            n = " ".join(nama.strip().lower().split())
            return n in ("tenaga ahli fraksi", "tenaga ahli fraksi dprd kota bitung") or n.startswith("tenaga ahli fraksi")

        pihak_wk = [p for p in pihak_raw if not _is_tenaga_ahli_fraksi(p)]
        ctx_pihak_wk = {f"pihak_terkait_wk_{i+1}": nama for i, nama in enumerate(pihak_wk)}

        ttd_raw = self.undangan_inputs["penandatanganan"].get()
        jab_ttd, nama_ttd = ttd_raw.split(" - ", 1) if " - " in ttd_raw else ("", ttd_raw)
        jab_ttd = format_jabatan_penandatanganan(jab_ttd)

        # Template rapat_biasa.docx berisi 4 blok halaman (nomor_surat_rapat_1
        # .. nomor_surat_rapat_4). Nomor yang diketik user jadi nomor awal,
        # lalu di-increment per halaman (sama seperti undangan paripurna).
        nomor_base = self.undangan_inputs["nomor_undangan"].get().strip()
        ctx_nomor = {
            f"nomor_surat_rapat_{i+1}": increment_nomor(nomor_base, i)
            for i in range(4)
        }

        # {{tujuan_surat_dprd}} (halaman 1) otomatis mengikuti kategori
        # Pelaksana Rapat yang dipilih pengguna, HURUF BESAR SEMUA + BOLD.
        # Halaman 2 (WALI KOTA BITUNG) & halaman 3 (TENAGA AHLI FRAKSI DPRD
        # KOTA BITUNG) SUDAH teks tetap di template master, tidak disentuh
        # di sini -- sesuai instruksi pengguna.
        pelaksana_pilihan = self.undangan_inputs["pelaksana_rapat"].get()
        if pelaksana_pilihan == PELAKSANA_RAPAT_CUSTOM:
            # Placeholder-nya sendiri dibiarkan terpilih tanpa diketik ulang
            # (combobox ini editable, tapi pengguna mungkin belum sempat
            # mengetik nama pelaksana yang sesungguhnya) -- JANGAN biarkan
            # literal "Masukan Nama Pelaksana" ikut tercetak di surat.
            pelaksana_pilihan = ""
        tujuan_dprd_text = TUJUAN_SURAT_DPRD_MAP.get(pelaksana_pilihan, pelaksana_pilihan.upper())
        tujuan_dprd_rt = RichText()
        tujuan_dprd_rt.add(tujuan_dprd_text, bold=True)

        # Struktur baru: daftar "halaman" (tujuan_halaman_groups), masing2
        # halaman punya daftar tujuan sendiri TANPA BATAS (sama seperti
        # pola form 6/Pihak Terkait). Halaman pertama (index 0) = halaman
        # ke-4 template utama; halaman ke-2 dst dicetak sbg halaman baru
        # oleh generate_undangan_biasa, masing2 dgn tujuan-nya sendiri.
        def _kumpulkan_tujuan(group):
            return [ent.get().strip() for ent in group["entries"] if ent.get().strip()]

        halaman_groups = self.undangan_inputs["tujuan_halaman_groups"]
        tujuan_halaman_4 = _kumpulkan_tujuan(halaman_groups[0]) if halaman_groups else []
        tujuan_lainnya_rt = build_tujuan_richtext(tujuan_halaman_4)
        # Tiap halaman tambahan (index 1 dst) jadi 1 elemen list of list-nama
        tujuan_lainnya_extra = [_kumpulkan_tujuan(g) for g in halaman_groups[1:]]
        tujuan_lainnya_extra = [names for names in tujuan_lainnya_extra if names]

        ctx = {
            "tgl_surat_rapat_biasa": tgl_surat,
            "pelaksana_rapat_skpd": pelaksana_pilihan,
            "jenis_rapat_biasa_skpd": self.undangan_inputs["jenis_rapat"].get(),
            "isi_surat_rapat_biasa": self.undangan_inputs["isi_surat"].get("1.0", tk.END).strip(),
            "isi_surat_rapat_biasa_skpd": self.undangan_inputs["isi_surat"].get("1.0", tk.END).strip(),
            "hari_rapat_biasa": self.undangan_inputs["hari_rapat"].get(),
            "tanggal_rapat_biasa": tgl_rapat,
            "jam_rapat_biasa": self.undangan_inputs["jam_pelaksanaan"].get(),
            "tujuan_surat_dprd": tujuan_dprd_rt,
            "tujuan_surat_lainnya": tujuan_lainnya_rt,
            "jabatan_ttd_rapat": jab_ttd,
            "nama_ttd_rapat": nama_ttd,
        }
        ctx.update(ctx_pihak)
        ctx.update(ctx_pihak_wk)
        ctx.update(ctx_nomor)
        ctx["_tujuan_lainnya_extra"] = tujuan_lainnya_extra
        ctx["_nomor_base"] = nomor_base
        ctx["_pihak_terkait_count"] = len(pihak_raw)
        ctx["_pihak_terkait_wk_count"] = len(pihak_wk)
        return ctx

    def _render_undangan_biasa_docx(self, ctx, out_path):
        """Merender surat Undangan Rapat Biasa lengkap dgn semua halaman
        tambahan (jika ada) ke satu file out_path. Dipakai BERSAMA oleh
        generate_undangan_biasa() dan _preview_worker() supaya preview dan
        hasil cetak akhir selalu identik -- sebelumnya preview punya jalur
        render sendiri yang tidak tahu soal halaman tambahan, itu sebabnya
        halaman tambahan tidak muncul di pratinjau."""
        template_original = TEMPLATE_RAPAT_BIASA
        ctx = dict(ctx)
        tujuan_lainnya_extra = ctx.pop("_tujuan_lainnya_extra", [])
        nomor_base = ctx.pop("_nomor_base", "")
        pihak_terkait_count = ctx.pop("_pihak_terkait_count", 0)
        pihak_terkait_wk_count = ctx.pop("_pihak_terkait_wk_count", 0)

        main_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
        doc = DocxTemplate(template_original)
        # Perluas slot placeholder pihak_terkait_N / pihak_terkait_wk_N di
        # template SEBELUM render, kalau jumlah pihak yang diisi pengguna
        # melebihi jumlah slot bawaan template -- lihat ensure_numbered_slots
        # untuk detail bug yang diperbaiki (entri berlebih dulu hilang diam-diam).
        # Dibungkus try/except: kalau API docxtpl versi tertentu berubah,
        # generate surat tetap lanjut jalan (fallback ke perilaku lama -
        # dibatasi jumlah slot bawaan template) daripada gagal total.
        try:
            ensure_numbered_slots(doc.get_docx(), "pihak_terkait", pihak_terkait_count)
            ensure_numbered_slots(doc.get_docx(), "pihak_terkait_wk", pihak_terkait_wk_count)
        except Exception as e:
            safe_log(f"Gagal memperluas slot pihak_terkait (lanjut dgn slot bawaan template): {e}")
        doc.render(ctx)
        doc.save(main_tmp)
        doc_clean = Document(main_tmp)
        cleanup_skenario_paripurna(doc_clean)
        doc_clean.save(main_tmp)

        files_to_combine = [main_tmp]
        if tujuan_lainnya_extra:
            cache_path = os.path.join(tempfile.gettempdir(), "sips_halaman_tujuan_lain_template.docx")
            halaman_lain_template = build_halaman_tujuan_lain_template(template_original, cache_path)
            for i, nama_list in enumerate(tujuan_lainnya_extra):
                extra_ctx = dict(ctx)
                extra_ctx["tujuan_surat_lainnya"] = build_tujuan_richtext(nama_list)
                extra_ctx["nomor_surat_rapat_4"] = increment_nomor(nomor_base, 4 + i)
                extra_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                extra_doc = DocxTemplate(halaman_lain_template)
                extra_doc.render(extra_ctx)
                extra_doc.save(extra_tmp)
                files_to_combine.append(extra_tmp)

        if len(files_to_combine) > 1:
            _combine_word_pages(files_to_combine, out_path)
        else:
            shutil.copy(main_tmp, out_path)

        for f in files_to_combine:
            try:
                os.remove(f)
            except OSError:
                pass

    def generate_undangan_biasa(self):
        if not hasattr(self, 'undangan_inputs'): return
        out_dir = filedialog.askdirectory(title="Pilih Folder untuk Menyimpan Surat Undangan Rapat Biasa")
        if not out_dir: return

        template_original = TEMPLATE_RAPAT_BIASA
        if not os.path.exists(template_original):
            messagebox.showerror("Error", f"Template tidak ditemukan:\n{template_original}")
            return

        try:
            ctx = self.get_undangan_biasa_context()
            out_path = os.path.join(out_dir, self._nama_file_undangan_biasa(ctx) + ".docx")
            self._render_undangan_biasa_docx(ctx, out_path)

            self._record_riwayat_undangan("biasa")

            messagebox.showinfo("Berhasil", f"Surat Undangan Rapat Biasa berhasil dibuat:\n{out_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Gagal membuat surat undangan rapat biasa:\n{str(e)}")


    # ------------------------------------------------------------------
    # DATA PERJALANAN DINAS & CONTEXT BUILDING
    # ------------------------------------------------------------------
    def show_perjalanan_dinas(self):
        self.current_view = "perjalanan_dinas"
        self.middle_frame.configure(label_text="Data Perjalanan Dinas")
        self.btn_generate_main.configure(text="⚡ CETAK SURAT & SPD", command=self.generate_documents_action)
        self.btn_back_to_perjalanan.grid_remove()
        self._set_nav_button_states(None)
        self._switch_to_perjalanan_layout()
        self.setup_perjalanan_dinas_form()
        self.combo_preview_jenis.set("Surat Tugas (DPRD)" if self.mode == "dprd" else "Surat Tugas (ASN)")
        self.after(350, lambda: self.schedule_preview_refresh(immediate=True))

    def setup_perjalanan_dinas_form(self):
        for widget in self.middle_frame.winfo_children(): widget.destroy()
        self.inputs = {}
        self.mode_specific_widgets = {}
        # Formulir dibangun ulang dari nol -> anggap ini surat BARU sampai
        # load_riwayat_perjalanan() bilang sebaliknya (lihat catatan di
        # build_context tentang _editing_perjalanan_key).
        self._editing_perjalanan_key = None
        # Setiap bagian formulir dibungkus frame tersendiri (sec_*) supaya
        # StepLockManager bisa menimpanya dengan overlay kunci saat bagian
        # SEBELUMNYA (termasuk pelaksana di panel kanan) belum lengkap.
        # Urutan pembuatan frame ini SAMA dengan urutan wajib pengisian.
        self.step_lock.reset()

        self.sec_nomor_surat = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_nomor_surat.pack(fill="x", padx=0, pady=0)
        for var_name, label in [("nomor_surat", "Nomor Surat Tugas DPRD"), ("nomor_surat_asn", "Nomor Surat Tugas Setwan")]:
            lbl = ctk.CTkLabel(self.sec_nomor_surat, text=label, anchor="w", font=themed_font(12, "bold"))
            lbl.pack(fill="x", padx=10, pady=(8, 2))
            ent = ctk.CTkEntry(self.sec_nomor_surat, placeholder_text=f"Masukkan {label.lower()}...")
            ent.pack(fill="x", padx=10, pady=(0, 6))
            ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
            self.inputs[var_name] = ent

        self.sec_pemberitahuan = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_pemberitahuan.pack(fill="x", padx=0, pady=0)
        self.mode_specific_widgets["lbl_pemberitahuan_dprd"] = ctk.CTkLabel(self.sec_pemberitahuan, text="Nomor Surat Pemberitahuan DPRD", anchor="w", font=themed_font(12, "bold"))
        self.mode_specific_widgets["lbl_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(8, 2))
        self.inputs["nomor_pemberitahuan_dprd"] = ctk.CTkEntry(self.sec_pemberitahuan, placeholder_text="Masukkan nomor surat pemberitahuan dprd...")
        self.inputs["nomor_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(0, 6))
        self.inputs["nomor_pemberitahuan_dprd"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        self.mode_specific_widgets["ent_pemberitahuan_dprd"] = self.inputs["nomor_pemberitahuan_dprd"]

        for var_name, label in [("nomor_pemberitahuan_asn", "Nomor Surat Pemberitahuan Setwan")]:
            lbl = ctk.CTkLabel(self.sec_pemberitahuan, text=label, anchor="w", font=themed_font(12, "bold"))
            lbl.pack(fill="x", padx=10, pady=(8, 2))
            ent = ctk.CTkEntry(self.sec_pemberitahuan, placeholder_text=f"Masukkan {label.lower()}...")
            ent.pack(fill="x", padx=10, pady=(0, 6))
            ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
            self.inputs[var_name] = ent

        self.sec_spd = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_spd.pack(fill="x", padx=0, pady=0)
        lbl_spd_title = ctk.CTkLabel(self.sec_spd, text="Nomor SPD", anchor="w", font=themed_font(12, "bold"))
        lbl_spd_title.pack(fill="x", padx=10, pady=(10, 2))

        spd_frame = ctk.CTkFrame(self.sec_spd, fg_color="transparent")
        spd_frame.pack(fill="x", padx=10, pady=(0, 6))
        spd_frame.grid_columnconfigure(0, weight=1)
        spd_frame.grid_columnconfigure(1, weight=1)

        self.mode_specific_widgets["lbl_spd_dprd"] = ctk.CTkLabel(spd_frame, text="Nomor SPD DPRD :", anchor="w", font=themed_font(11), text_color="#1E3A8A")
        self.mode_specific_widgets["lbl_spd_dprd"].grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.inputs["nomor_spd_dprd"] = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 10/SPD/X/2026/")
        self.inputs["nomor_spd_dprd"].grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.inputs["nomor_spd_dprd"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        self.mode_specific_widgets["ent_spd_dprd"] = self.inputs["nomor_spd_dprd"]

        self.mode_specific_widgets["lbl_spd_setwan"] = ctk.CTkLabel(spd_frame, text="Nomor SPD Setwan (ASN) :", anchor="w", font=themed_font(11), text_color="#1E3A8A")
        self.mode_specific_widgets["lbl_spd_setwan"].grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.inputs["nomor_spd_asn"] = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 20/SPD/X/2026/")
        self.inputs["nomor_spd_asn"].grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.inputs["nomor_spd_asn"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        self.lbl_spd_pelaksana = ctk.CTkLabel(spd_frame, text="Nomor SPD Pelaksana ASN :", anchor="w", font=themed_font(11), text_color="#1E3A8A")
        self.ent_spd_pelaksana = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 10/SPD-PL/X/2026/")
        self.ent_spd_pelaksana.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        self.lbl_spd_pendamping = ctk.CTkLabel(spd_frame, text="Nomor SPD Pendamping ASN :", anchor="w", font=themed_font(11), text_color="#059669")
        self.ent_spd_pendamping = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 20/SPD-PD/X/2026/")
        self.ent_spd_pendamping.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_spd_info = ctk.CTkLabel(self.sec_spd, text="ℹ️  SPD DPRD: semua anggota pakai nomor sama  |  SPD ASN: nomor otomatis berurutan", anchor="w", font=themed_font(10), text_color="gray")
        lbl_spd_info.pack(fill="x", padx=10, pady=(0, 6))
        self.mode_specific_widgets["lbl_spd_info"] = lbl_spd_info

        self.sec_dasar = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_dasar.pack(fill="x", padx=0, pady=0)
        lbl_dasar_title = ctk.CTkLabel(self.sec_dasar, text="Dasar Surat Tugas", anchor="w", font=themed_font(12, "bold"))
        lbl_dasar_title.pack(fill="x", padx=10, pady=(10, 2))
        dasar_frame = ctk.CTkFrame(self.sec_dasar, fg_color="transparent")
        dasar_frame.pack(fill="x", padx=10, pady=(0, 6))
        dasar_frame.grid_columnconfigure(0, weight=1)
        dasar_frame.grid_columnconfigure(1, weight=1)
        lbl_dasar_dprd = ctk.CTkLabel(dasar_frame, text="Dasar Surat Tugas DPRD :", anchor="w", font=themed_font(11), text_color="#1E3A8A")
        lbl_dasar_dprd.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.txt_dasar_dprd = ctk.CTkTextbox(dasar_frame, height=60, wrap="word")
        self.txt_dasar_dprd.grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.txt_dasar_dprd.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        lbl_dasar_asn = ctk.CTkLabel(dasar_frame, text="Dasar Surat Tugas ASN :", anchor="w", font=themed_font(11), text_color="#1E3A8A")
        lbl_dasar_asn.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.txt_dasar_asn = ctk.CTkTextbox(dasar_frame, height=60, wrap="word")
        self.txt_dasar_asn.grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.txt_dasar_asn.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        self.sec_materi = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_materi.pack(fill="x", padx=0, pady=0)
        lbl_materi_title = ctk.CTkLabel(self.sec_materi, text="Materi / Agenda Kegiatan", anchor="w", font=themed_font(12, "bold"))
        lbl_materi_title.pack(fill="x", padx=10, pady=(10, 2))
        materi_frame = ctk.CTkFrame(self.sec_materi, fg_color="transparent")
        materi_frame.pack(fill="x", padx=10, pady=(0, 6))
        materi_frame.grid_columnconfigure(0, weight=1)
        materi_frame.grid_columnconfigure(1, weight=1)
        lbl_mt_st = ctk.CTkLabel(materi_frame, text="Surat Tugas & SPPD :", anchor="w", font=themed_font(11), text_color="#1E3A8A")
        lbl_mt_st.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.txt_materi_st = ctk.CTkTextbox(materi_frame, height=70, wrap="word")
        self.txt_materi_st.grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.txt_materi_st.bind("<KeyRelease>", lambda e: (self.schedule_preview_refresh(), self._schedule_cek_judul_duplikat()))
        lbl_mt_pb = ctk.CTkLabel(materi_frame, text="Surat Pemberitahuan :", anchor="w", font=themed_font(11), text_color="#1E3A8A")
        lbl_mt_pb.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.txt_materi_pb = ctk.CTkTextbox(materi_frame, height=70, wrap="word")
        self.txt_materi_pb.grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.txt_materi_pb.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        self.sec_jenis = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_jenis.pack(fill="x", padx=0, pady=0)
        lbl_jp = ctk.CTkLabel(self.sec_jenis, text="Jenis Perjalanan", anchor="w", font=themed_font(12, "bold"))
        lbl_jp.pack(fill="x", padx=10, pady=(8, 2))
        self.combo_jenis = ctk.CTkComboBox(self.sec_jenis, values=["Kunjungan Kerja", "Kunjungan Konsultasi", "Bimbingan Teknis"], command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_jenis.pack(fill="x", padx=10, pady=(0, 6))

        self.sec_tanggal = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_tanggal.pack(fill="x", padx=0, pady=0)
        lbl_tgl_surat = ctk.CTkLabel(self.sec_tanggal, text="Tanggal Surat", anchor="w", font=themed_font(12, "bold"))
        lbl_tgl_surat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_surat = DateEntry(self.sec_tanggal, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_surat.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_surat.bind("<<DateEntrySelected>>", lambda e: self.schedule_preview_refresh())
        else:
            self.dp_surat = ctk.CTkEntry(self.sec_tanggal)
            self.dp_surat.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_surat.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_tgl_mulai = ctk.CTkLabel(self.sec_tanggal, text="Tanggal Mulai Tugas", anchor="w", font=themed_font(12, "bold"))
        lbl_tgl_mulai.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_mulai = DateEntry(self.sec_tanggal, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_mulai.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_mulai.bind("<<DateEntrySelected>>", lambda e: self.calculate_duration())
        else:
            self.dp_mulai = ctk.CTkEntry(self.sec_tanggal)
            self.dp_mulai.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_mulai.bind("<KeyRelease>", lambda e: self.calculate_duration())

        lbl_tgl_akhir = ctk.CTkLabel(self.sec_tanggal, text="Tanggal Akhir Tugas", anchor="w", font=themed_font(12, "bold"))
        lbl_tgl_akhir.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_akhir = DateEntry(self.sec_tanggal, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_akhir.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_akhir.bind("<<DateEntrySelected>>", lambda e: self.calculate_duration())
        else:
            self.dp_akhir = ctk.CTkEntry(self.sec_tanggal)
            self.dp_akhir.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_akhir.bind("<KeyRelease>", lambda e: self.calculate_duration())

        lbl_lama_hari = ctk.CTkLabel(self.sec_tanggal, text="Lama Hari Perjalanan (Otomatis)", anchor="w", font=themed_font(12, "bold"))
        lbl_lama_hari.pack(fill="x", padx=10, pady=(8, 2))
        self.ent_lama_hari = ctk.CTkEntry(self.sec_tanggal, fg_color="#F3F4F6")
        self.ent_lama_hari.pack(fill="x", padx=10, pady=(0, 6))

        self.sec_tujuan = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_tujuan.pack(fill="x", padx=0, pady=0)
        lbl_tujuan = ctk.CTkLabel(self.sec_tujuan, text="Kota Tujuan Bertugas (Multi Lokasi)", anchor="w", font=themed_font(12, "bold"))
        lbl_tujuan.pack(fill="x", padx=10, pady=(8, 2))
        tujuan_input_frame = ctk.CTkFrame(self.sec_tujuan, fg_color="transparent")
        tujuan_input_frame.pack(fill="x", padx=10, pady=(0, 2))
        tujuan_input_frame.grid_columnconfigure(0, weight=1)
        self.ent_tujuan = ctk.CTkEntry(tujuan_input_frame, placeholder_text="Ketik nama kota lalu klik Tambah...")
        self.ent_tujuan.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self.ent_tujuan.bind("<KeyRelease>", self.on_tujuan_key_release)
        self.btn_tambah_tujuan = ctk.CTkButton(tujuan_input_frame, text="+ Tambah", width=80, command=self.tambah_tujuan)
        self.btn_tambah_tujuan.grid(row=0, column=1, sticky="e")

        self.suggestion_frame = ctk.CTkScrollableFrame(self.sec_tujuan, height=110, fg_color="#F3F4F6")
        lbl_tujuan_terpilih = ctk.CTkLabel(self.sec_tujuan, text="Lokasi yang dipilih:", anchor="w", font=themed_font(11), text_color="gray")
        lbl_tujuan_terpilih.pack(fill="x", padx=10, pady=(4, 1))
        self.tujuan_list_frame = ctk.CTkScrollableFrame(self.sec_tujuan, height=80, fg_color="#F0F4FF")
        self.tujuan_list_frame.pack(fill="x", padx=10, pady=(0, 6))
        lbl_tujuan_hint = ctk.CTkLabel(self.sec_tujuan, text="ℹ️  Klik ✕ pada lokasi untuk menghapus. Urutan sesuai tampilan.", anchor="w", font=themed_font(10), text_color="gray")
        lbl_tujuan_hint.pack(fill="x", padx=10, pady=(0, 6))

        self.sec_ttd = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.sec_ttd.pack(fill="x", padx=0, pady=0)
        self.lbl_sign_dprd = ctk.CTkLabel(self.sec_ttd, text="Penandatangan DPRD:", font=themed_font(12, "bold"))
        self.lbl_sign_dprd.pack(fill="x", padx=10, pady=(15, 2))
        self.combo_ttd_dprd = ctk.CTkComboBox(self.sec_ttd, values=["-"], height=32, command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_ttd_dprd.pack(fill="x", padx=10, pady=(0, 8))

        self.lbl_sign_asn = ctk.CTkLabel(self.sec_ttd, text="Penandatangan ASN / SPPD:", font=themed_font(12, "bold"))
        self.lbl_sign_asn.pack(fill="x", padx=10, pady=(10, 2))
        self.combo_ttd_asn = ctk.CTkComboBox(self.sec_ttd, values=["-"], height=32, command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_ttd_asn.pack(fill="x", padx=10, pady=(0, 8))

        self.refresh_signer_dropdowns()

        # DEFAULTS
        self.inputs["nomor_surat"].insert(0, "170/DPRD/X/2026")
        self.inputs["nomor_surat_asn"].insert(0, "170/SEK-DPRD/X/2026")
        self.inputs["nomor_pemberitahuan_dprd"].insert(0, "180/DPRD/X/2026")
        self.inputs["nomor_pemberitahuan_asn"].insert(0, "181/SEK-DPRD/X/2026")
        self.inputs["nomor_spd_dprd"].insert(0, "090/SPD/")
        self.inputs["nomor_spd_asn"].insert(0, "091/SPD/")
        self.txt_materi_st.insert("1.0", "Studi Banding terkait Pembahasan Peraturan Daerah")
        self.txt_materi_pb.insert("1.0", "Pimpinan dan Anggota DPRD Kota Bitung akan melakukan Studi Banding terkait Pembahasan Peraturan Daerah")
        self.txt_dasar_dprd.insert("1.0", "Keputusan Pimpinan DPRD Kota Bitung")
        self.txt_dasar_asn.insert("1.0", "Surat Perintah Sekretaris DPRD Kota Bitung")

        self.tujuan_terpilih = ["Kota Manado"]
        self.refresh_tujuan_list_ui()

        if self.mode == "setwan":
            self.lbl_sign_dprd.pack_forget()
            self.combo_ttd_dprd.pack_forget()
            self.inputs["nomor_surat"].configure(placeholder_text="Masukkan nomor surat tugas setwan...")
            self.mode_specific_widgets["lbl_pemberitahuan_dprd"].pack_forget()
            self.mode_specific_widgets["ent_pemberitahuan_dprd"].pack_forget()
            self.mode_specific_widgets["lbl_spd_dprd"].grid_forget()
            self.mode_specific_widgets["ent_spd_dprd"].grid_forget()
            self.mode_specific_widgets["lbl_spd_setwan"].grid_forget()
            self.mode_specific_widgets["lbl_spd_info"].configure(text="ℹ️  SPD Pelaksana & Pendamping: nomor otomatis berurutan per kategori")
            self.lbl_spd_pelaksana.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
            self.ent_spd_pelaksana.grid(row=1, column=0, padx=(0, 4), sticky="ew")
            self.lbl_spd_pendamping.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
            self.ent_spd_pendamping.grid(row=1, column=1, padx=(4, 0), sticky="ew")
            self.combo_jenis.configure(values=["Studi Komparasi", "Kunjungan Konsultasi", "Bimbingan Teknis"])
        else:
            self.lbl_sign_dprd.pack(fill="x", padx=10, pady=(15, 2), before=self.lbl_sign_asn)
            self.combo_ttd_dprd.pack(fill="x", padx=10, pady=(0, 8), before=self.lbl_sign_asn)
            self.inputs["nomor_surat"].configure(placeholder_text="Masukkan nomor surat tugas dprd...")
            self.mode_specific_widgets["lbl_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(8, 2))
            self.mode_specific_widgets["ent_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(0, 6))
            self.mode_specific_widgets["lbl_spd_dprd"].grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
            self.mode_specific_widgets["ent_spd_dprd"].grid(row=1, column=0, padx=(0, 4), sticky="ew")
            self.mode_specific_widgets["lbl_spd_setwan"].grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
            self.mode_specific_widgets["lbl_spd_info"].configure(text="ℹ️  SPD DPRD: semua anggota pakai nomor sama  |  SPD ASN: nomor otomatis berurutan")
            self.lbl_spd_pelaksana.grid_forget()
            self.ent_spd_pelaksana.grid_forget()
            self.lbl_spd_pendamping.grid_forget()
            self.ent_spd_pendamping.grid_forget()
            self.combo_jenis.configure(values=["Kunjungan Kerja", "Kunjungan Konsultasi", "Bimbingan Teknis"])

        self.setup_category_checkboxes()
        self.refresh_personnel_list()

        # --------------------------------------------------------------
        # VALIDASI BERURUTAN: pelaksana wajib dipilih dulu (panel kanan),
        # baru bagian-bagian di bawah ini boleh diisi SATU PER SATU sesuai
        # urutan tampilan. Mencoba menyentuh bagian yang masih terkunci
        # akan memunculkan toast peringatan dan klik dibatalkan.
        # --------------------------------------------------------------
        self.step_lock.register(
            "pelaksana", "Pelaksana Perjalanan Dinas", None,
            self._has_pelaksana_selected,
        )
        self.step_lock.register(
            "nomor_surat", "Nomor Surat Tugas", self.sec_nomor_surat,
            lambda: bool(self.inputs["nomor_surat"].get().strip()) and bool(self.inputs["nomor_surat_asn"].get().strip()),
            widgets=[self.inputs["nomor_surat"], self.inputs["nomor_surat_asn"]],
        )
        self.step_lock.register(
            "pemberitahuan", "Nomor Surat Pemberitahuan", self.sec_pemberitahuan,
            lambda: (self.mode != "dprd" or bool(self.inputs["nomor_pemberitahuan_dprd"].get().strip()))
                    and bool(self.inputs["nomor_pemberitahuan_asn"].get().strip()),
            widgets=[self.inputs["nomor_pemberitahuan_dprd"], self.inputs["nomor_pemberitahuan_asn"]],
        )
        self.step_lock.register(
            "spd", "Nomor SPD", self.sec_spd,
            lambda: (bool(self.ent_spd_pelaksana.get().strip()) and bool(self.ent_spd_pendamping.get().strip()))
                    if self.mode == "setwan" else
                    (bool(self.inputs["nomor_spd_dprd"].get().strip()) and bool(self.inputs["nomor_spd_asn"].get().strip())),
            widgets=[self.inputs["nomor_spd_dprd"], self.inputs["nomor_spd_asn"], self.ent_spd_pelaksana, self.ent_spd_pendamping],
        )
        self.step_lock.register(
            "dasar", "Dasar Surat Tugas", self.sec_dasar,
            lambda: bool(self.txt_dasar_dprd.get("1.0", "end").strip()) and bool(self.txt_dasar_asn.get("1.0", "end").strip()),
            widgets=[self.txt_dasar_dprd, self.txt_dasar_asn],
        )
        self.step_lock.register(
            "materi", "Materi / Agenda Kegiatan", self.sec_materi,
            lambda: bool(self.txt_materi_st.get("1.0", "end").strip()) and bool(self.txt_materi_pb.get("1.0", "end").strip()),
            widgets=[self.txt_materi_st, self.txt_materi_pb],
        )
        self.step_lock.register(
            "jenis", "Jenis Perjalanan", self.sec_jenis,
            lambda: bool(self.combo_jenis.get().strip()),
            widgets=[self.combo_jenis],
        )
        self.step_lock.register(
            "tanggal", "Tanggal Surat / Mulai / Akhir Tugas", self.sec_tanggal,
            lambda: bool(self.dp_surat.get().strip()) and bool(self.dp_mulai.get().strip()) and bool(self.dp_akhir.get().strip()),
            widgets=[self.dp_surat, self.dp_mulai, self.dp_akhir],
        )
        self.step_lock.register(
            "tujuan", "Kota Tujuan Bertugas", self.sec_tujuan,
            lambda: len(self.tujuan_terpilih) > 0,
            widgets=[self.ent_tujuan, self.btn_tambah_tujuan],
        )
        self.step_lock.register(
            "ttd", "Penandatangan", self.sec_ttd,
            lambda: (self.mode != "dprd" or self.combo_ttd_dprd.get().strip() not in ("", "-"))
                    and self.combo_ttd_asn.get().strip() not in ("", "-"),
            widgets=[self.combo_ttd_dprd, self.combo_ttd_asn],
        )
        self.step_lock.refresh()


    def build_context(self, record_history=True):
        try:
            jml_angka = int(self.ent_lama_hari.get())
            jml_teks = self.terbilang(jml_angka)
        except:
            jml_angka = self.ent_lama_hari.get()
            jml_teks = "-"

        ttd_dprd_raw = self.combo_ttd_dprd.get() if hasattr(self, 'combo_ttd_dprd') else "-"
        ttd_asn_raw = self.combo_ttd_asn.get() if hasattr(self, 'combo_ttd_asn') else "-"
        jab_dprd, nama_dprd = ttd_dprd_raw.split(" - ", 1) if " - " in ttd_dprd_raw else ("KETUA", ttd_dprd_raw)
        jab_asn, nama_asn = ttd_asn_raw.split(" - ", 1) if " - " in ttd_asn_raw else ("SEKRETARIS DPRD", ttd_asn_raw)
        jab_dprd = format_jabatan_penandatanganan(jab_dprd)
        jab_asn = format_jabatan_penandatanganan(jab_asn)

        if self.mode == "dprd":
            selected_dprd = [p for p in self.db_dprd if self.dprd_vars.get((p.get('nama', ''), str(p.get('kategori', '')).strip())) and self.dprd_vars[(p.get('nama', ''), str(p.get('kategori', '')).strip())].get()]
            selected_asn = [p for p in self.db_asn if self.asn_vars.get(p.get('nama', '')) and self.asn_vars[p.get('nama', '')].get()]
            selected_pelaksana = []
            selected_pendamping = []
        else:
            selected_dprd = []
            selected_asn = []
            selected_pelaksana = [p for p in self.db_asn if self.pelaksana_vars.get(f"pelaksana_{p.get('nama', '')}") and self.pelaksana_vars[f"pelaksana_{p.get('nama', '')}"].get()]
            selected_pendamping = [p for p in self.db_asn if self.pendamping_vars.get(f"pendamping_{p.get('nama', '')}") and self.pendamping_vars[f"pendamping_{p.get('nama', '')}"].get()]

        if HAS_TKCALENDAR:
            tanggal_surat = self.format_indonesian_date(self.dp_surat.get_date())
            tanggal_mulai = self.format_indonesian_date(self.dp_mulai.get_date())
            tanggal_akhir = self.format_indonesian_date(self.dp_akhir.get_date())
        else:
            tanggal_surat = self.dp_surat.get()
            tanggal_mulai = self.dp_mulai.get()
            tanggal_akhir = self.dp_akhir.get()

        tujuan_list = self.tujuan_terpilih if self.tujuan_terpilih else ["(belum diisi)"]
        tujuan_str = " / ".join(tujuan_list)

        jenis_perjalanan = self.combo_jenis.get()
        materi_st = self.txt_materi_st.get("1.0", tk.END).strip()
        materi_pb = self.txt_materi_pb.get("1.0", tk.END).strip()
        dasar_surat_dprd = self.txt_dasar_dprd.get("1.0", tk.END).strip()
        dasar_surat_asn = self.txt_dasar_asn.get("1.0", tk.END).strip()

        nomor_spd_dprd = self.inputs["nomor_spd_dprd"].get()
        nomor_spd_asn = self.inputs["nomor_spd_asn"].get()
        nomor_spd_pelaksana = self.ent_spd_pelaksana.get() if self.mode == "setwan" else ""
        nomor_spd_pendamping = self.ent_spd_pendamping.get() if self.mode == "setwan" else ""

        city_names = [extract_city_name(d) for d in tujuan_list]
        transport = "Pesawat / Mobil / Kereta" if any(not is_in_sulawesi_utara(c) for c in city_names) else "Mobil"

        ctx = {
            "nomor_surat": self.inputs["nomor_surat"].get(),
            "nomor_surat_asn": self.inputs["nomor_surat_asn"].get(),
            "nomor_pemberitahuan_dprd": self.inputs["nomor_pemberitahuan_dprd"].get(),
            "nomor_pemberitahuan_asn": self.inputs["nomor_pemberitahuan_asn"].get(),
            "nomor_spd_dprd": nomor_spd_dprd,
            "nomor_spd_asn": nomor_spd_asn,
            "nomor_spd_pelaksana": nomor_spd_pelaksana,
            "nomor_spd_pendamping": nomor_spd_pendamping,
            "nomor_spd": nomor_spd_dprd,
            "tanggal_surat": tanggal_surat,
            "tanggal_surat_asn": tanggal_surat,
            "jenis_perjalanan": jenis_perjalanan,
            "tujuan_bertugas": tujuan_str,
            "tujuan_bertugas_list": tujuan_list,
            "dasar_surat_dprd": dasar_surat_dprd,
            "dasar_surat_asn": dasar_surat_asn,
            "materi_tugas": materi_st,
            "materi_tugas_asn": materi_st,
            "isi_surat_pemberitahuan": materi_pb,
            "isi_surat_izin": materi_st,
            "tanggal_mulai": tanggal_mulai,
            "tanggal_akhir": tanggal_akhir,
            "jumlah_angka": jml_angka,
            "jumlah_teks": jml_teks,
            "jabatan_ttd": jab_dprd.strip(),
            "nama_ttd": nama_dprd.strip(),
            "jabatan_ttd_asn": jab_asn.strip(),
            "nama_ttd_asn": nama_asn.strip(),
            "transportasi_otomatis": transport,
            "tanggal_surat_info": tanggal_surat,
            "tujuan_surat_info": tujuan_str,
            "pelaksana_tugas_dprd_info": "Pimpinan dan Anggota",
            "jenis_perjalanan_info": jenis_perjalanan,
            "tujuan_bertugas_info": tujuan_str,
            "materi_tugas_info": materi_pb,
            "hari_info": "Sesuai Jadwal",
            "tanggal_bertugas_info": f"{tanggal_mulai} s/d {tanggal_akhir}",
            "pelaksana_tugas_info": "Anggota DPRD",
            "jlh_pelaksana_dprd": len(selected_dprd),
            "pelaksana_tugas_asn_info": "Pendamping ASN" if self.mode == "dprd" else "Pelaksana ASN",
            "jlh_pelaksana_asn": len(selected_asn),
            "jlh_pelaksana": len(selected_pelaksana),
            "jlh_pendamping": len(selected_pendamping),
            "jabatan_ttd_info": jab_dprd.strip(),
            "nama_ttd_info": nama_dprd.strip(),
            "pelaksana_dprd": selected_dprd,
            "pelaksana_asn": selected_asn,
            "pelaksana_list": selected_pelaksana,
            "pendamping_list": selected_pendamping,
        }

        if record_history:
            # Ringkasan pelaksana tugas untuk kolom "Pelaksana Tugas" di
            # tabel Riwayat Surat (maks. 3 nama ditampilkan + "dst.").
            if self.mode == "dprd":
                nama_pelaksana = [p.get("nama", "") for p in selected_dprd] + [p.get("nama", "") for p in selected_asn]
            else:
                nama_pelaksana = [p.get("nama", "") for p in selected_pelaksana] + [p.get("nama", "") for p in selected_pendamping]
            nama_pelaksana = [n for n in nama_pelaksana if n]
            if len(nama_pelaksana) > 3:
                pelaksana_display = ", ".join(nama_pelaksana[:3]) + f", dan {len(nama_pelaksana) - 3} lainnya"
            elif nama_pelaksana:
                pelaksana_display = ", ".join(nama_pelaksana)
            else:
                pelaksana_display = "-"

            nama_surat_display = f"Surat Tugas & SPD - {jenis_perjalanan or 'Perjalanan Dinas'} - {tujuan_str}"

            # Kategori pelaksana kegiatan, dipakai kolom "Kategori Pelaksana"
            # di menu Judul Perjadin -- ambil dari kategori masing-masing
            # anggota DPRD terpilih (mis. "Pimpinan DPRD", "Anggota"), atau
            # label peran Pelaksana/Pendamping ASN untuk mode Setwan.
            if self.mode == "dprd":
                kategori_set = {str(p.get("kategori", "")).strip() for p in selected_dprd if p.get("kategori")}
                if selected_asn:
                    kategori_set.add("Pendamping ASN")
                kategori_pelaksana_display = ", ".join(sorted(kategori_set)) or "-"
            else:
                bagian = []
                if selected_pelaksana:
                    bagian.append("Pelaksana ASN")
                if selected_pendamping:
                    bagian.append("Pendamping ASN")
                kategori_pelaksana_display = ", ".join(bagian) or "-"

            tanggal_pelaksanaan_display = format_rentang_tanggal(tanggal_mulai, tanggal_akhir)

            # --- BUGFIX PENTING --------------------------------------------
            # Dulu key riwayat = ctx["nomor_surat"] (teks yang diketik
            # pengguna) apa adanya. Masalahnya field ini SELALU terisi nilai
            # CONTOH default ("170/DPRD/X/2026", lihat setup_perjalanan_dinas_form)
            # setiap kali formulir dibuka -- kalau pengguna lupa
            # menggantinya (sangat sering terjadi utk surat mode Setwan yang
            # nomor DPRD-nya memang tidak relevan), maka SETIAP surat baru
            # tersimpan dengan key yang SAMA dan MENIMPA riwayat surat
            # sebelumnya alih-alih menambah entri baru. Akibatnya: tabel
            # "Judul Perjadin" & "Riwayat Surat" cuma menampilkan 1 baris
            # (yang tersisa cuma surat TERAKHIR), dan pengecekan judul
            # duplikat pun "tidak berfungsi" karena riwayat lamanya sudah
            # keburu tertimpa sebelum sempat dibandingkan.
            #
            # Perbaikan: key riwayat sekarang ID internal STABIL yang TIDAK
            # PERNAH bergantung pada teks nomor surat:
            #   - Kalau formulir sedang MENGEDIT riwayat lama (dibuka lewat
            #     tombol Edit di Riwayat Surat -> _editing_perjalanan_key
            #     terisi), pakai key yang sama itu supaya hasil edit
            #     MENIMPA record aslinya (bukan surat baru berduplikat).
            #   - Kalau ini surat BARU, buat UUID baru sekali, lalu SIMPAN
            #     ke _editing_perjalanan_key supaya kalau tombol cetak
            #     ditekan lagi tanpa pindah halaman (mis. cetak ulang PDF
            #     surat yang sama), tetap menimpa record yang sama, bukan
            #     membuat entri baru lagi.
            # Nomor surat yang diketik pengguna TETAP disimpan seperti biasa
            # sebagai field "nomor_surat" biasa di dalam record (dipakai
            # utk ditampilkan & utk validasi anti-duplikat nomor surat) --
            # yang berubah HANYA cara menentukan KEY dict-nya.
            record_key = self._editing_perjalanan_key or uuid.uuid4().hex
            self._editing_perjalanan_key = record_key

            self.history_data.setdefault("perjalanan_dinas", {})[record_key] = {
                "nomor_surat": ctx["nomor_surat"],
                "nomor_surat_asn": ctx["nomor_surat_asn"],
                "nomor_pemberitahuan_dprd": ctx["nomor_pemberitahuan_dprd"],
                "nomor_pemberitahuan_asn": ctx["nomor_pemberitahuan_asn"],
                "nomor_spd_dprd": nomor_spd_dprd,
                "nomor_spd_asn": nomor_spd_asn,
                "jenis_perjalanan": ctx["jenis_perjalanan"],
                "tujuan_bertugas": tujuan_str,
                "tujuan_bertugas_list": tujuan_list,
                "materi_tugas": materi_st,
                "materi_tugas_pb": materi_pb,
                "dasar_surat_dprd": dasar_surat_dprd,
                "dasar_surat_asn": dasar_surat_asn,
                "dprd_terpilih": [f"{p.get('nama')}||{p.get('kategori')}" for p in selected_dprd],
                "asn_terpilih": [p.get('nama') for p in selected_asn],
                "pelaksana_terpilih": [p.get('nama') for p in selected_pelaksana] if self.mode == "setwan" else [],
                "pendamping_terpilih": [p.get('nama') for p in selected_pendamping] if self.mode == "setwan" else [],
                "nama_surat": nama_surat_display,
                "tanggal_surat": tanggal_surat,
                "tanggal_mulai": tanggal_mulai,
                "tanggal_akhir": tanggal_akhir,
                "tanggal_pelaksanaan_display": tanggal_pelaksanaan_display,
                "kategori_pelaksana_display": kategori_pelaksana_display,
                "pelaksana_display": pelaksana_display,
                "dibuat_oleh": getattr(self, "_current_user", "-"),
                "tanggal_dibuat": datetime.now().strftime("%d-%m-%Y %H:%M"),
            }
            self.save_history()
            self._refresh_riwayat_window_if_open()

        return ctx, selected_dprd, selected_asn, selected_pelaksana, selected_pendamping

    def generate_documents_action(self):
        if self.current_view == "perjalanan_dinas" and hasattr(self, "step_lock"):
            if self.step_lock.notify_if_incomplete():
                return

        # Cek nomor surat terhadap RIWAYAT BERSAMA (bisa berisi surat yang
        # dibuat komputer lain kalau mode jaringan aktif) SEBELUM data
        # ditulis ke riwayat oleh build_context(). Nomor yang sudah pernah
        # dipakai untuk surat lain ditolak -- kecuali memang sedang
        # mengedit/merevisi surat yang nomornya itu sendiri (nomor lama
        # yang sama masih boleh dipakai ulang untuk surat yang sama).
        # Dikecualikan dari pengecekan: KEY internal riwayat yang sedang
        # diedit (bukan teks nomor surat -- lihat catatan di build_context).
        # None kalau ini surat baru, jadi tidak mengecualikan apa pun.
        nomor_induk_form = getattr(self, "_editing_perjalanan_key", None)
        field_nomor_untuk_dicek = [
            ("Nomor Surat Tugas DPRD", self.inputs["nomor_surat"].get().strip()),
            ("Nomor Surat Tugas Setwan", self.inputs["nomor_surat_asn"].get().strip()),
            ("Nomor Surat Pemberitahuan DPRD", self.inputs["nomor_pemberitahuan_dprd"].get().strip()),
            ("Nomor Surat Pemberitahuan Setwan", self.inputs["nomor_pemberitahuan_asn"].get().strip()),
            ("Nomor SPD DPRD", self.inputs["nomor_spd_dprd"].get().strip()),
            ("Nomor SPD Setwan", self.inputs["nomor_spd_asn"].get().strip()),
        ]
        for label, nomor in field_nomor_untuk_dicek:
            if not nomor:
                continue
            try:
                sudah_dipakai = nomor_surat_sudah_dipakai(nomor, kecuali_nomor_induk=nomor_induk_form)
            except Exception:
                sudah_dipakai = False  # jangan blokir pengguna hanya karena gagal mengecek riwayat
            if sudah_dipakai:
                messagebox.showerror(
                    "Nomor Surat Sudah Dipakai",
                    f"{label} '{nomor}' SUDAH PERNAH digunakan untuk surat lain "
                    "(bisa jadi dibuat dari komputer lain).\n\n"
                    "Nomor yang sudah dipakai hanya boleh dipakai lagi untuk "
                    "mengedit/merevisi surat yang SAMA -- gunakan menu Riwayat "
                    "Surat untuk membuka surat tsb, atau ganti nomornya.",
                )
                return

        ctx, sel_dprd, sel_asn, sel_pelaksana, sel_pendamping = self.build_context()
        if self.mode == "dprd":
            if not sel_dprd and not sel_asn:
                messagebox.showwarning("Peringatan", "Pilih minimal satu pelaksana!")
                return
        else:
            if not sel_pelaksana and not sel_pendamping:
                messagebox.showwarning("Peringatan", "Pilih minimal satu Pelaksana ASN atau Pendamping ASN!")
                return
        if not self.tujuan_terpilih:
            messagebox.showwarning("Peringatan", "Tambahkan minimal satu kota tujuan bertugas!")
            return

        out_dir = filedialog.askdirectory(title="Pilih Folder untuk Menyimpan Dokumen")
        if not out_dir: return

        # Jalankan proses cetak (I/O berat: docxtpl/python-docx) di background
        # thread supaya jendela aplikasi tidak freeze ("Not Responding") saat
        # mencetak banyak dokumen sekaligus. UI hanya boleh disentuh lewat
        # self.after(...) dari thread ini.
        progress = self._show_progress_dialog("Mencetak Dokumen", "Mempersiapkan dokumen...")
        self.btn_generate_main.configure(state="disabled")

        def worker():
            try:
                success_count, missing, errors = self._generate_documents_worker(
                    ctx, sel_dprd, sel_asn, sel_pelaksana, sel_pendamping, out_dir,
                    progress_cb=lambda msg: self.after(0, lambda: self._update_progress_dialog(progress, msg)),
                )
            except Exception as e:
                self.after(0, lambda: self._finish_generate_documents(progress, 0, [], out_dir, error=str(e)))
                return
            self.after(0, lambda: self._finish_generate_documents(progress, success_count, missing, out_dir, errors=errors))

        threading.Thread(target=worker, daemon=True).start()

    def _show_progress_dialog(self, title, message):
        dlg = ctk.CTkToplevel(self)
        dlg.title(title)
        dlg.geometry("360x120")
        dlg.resizable(False, False)
        dlg.transient(self)
        dlg.grab_set()
        dlg.protocol("WM_DELETE_WINDOW", lambda: None)  # cegah ditutup paksa saat proses jalan
        label = ctk.CTkLabel(dlg, text=message, wraplength=320)
        label.pack(pady=(20, 10), padx=20)
        bar = ctk.CTkProgressBar(dlg, mode="indeterminate", width=300)
        bar.pack(pady=10, padx=20)
        bar.start()
        dlg.update_idletasks()
        # Posisikan di tengah window utama
        x = self.winfo_x() + (self.winfo_width() // 2) - 180
        y = self.winfo_y() + (self.winfo_height() // 2) - 60
        dlg.geometry(f"+{max(x, 0)}+{max(y, 0)}")
        dlg._label = label
        return dlg

    def _update_progress_dialog(self, dlg, message):
        try:
            dlg._label.configure(text=message)
        except Exception:
            pass

    def _finish_generate_documents(self, dlg, success_count, missing, out_dir, error=None, errors=None):
        try:
            dlg.grab_release()
            dlg.destroy()
        except Exception:
            pass
        self.btn_generate_main.configure(state="normal")

        if error:
            messagebox.showerror("Error", f"Terjadi kesalahan saat mencetak dokumen:\n{error}")
            return

        self._refresh_riwayat_window_if_open()

        if success_count > 0:
            info_lines = [f"Berhasil mencetak {success_count} file dokumen di:\n{out_dir}\n"]
            if missing: info_lines.append("\nTemplate tidak ditemukan:\n" + "\n".join(missing))
            if errors: info_lines.append("\nSebagian dokumen GAGAL dibuat:\n" + "\n".join(errors))
            messagebox.showinfo("Cetak Berhasil", "\n".join(info_lines))
        else:
            pesan = "Gagal mencetak. Pastikan file template tersedia dan ada pelaksana terpilih."
            if errors: pesan += "\n\nDetail error:\n" + "\n".join(errors)
            messagebox.showerror("Error", pesan)

    def _generate_documents_worker(self, ctx, sel_dprd, sel_asn, sel_pelaksana, sel_pendamping, out_dir, progress_cb):
        success_count = 0
        missing = []
        # PENTING: dulu kegagalan tiap jenis dokumen hanya di-print() lalu
        # dilupakan -- di build .exe mode windowed (console=False) print()
        # bisa melempar exception baru dan MEMBUNGKAM error asli, sehingga
        # dokumen gagal dibuat TANPA jejak sama sekali bagi pengguna (cuma
        # kelihatan sebagai jumlah file yang lebih sedikit dari yang
        # diharapkan). Sekarang tiap kegagalan dicatat ke sini supaya bisa
        # ditampilkan di dialog "Cetak Berhasil/Gagal" di akhir proses.
        errors = []

        def catat_gagal(label, exc):
            pesan = f"{label}: {exc}"
            safe_log(pesan)
            errors.append(pesan)

        if self.mode == "dprd":
            if sel_dprd:
                template_st_dprd = TEMPLATE_ST_DPRD_BIASA if len(sel_dprd) <= 3 else TEMPLATE_ST_DPRD_TABEL
                if os.path.exists(template_st_dprd):
                    try:
                        out_path = os.path.join(out_dir, self._nama_file_perjalanan("surat-tugas") + ".docx")
                        buat_surat_tugas_dprd(ctx, sel_dprd, out_path)
                        success_count += 1
                    except Exception as e:
                        catat_gagal("Gagal buat surat tugas DPRD", e)
                else: missing.append(template_st_dprd)
            if sel_asn:
                template_st_asn = TEMPLATE_ST_ASN_BIASA if len(sel_asn) <= 3 else TEMPLATE_ST_ASN_TABEL
                if os.path.exists(template_st_asn):
                    try:
                        out_path = os.path.join(out_dir, self._nama_file_perjalanan("surat-tugas-pendamping") + ".docx")
                        buat_surat_tugas_asn(ctx, sel_asn, out_path)
                        success_count += 1
                    except Exception as e:
                        catat_gagal("Gagal buat surat tugas ASN", e)
                else: missing.append(template_st_asn)
            if os.path.exists(TEMPLATE_PEMBERITAHUAN):
                try:
                    base_number = ctx.get("nomor_pemberitahuan_dprd", ctx.get("nomor_surat", ""))
                    out_path = os.path.join(out_dir, self._nama_file_perjalanan("surat-pemberitahuan") + ".docx")
                    buat_surat_pemberitahuan_multi(TEMPLATE_PEMBERITAHUAN, ctx, sel_dprd, sel_asn, self.tujuan_terpilih, base_number, out_path, label_asn="Pendamping ASN")
                    success_count += 1
                except Exception as e:
                    catat_gagal("Gagal buat pemberitahuan", e)
            else: missing.append(TEMPLATE_PEMBERITAHUAN)
            if sel_dprd:
                try:
                    buat_sppd_dprd(TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, ctx, sel_dprd, self.tujuan_terpilih, os.path.join(out_dir, self._nama_file_perjalanan("spd-depan") + ".docx"), os.path.join(out_dir, self._nama_file_perjalanan("spd-belakang") + ".docx"))
                    success_count += 2
                except Exception as e:
                    catat_gagal("Gagal buat SPD DPRD", e)
            if sel_asn:
                try:
                    buat_sppd_asn(TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, ctx, sel_asn, self.tujuan_terpilih, os.path.join(out_dir, self._nama_file_perjalanan("spd-pendamping-depan") + ".docx"), os.path.join(out_dir, self._nama_file_perjalanan("spd-pendamping-belakang") + ".docx"))
                    success_count += 2
                except Exception as e:
                    catat_gagal("Gagal buat SPD ASN", e)
        else:
            sel_asn = sel_pelaksana + sel_pendamping
            if sel_pelaksana:
                template_st_asn = TEMPLATE_ST_ASN_BIASA if len(sel_pelaksana) <= 3 else TEMPLATE_ST_ASN_TABEL
                if os.path.exists(template_st_asn):
                    try:
                        ctx_pelaksana = ctx.copy()
                        ctx_pelaksana["pelaksana_tugas_asn_info"], ctx_pelaksana["jlh_pelaksana_asn"] = "Pelaksana ASN", len(sel_pelaksana)
                        buat_surat_tugas_asn(ctx_pelaksana, sel_pelaksana, os.path.join(out_dir, self._nama_file_perjalanan("surat-tugas-pelaksana") + ".docx"))
                        success_count += 1
                    except Exception as e: catat_gagal("Gagal buat surat tugas Pelaksana ASN", e)
                else: missing.append(template_st_asn)
            if sel_pendamping:
                template_st_asn = TEMPLATE_ST_ASN_BIASA if len(sel_pendamping) <= 3 else TEMPLATE_ST_ASN_TABEL
                if os.path.exists(template_st_asn):
                    try:
                        ctx_pendamping = ctx.copy()
                        ctx_pendamping["pelaksana_tugas_asn_info"], ctx_pendamping["jlh_pelaksana_asn"] = "Pendamping ASN", len(sel_pendamping)
                        buat_surat_tugas_asn(ctx_pendamping, sel_pendamping, os.path.join(out_dir, self._nama_file_perjalanan("surat-tugas-pendamping") + ".docx"))
                        success_count += 1
                    except Exception as e: catat_gagal("Gagal buat surat tugas Pendamping ASN", e)
                else: missing.append(template_st_asn)
            if os.path.exists(TEMPLATE_PEMBERITAHUAN):
                base_number = ctx.get("nomor_pemberitahuan_asn", ctx.get("nomor_surat_asn", ""))
                if sel_pelaksana:
                    try:
                        buat_surat_pemberitahuan_multi(TEMPLATE_PEMBERITAHUAN, ctx, [], sel_pelaksana, self.tujuan_terpilih, base_number, os.path.join(out_dir, self._nama_file_perjalanan("surat-pemberitahuan-pelaksana") + ".docx"), label_asn="Pelaksana ASN")
                        success_count += 1
                    except Exception as e: catat_gagal("Gagal buat pemberitahuan Pelaksana", e)
                if sel_pendamping:
                    try:
                        buat_surat_pemberitahuan_multi(TEMPLATE_PEMBERITAHUAN, ctx, [], sel_pendamping, self.tujuan_terpilih, base_number, os.path.join(out_dir, self._nama_file_perjalanan("surat-pemberitahuan-pendamping") + ".docx"), label_asn="Pendamping ASN")
                        success_count += 1
                    except Exception as e: catat_gagal("Gagal buat pemberitahuan Pendamping", e)
            else: missing.append(TEMPLATE_PEMBERITAHUAN)
            if sel_pelaksana:
                ctx_pelaksana_spd = ctx.copy()
                ctx_pelaksana_spd["nomor_spd_asn"] = ctx.get("nomor_spd_pelaksana", ctx.get("nomor_spd_asn", ""))
                try:
                    buat_sppd_asn(TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, ctx_pelaksana_spd, sel_pelaksana, self.tujuan_terpilih, os.path.join(out_dir, self._nama_file_perjalanan("spd-pelaksana-depan") + ".docx"), os.path.join(out_dir, self._nama_file_perjalanan("spd-pelaksana-belakang") + ".docx"))
                    success_count += 2
                except Exception as e: catat_gagal("Gagal buat SPD Pelaksana ASN", e)
            if sel_pendamping:
                ctx_pendamping_spd = ctx.copy()
                ctx_pendamping_spd["nomor_spd_asn"] = ctx.get("nomor_spd_pendamping", ctx.get("nomor_spd_asn", ""))
                try:
                    buat_sppd_asn(TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, ctx_pendamping_spd, sel_pendamping, self.tujuan_terpilih, os.path.join(out_dir, self._nama_file_perjalanan("spd-pendamping-depan") + ".docx"), os.path.join(out_dir, self._nama_file_perjalanan("spd-pendamping-belakang") + ".docx"))
                    success_count += 2
                except Exception as e: catat_gagal("Gagal buat SPD Pendamping ASN", e)

        try:
            out_daftar = os.path.join(out_dir, self._nama_file_perjalanan("daftar-hadir") + ".docx")
            pelaksana_daftar = sel_dprd if self.mode == "dprd" else sel_pelaksana + sel_pendamping
            if pelaksana_daftar:
                self.buat_daftar_hadir(ctx, pelaksana_daftar, self.tujuan_terpilih, self.mode, out_daftar)
                success_count += 1
        except Exception as e: catat_gagal("Gagal buat daftar hadir", e)



        return success_count, missing, errors

    def buat_daftar_hadir(self, ctx, pelaksana_list, destinations, mode, out_path):
        template_path = TEMPLATE_DAFTAR_HADIR
        if not os.path.exists(template_path): raise FileNotFoundError(f"Template {template_path} tidak ditemukan.")
        periods = generate_periods(ctx.get("tanggal_mulai", ""), destinations)
        jenis_perjalanan = ctx.get("jenis_perjalanan", "").strip()
        temp_files = []
        for period in periods:
            pelaku_str, instansi_tujuan = self.build_judul_daftar_hadir(pelaksana_list, period["tujuan"], mode)
            materi_upper = strip_jenis_perjalanan_prefix(ctx.get("materi_tugas", ""), jenis_perjalanan).upper().strip()
            zona = detect_zona_waktu(period["tujuan"])
            render_ctx = {
                "pelaksana_tugas_dprd": pelaku_str,
                "JENIS_PERJALANAN_DAFTAR_HADIR": jenis_perjalanan.upper(),
                "TEMPAT_TUGAS_DPRD_DAFTAR_HADIR": instansi_tujuan,
                "MATERI_TUGAS_DPRD_DAFTAR_HADIR": materi_upper,
                "HARI": period["hari"],
                "TANGGAL_DAFTAR_HADIR": period["tanggal"],
                "TEMPAT_DAFTAR_HADIR": period["tujuan"],
                "zona": zona,
                "loop": {"index": ""}, "tabel": {"NAMA_DAFTAR_HADIR": "", "jabatan_daftar_hadir": ""}
            }
            doc_tpl = DocxTemplate(template_path)
            doc_tpl.render(render_ctx)
            tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
            doc_tpl.save(tmp_file)
            doc = Document(tmp_file)
            rows_data = [[str(i + 1), p.get('nama', ''), p.get('jabatan', '')] for i, p in enumerate(pelaksana_list)]
            # max_tables=1: HANYA tabel format keberangkatan yang diisi nama
            # & jabatan. Tabel format tempat tugas SENGAJA dibiarkan kosong
            # sesuai master template -- akan ditandatangani manual di
            # tempat tugas, bukan diisi otomatis oleh program.
            _fill_table_rows_from_master(doc, ["no", "nama", "jabatan", "tanda tangan"], rows_data, max_tables=1)
            _force_daftar_hadir_page_break(doc)
            doc.save(tmp_file)
            temp_files.append(tmp_file)
        if temp_files: _combine_word_pages(temp_files, out_path)
        for f in temp_files:
            try: os.unlink(f)
            except: pass

    def build_judul_daftar_hadir(self, pelaksana_list, tujuan, mode):
        """Mengembalikan nilai ATOMIK (pelaku_str, instansi_tujuan) -- bukan
        kalimat yang sudah digabung -- karena penyusunan kalimat
        "PADA ... KE ... ..." sekarang dilakukan di dalam template Word itu
        sendiri (lihat DAFTAR_HADIR_DPRD.docx).

        pelaku_str dibangun per kelompok kategori (Pimpinan DPRD, Komisi I/II/III, dst):
        - Kelompok "Pimpinan DPRD" dengan HANYA 1 orang -> pakai jabatan spesifik
          orang itu apa adanya (mis. "Wakil Ketua DPRD Kota Bitung",
          "Ketua DPRD Kota Bitung"), TIDAK ditambah kata "Pimpinan" lagi
          supaya tidak dobel seperti "Pimpinan Pimpinan DPRD".
        - Kelompok "Pimpinan DPRD" dengan LEBIH dari 1 orang (mis. gabungan
          Ketua + Wakil Ketua) -> digeneralisasi jadi "Pimpinan DPRD Kota Bitung".
        - Kelompok Komisi (Komisi I/II/III) -> dideteksi otomatis dari jabatan
          tiap orang: jika ada yang berjabatan ketua/wakil ketua/sekretaris
          komisi DAN ada yang anggota -> "Pimpinan dan Anggota Komisi X";
          jika hanya pimpinan komisi -> "Pimpinan Komisi X"; jika hanya
          anggota -> "Anggota Komisi X".
        - Jika pelaksana berasal dari LEBIH dari satu kelompok kategori
          (mis. Pimpinan DPRD bersama Komisi I), label tiap kelompok
          digabung dengan kata "BERSAMA", mis. "WAKIL KETUA DPRD KOTA
          BITUNG BERSAMA PIMPINAN DAN ANGGOTA KOMISI I DPRD KOTA BITUNG".
        """
        # Kelompokkan pelaksana per kategori, urut sesuai kemunculan pertama.
        groups, order = {}, []
        for p in pelaksana_list:
            kat = p.get('kategori', '').strip() or "Lainnya"
            if kat not in groups:
                groups[kat] = []
                order.append(kat)
            groups[kat].append(p)

        def label_for_group(kategori, members):
            if kategori == "Pimpinan DPRD":
                if len(members) == 1:
                    # 1 orang pimpinan -> pakai jabatan spesifiknya apa adanya
                    jabatan = members[0].get('jabatan', '').strip()
                    return jabatan if jabatan else "Pimpinan DPRD Kota Bitung"
                # lebih dari 1 orang pimpinan (mis. Ketua + Wakil Ketua) -> generik
                return "Pimpinan DPRD Kota Bitung"
            # Kategori lain (Komisi I/II/III, dll) -> deteksi otomatis dari jabatan
            jabatans_lower = [m.get('jabatan', '').lower() for m in members]
            has_pimpinan = any(("ketua" in j or "sekretaris" in j) for j in jabatans_lower)
            has_anggota = any("anggota" in j for j in jabatans_lower)
            if has_pimpinan and has_anggota: label = f"Pimpinan dan Anggota {kategori}"
            elif has_pimpinan: label = f"Pimpinan {kategori}"
            elif has_anggota: label = f"Anggota {kategori}"
            else: label = kategori
            if "DPRD" not in label.upper(): label = f"{label} DPRD Kota Bitung"
            return label

        labels = [label_for_group(kat, groups[kat]) for kat in order]
        if len(labels) == 1:
            pelaku_str = labels[0].upper()
        elif len(labels) <= 3:
            pelaku_str = " BERSAMA ".join(labels).upper()
        else:
            pelaku_str = "PIMPINAN DAN ANGGOTA DPRD KOTA BITUNG"

        # Tujuan bisa berupa nama kota/kabupaten/provinsi polos (mis. "Kota
        # Bekasi") -> berarti maksudnya DPRD setempat, jadi perlu diberi
        # awalan "DPRD ". TAPI kalau tujuan sudah berupa nama instansi
        # spesifik (mis. "BPBD Provinsi DKI Jakarta", "Universitas Sam
        # Ratulangi") -> pakai apa adanya, JANGAN ditambah "DPRD" di depan.
        # PENTING: keputusan ini HARUS berdasarkan apakah seluruh teks
        # tujuan murni nama wilayah administratif (is_plain_region_name),
        # bukan sekadar membandingkan hasil extract_city_name dengan teks
        # aslinya -- karena nama instansi yang KEBETULAN tidak mengandung
        # kata Kota/Kabupaten/Provinsi (mis. "Universitas Sam Ratulangi")
        # akan lolos perbandingan tersebut dan salah ditempeli "DPRD ".
        tujuan_stripped = tujuan.strip()
        if is_plain_region_name(tujuan_stripped):
            instansi_tujuan = tujuan_stripped if "DPRD" in tujuan_stripped.upper() else f"DPRD {tujuan_stripped}"
        else:
            # tujuan sudah berupa nama instansi spesifik -> pakai apa adanya
            instansi_tujuan = tujuan_stripped
        return pelaku_str, instansi_tujuan

    # ------------------------------------------------------------------
    # LIVE PREVIEW LOGIC
    # ------------------------------------------------------------------
    def schedule_preview_refresh(self, *args, immediate=False, **kwargs):
        # Refresh status kunci-berurutan formulir Perjalanan Dinas setiap
        # kali ada perubahan input apa pun (ini dipanggil dari hampir semua
        # event handler widget) -- disengaja SINKRON (bukan didebounce
        # seperti render preview di bawah) supaya overlay kunci terbuka /
        # tertutup terasa instan mengikuti kecepatan mengetik/klik pengguna.
        if self.current_view == "perjalanan_dinas" and hasattr(self, "step_lock"):
            try:
                self.step_lock.refresh()
            except Exception:
                pass
        if not hasattr(self, "combo_preview_jenis"): return
        if self._preview_after_id is not None:
            try: self.after_cancel(self._preview_after_id)
            except Exception: pass
            self._preview_after_id = None
        delay = 100 if immediate else 700
        self._preview_after_id = self.after(delay, self._launch_preview_render)

    def _launch_preview_render(self):
        self._preview_after_id = None
        start_thread = False
        with self._preview_lock:
            if self._preview_busy: self._preview_pending = True
            else:
                self._preview_busy = True
                start_thread = True
        if start_thread:
            threading.Thread(target=self._preview_worker, daemon=True).start()

    def _set_preview_status(self, text, color="gray"):
        def _do():
            try: self.preview_status_lbl.configure(text=text, text_color=color)
            except Exception: pass
        self.after(0, _do)

    def _preview_worker(self):
        try:
            label = self.combo_preview_jenis.get()
            template_file, mode = None, "ctx"
            for name, fname, m in PREVIEW_TEMPLATES:
                if name == label:
                    template_file, mode = fname, m
                    break

            if not template_file:
                self._set_preview_status("Pilih jenis surat untuk pratinjau.")
                return
            if template_file not in ("__surat_tugas_dprd__", "__surat_tugas_asn__", "__daftar_hadir__", "__undangan_paripurna__", "__undangan_biasa__") and not os.path.exists(template_file):
                self._set_preview_status(f"Template tidak ditemukan: {template_file}")
                return

            self._set_preview_status("Memproses preview...")

            tmp_docx = os.path.join(self.preview_dir, "preview_render.docx")
            tmp_pdf = os.path.join(self.preview_dir, "preview_render.pdf")

            # --- PREVIEW UNDANGAN PARIPURNA ---
            if self.current_view == "undangan_paripurna" and template_file == "__undangan_paripurna__":
                template_original = TEMPLATE_PARIPURNA
                if not os.path.exists(template_original):
                    self._set_preview_status(f"Template tidak ditemukan: {template_original}")
                    return
                
                # Gunakan context generator yang sudah mencakup penomoran Skenario
                ctx = self.get_undangan_context()
                
                # Render template asli seutuhnya (tanpa memotong halaman yang merusak format)
                doc = DocxTemplate(template_original)
                doc.render(ctx)
                doc.save(tmp_docx)
                # Hapus paragraf skenario kosong & pasang page-break eksplisit
                # antar-blok, supaya preview cerminan persis hasil cetak akhir.
                doc_clean = Document(tmp_docx)
                cleanup_skenario_paripurna(doc_clean)
                _force_paripurna_page_break(doc_clean)
                doc_clean.save(tmp_docx)

            # --- PREVIEW UNDANGAN RAPAT BIASA ---
            elif self.current_view == "undangan_biasa" and template_file == "__undangan_biasa__":
                template_original = TEMPLATE_RAPAT_BIASA
                if not os.path.exists(template_original):
                    self._set_preview_status(f"Template tidak ditemukan: {template_original}")
                    return

                ctx = self.get_undangan_biasa_context()
                self._render_undangan_biasa_docx(ctx, tmp_docx)

            # --- PREVIEW PERJALANAN DINAS ---
            else:
                if self.current_view != "perjalanan_dinas":
                    self._set_preview_status("Preview hanya tersedia di mode Perjalanan Dinas.")
                    return
                ctx, sel_dprd, sel_asn, sel_pelaksana, sel_pendamping = self.build_context(record_history=False)

                if mode == "person_dprd":
                    if not sel_dprd:
                        self._set_preview_status("Pilih minimal satu anggota DPRD untuk melihat pratinjau SPD DPRD.")
                        return
                    # Render semua orang (bukan hanya sel_dprd[0]) supaya
                    # preview menampilkan seluruh halaman seperti dokumen asli.
                    from app.core.sppd_generators import buat_sppd_dprd, buat_sppd_asn
                    out_belakang_tmp = os.path.join(self.preview_dir, "preview_spd_belakang_dummy.docx")
                    if template_file == TEMPLATE_SPD_DEPAN:
                        out_depan_tmp = tmp_docx
                        buat_sppd_dprd(TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, ctx, sel_dprd,
                                       self.tujuan_terpilih or ["Manado"],
                                       out_depan_tmp, out_belakang_tmp)
                    else:  # SPD Belakang
                        out_depan_tmp = os.path.join(self.preview_dir, "preview_spd_depan_dummy.docx")
                        buat_sppd_dprd(TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, ctx, sel_dprd,
                                       self.tujuan_terpilih or ["Manado"],
                                       out_depan_tmp, tmp_docx)
                elif mode == "person_asn":
                    if not sel_asn:
                        self._set_preview_status("Pilih minimal satu ASN untuk melihat pratinjau SPD ASN.")
                        return
                    from app.core.sppd_generators import buat_sppd_asn
                    out_belakang_tmp = os.path.join(self.preview_dir, "preview_spd_belakang_dummy.docx")
                    if template_file == TEMPLATE_SPD_DEPAN:
                        buat_sppd_asn(TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, ctx, sel_asn,
                                      self.tujuan_terpilih or ["Manado"],
                                      tmp_docx, out_belakang_tmp)
                    else:  # SPD Belakang
                        out_depan_tmp = os.path.join(self.preview_dir, "preview_spd_depan_dummy.docx")
                        buat_sppd_asn(TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, ctx, sel_asn,
                                      self.tujuan_terpilih or ["Manado"],
                                      out_depan_tmp, tmp_docx)
                elif template_file == "__surat_tugas_dprd__":
                    if not sel_dprd:
                        self._set_preview_status("Pilih minimal satu pelaksana DPRD untuk pratinjau Surat Tugas DPRD.")
                        return
                    buat_surat_tugas_dprd(ctx, sel_dprd, tmp_docx)
                elif template_file == "__surat_tugas_asn__":
                    if not sel_asn:
                        self._set_preview_status("Pilih minimal satu pendamping ASN untuk pratinjau Surat Tugas ASN.")
                        return
                    buat_surat_tugas_asn(ctx, sel_asn, tmp_docx)
                elif template_file == TEMPLATE_PEMBERITAHUAN:
                    if not self.tujuan_terpilih:
                        self._set_preview_status("Tambahkan minimal satu tujuan untuk pratinjau pemberitahuan.")
                        return
                    base_number = ctx.get("nomor_pemberitahuan_dprd", ctx.get("nomor_surat", ""))
                    pelaksana_dprd = sel_dprd if self.mode == "dprd" else []
                    label_asn = "Pendamping ASN" if self.mode == "dprd" else "Pelaksana Tugas ASN"
                    # Gunakan SEMUA tujuan (bukan hanya [:1]) supaya preview
                    # menampilkan seluruh halaman seperti dokumen asli.
                    buat_surat_pemberitahuan_multi(template_file, ctx, pelaksana_dprd, sel_asn,
                                                   self.tujuan_terpilih, base_number, tmp_docx,
                                                   label_asn=label_asn)
                elif template_file == "__daftar_hadir__":
                    pelaksana = sel_dprd if self.mode == "dprd" else sel_asn
                    if not pelaksana:
                        self._set_preview_status("Pilih minimal satu pelaksana untuk pratinjau daftar hadir.")
                        return
                    if not self.tujuan_terpilih:
                        self._set_preview_status("Tambahkan minimal satu tujuan untuk pratinjau daftar hadir.")
                        return
                    try: self.buat_daftar_hadir(ctx, pelaksana, self.tujuan_terpilih, self.mode, tmp_docx)
                    except Exception as e:
                        self._set_preview_status(f"Gagal membuat daftar hadir: {e}")
                        return
                else:
                    doc = DocxTemplate(template_file)
                    doc.render(ctx)
                    doc.save(tmp_docx)

            # ===== KONVERSI KE PDF =====
            if not HAS_FITZ:
                self._set_preview_status("Pratinjau visual butuh paket 'PyMuPDF'.\nJalankan: pip install pymupdf pillow")
                return

            if os.path.exists(tmp_pdf):
                try: os.remove(tmp_pdf)
                except Exception: pass

            # Prioritaskan Word (COM) dulu kalau tersedia: setelah panggilan
            # pertama, instance Word dibiarkan tetap terbuka (lihat
            # _get_persistent_word_app), jadi refresh berikutnya jauh lebih
            # cepat dibanding LibreOffice yang start-up ulang dari nol setiap
            # kali dipanggil sebagai proses baru.
            ok = False
            if HAS_DOCX2PDF and sys.platform == "win32":
                ok = self._convert_with_docx2pdf_safe(tmp_docx, tmp_pdf)
            if not ok:
                soffice_bin = self._find_soffice_binary()
                if soffice_bin:
                    ok = self._convert_with_soffice(soffice_bin, tmp_docx, tmp_pdf)

            if not ok:
                # Bantu diagnosis: tampilkan info konverter yang dicoba
                detail = []
                if soffice_bin: detail.append(f"LibreOffice: {soffice_bin} (gagal)")
                else: detail.append("LibreOffice: tidak ditemukan di PATH maupun lokasi instalasi umum")
                if HAS_DOCX2PDF: detail.append("Microsoft Word (COM): gagal/timeout")
                else: detail.append("Microsoft Word (docx2pdf): tidak terinstall")
                self._set_preview_status(
                    "Gagal konversi ke PDF untuk pratinjau.\n" + " | ".join(detail),
                    color="#EF4444"
                )
                return

            pdf_doc = fitz.open(tmp_pdf)
            page_images = []
            for page_idx in range(pdf_doc.page_count):
                page = pdf_doc.load_page(page_idx)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.3, 1.3))
                img_mode = "RGBA" if pix.alpha else "RGB"
                img = Image.frombytes(img_mode, [pix.width, pix.height], pix.samples)
                page_images.append(img)
            pdf_doc.close()

            combined_img = self._combine_preview_pages(page_images)

            self.after(0, lambda: self._apply_preview_image(combined_img))
        except Exception as e:
            self._set_preview_status(f"Gagal membuat preview: {str(e)}", color="#EF4444")
        finally:
            with self._preview_lock:
                self._preview_busy = False
                need_again = self._preview_pending
                self._preview_pending = False
            if need_again:
                self.after(50, self._launch_preview_render)

    def _combine_preview_pages(self, page_images, gap=20, bg_color=(229, 231, 235)):
        """Gabungkan semua halaman PDF (list gambar PIL) jadi SATU gambar
        panjang, ditumpuk vertikal dengan jarak antar-halaman, supaya live
        preview menampilkan SEMUA halaman (scroll ke bawah), bukan cuma
        halaman pertama."""
        if not page_images:
            return None
        if len(page_images) == 1:
            img = page_images[0]
            return img.convert("RGB") if img.mode != "RGB" else img
        max_width = max(im.width for im in page_images)
        total_height = sum(im.height for im in page_images) + gap * (len(page_images) - 1)
        combined = Image.new("RGB", (max_width, total_height), bg_color)
        y = 0
        for im in page_images:
            if im.mode != "RGB":
                im = im.convert("RGB")
            x = (max_width - im.width) // 2
            combined.paste(im, (x, y))
            y += im.height + gap
        return combined

    def _apply_preview_image(self, pil_img):
        if pil_img is None:
            return
        try:
            self.preview_canvas_frame.update_idletasks()
            frame_width = self.preview_canvas_frame.winfo_width()
            if frame_width < 100: frame_width = 480
        except Exception:
            frame_width = 480
        avail = max(frame_width - 30, 200)
        scale = avail / pil_img.width
        scale = max(0.12, min(scale, 1.2))
        disp_w = max(int(pil_img.width * scale), 50)
        disp_h = max(int(pil_img.height * scale), 50)
        ctk_img = ctk.CTkImage(light_image=pil_img, size=(disp_w, disp_h))
        self._preview_ctk_image = ctk_img
        self.preview_image_label.configure(image=ctk_img, text="")
        self.preview_status_lbl.configure(text="Pratinjau terkini.", text_color="gray")

        # PENTING: setelah pindah mode (Perjalanan Dinas <-> Undangan), panel
        # preview memindahkan preview_frame ke KOLOM GRID yang berbeda
        # (lihat _switch_to_undangan_layout / _switch_to_perjalanan_layout).
        # CTkScrollableFrame membungkus sebuah Canvas internal yang punya
        # scrollregion & posisi scroll SENDIRI -- kalau widget itu sempat
        # berukuran 0 (saat parent-nya baru dipindah/di-grid ulang) lalu
        # gambar baru dipasang, scrollregion & posisi scroll lama bisa
        # nyangkut sehingga gambar sebenarnya SUDAH terpasang tapi berada di
        # luar area yang sedang terlihat -- kelihatannya "kosong/blank".
        # Perbaikannya: paksa refresh geometry lagi setelah gambar terpasang
        # (bukan cuma sebelum menghitung lebar), lalu scroll paksa ke atas.
        def _force_refresh_scroll():
            try:
                self.preview_canvas_frame.update_idletasks()
                parent_canvas = getattr(self.preview_canvas_frame, "_parent_canvas", None)
                if parent_canvas is not None:
                    parent_canvas.configure(scrollregion=parent_canvas.bbox("all"))
                    parent_canvas.yview_moveto(0.0)
            except Exception:
                pass
        self.after(30, _force_refresh_scroll)

    def _find_soffice_binary(self):
        """Cari executable LibreOffice (soffice). Hasil di-cache di
        self._soffice_bin_cache supaya tidak scan ulang PATH + cek beberapa
        lokasi disk di SETIAP refresh Live Preview (itu menambah beberapa
        puluh-ratus ms per render kalau LibreOffice yang dipakai, bukan Word).

        shutil.which() SAJA sering tidak cukup di Windows: installer resmi
        LibreOffice TIDAK menambahkan dirinya ke PATH sistem secara
        otomatis (beda dengan kebiasaan installer di Linux/dev machine),
        jadi walau LibreOffice sudah terpasang, shutil.which("soffice")
        bisa saja tetap mengembalikan None. Untuk itu, kalau pencarian PATH
        gagal, coba juga lokasi instalasi default di Windows sebelum benar-
        benar menyerah.
        """
        cached = getattr(self, "_soffice_bin_cache", "__unset__")
        if cached != "__unset__":
            return cached
        found = shutil.which("soffice") or shutil.which("libreoffice")
        if not found:
            kandidat = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            for path in kandidat:
                if os.path.exists(path):
                    found = path
                    break
        self._soffice_bin_cache = found
        return found

    def _convert_with_soffice(self, soffice_bin, docx_path, pdf_path):
        out_dir = os.path.dirname(os.path.abspath(pdf_path)) or "."
        try:
            # CREATE_NO_WINDOW mencegah jendela CMD muncul sejenak di Windows,
            # dan mencegah subprocess hang saat parent exe tidak punya console
            # (PyInstaller console=False). stdin=DEVNULL wajib agar subprocess
            # tidak menunggu input dari stdin yang tidak ada.
            creation_flags = 0
            if sys.platform == "win32":
                creation_flags = subprocess.CREATE_NO_WINDOW
            subprocess.run(
                [soffice_bin, "--headless", "--norestore", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                check=True,
                timeout=60,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                creationflags=creation_flags,
            )
            generated = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if os.path.exists(generated):
                if os.path.abspath(generated) != os.path.abspath(pdf_path):
                    if os.path.exists(pdf_path): os.remove(pdf_path)
                    os.replace(generated, pdf_path)
                return True
            return False
        except Exception:
            return False

    def _get_persistent_word_app(self):
        """Ambil instance Word COM MILIK APLIKASI SENDIRI yang sudah terbuka,
        atau buka SEKALI kalau belum ada. Instance ini SENGAJA tidak di-Quit()
        setelah tiap konversi -- membuka Word butuh 1-3 detik, jadi kalau
        dibuka-tutup di setiap refresh Live Preview (yang bisa terpicu
        berkali-kali per menit saat mengetik/checklist), itu jadi penyebab
        utama preview terasa lambat. Word baru benar-benar ditutup saat
        aplikasi ditutup (lihat on_close).

        PENTING: pakai DispatchEx, BUKAN Dispatch biasa. win32com.client.
        Dispatch() akan MENYAMBUNG ke instance Word yang SUDAH TERBUKA di
        komputer user kalau ada (perilaku default COM Windows) -- artinya
        kalau user sedang kerja di Word lain, aplikasi ini bisa mengambil
        alih Word tsb (menyembunyikannya lewat Visible=False, lalu ikut
        Quit() dia berikut semua dokumen user yang belum disimpan saat
        aplikasi ditutup). DispatchEx memaksa Windows membuat proses Word
        BARU yang terpisah & independen, supaya Word milik user tidak pernah
        disentuh sama sekali.

        WAJIB dipanggil HANYA dari _word_worker_thread (satu-satunya thread
        yang boleh menyentuh objek COM ini -- lihat catatan di __init__)."""
        if self._word_app is not None:
            try:
                _ = self._word_app.Visible  # ping: masih hidup?
                return self._word_app
            except Exception:
                self._word_app = None  # instance lama sudah mati/crash, buat baru
                self._word_pid = None

        import pythoncom
        pythoncom.CoInitialize()
        import win32com.client
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        self._word_app = word
        try:
            import win32process
            self._word_pid = win32process.GetWindowThreadProcessId(word.Hwnd)[1]
        except Exception:
            self._word_pid = None
        return word

    def _kill_word_process(self, pid):
        """Hard-kill proses winword.exe tertentu lewat taskkill. Dipakai
        sebagai jaring pengaman terakhir kalau panggilan COM ke Word
        menggantung tanpa batas waktu (mis. Word diam-diam menampilkan
        dialog perbaikan dokumen yang tidak pernah bisa diklik siapapun
        karena instance-nya headless/Visible=False) -- daripada membiarkan
        proses itu jadi 'hantu' selamanya di Task Manager, kita paksa
        tutup supaya tidak menumpuk."""
        if not pid:
            return
        try:
            subprocess.run(
                ["taskkill", "/F", "/PID", str(pid), "/T"],
                stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                creationflags=subprocess.CREATE_NO_WINDOW,
                timeout=5,
            )
        except Exception:
            pass

    def _quit_persistent_word_app(self):
        if self._word_app is not None:
            try: self._word_app.Quit()
            except Exception: pass
            # Quit() sendiri bisa saja gagal/hang kalau instance dalam
            # keadaan macet -- pastikan proses benar-benar hilang.
            self._kill_word_process(self._word_pid)
            self._word_app = None
            self._word_pid = None

    def _get_word_worker_thread(self):
        """Thread TUNGGAL yang hidup selama aplikasi berjalan, satu-satunya
        yang boleh menyentuh self._word_app (lihat penjelasan panjang di
        __init__ soal kenapa 'satu thread baru per konversi' adalah akar
        masalah proses Word menumpuk). Tugas dikirim lewat queue, hasil
        dikembalikan lewat queue terpisah per tugas."""
        if self._word_worker_thread is not None and self._word_worker_thread.is_alive():
            return self._word_worker_thread

        self._word_task_queue = queue.Queue()

        def _worker_loop():
            import pythoncom
            pythoncom.CoInitialize()
            while True:
                task = self._word_task_queue.get()
                if task is None:  # sinyal berhenti
                    break
                docx_path, pdf_path, result_queue = task
                try:
                    with self._word_lock:
                        word = self._get_persistent_word_app()
                        abs_docx = os.path.abspath(docx_path).replace("/", "\\")
                        abs_pdf = os.path.abspath(pdf_path).replace("/", "\\")
                        doc = word.Documents.Open(abs_docx, ReadOnly=True, AddToRecentFiles=False, ConfirmConversions=False)
                        try:
                            doc.SaveAs(abs_pdf, FileFormat=17)  # 17 = wdFormatPDF
                        finally:
                            doc.Close(SaveChanges=False)
                    result_queue.put(("ok", None))
                except Exception as e:
                    # Instance mungkin sudah rusak -- buang supaya percobaan
                    # berikutnya membuka instance baru yang sehat.
                    self._kill_word_process(self._word_pid)
                    self._word_app = None
                    self._word_pid = None
                    result_queue.put(("error", str(e)))

        t = threading.Thread(target=_worker_loop, daemon=True, name="SIPS-WordCOM-Worker")
        t.start()
        self._word_worker_thread = t
        return t

    def _convert_with_docx2pdf_safe(self, docx_path, pdf_path):
        """Konversi docx -> pdf via Word COM, pakai instance Word yang SUDAH
        terbuka (persisten) supaya tidak menanggung biaya start-up Word di
        setiap refresh preview. Tugas dikirim ke satu-satunya worker thread
        (lihat _get_word_worker_thread) dengan timeout supaya tidak
        memblokir preview thread kalau Word hang -- dan kalau memang timeout,
        proses Word yang macet akan DI-KILL PAKSA (bukan cuma ditinggal)
        supaya tidak menumpuk di Task Manager."""
        self._get_word_worker_thread()
        result_queue = queue.Queue()
        self._word_task_queue.put((docx_path, pdf_path, result_queue))
        try:
            status, err = result_queue.get(timeout=20)
            return status == "ok"
        except queue.Empty:
            # Worker masih macet di panggilan COM -- paksa tutup proses Word
            # yang sedang dipegangnya supaya tidak jadi proses hantu, lalu
            # buang juga worker thread-nya (sudah pasti macet permanen di
            # panggilan COM yang blocking, tidak bisa "dibatalkan" dari
            # luar selain dengan mematikan prosesnya).
            self._kill_word_process(self._word_pid)
            self._word_app = None
            self._word_pid = None
            self._word_worker_thread = None
            self._word_task_queue = None
            return False

