"""
Generator dokumen surat: Surat Tugas (DPRD & ASN, format biasa/tabel
otomatis tergantung jumlah personel), ringkasan jumlah pelaksana DPRD per
kategori, dan Surat Pemberitahuan multi-halaman (satu halaman per tujuan).
"""
import os
import re
import tempfile

from docx import Document
from docxtpl import DocxTemplate

from sekretariat_app.sips.constants import AKD_LAINNYA_DISPLAY_NAMES, KATEGORI_DPRD_ORDER
from sekretariat_app.sips.settings import (
    TEMPLATE_ST_DPRD_BIASA,
    TEMPLATE_ST_DPRD_TABEL,
    TEMPLATE_ST_ASN_BIASA,
    TEMPLATE_ST_ASN_TABEL,
)
from sekretariat_app.sips.docx_utils import _combine_word_pages, _fill_table_rows_from_master, cleanup_surat_tugas_biasa
from sekretariat_app.sips.text_utils import increment_nomor, generate_periods


def buat_surat_tugas_dprd(ctx, selected_dprd, out_path):
    n = len(selected_dprd)
    render_ctx = ctx.copy()

    if n <= 3:
        template_path = TEMPLATE_ST_DPRD_BIASA
        for i in range(1, 4):
            p = selected_dprd[i - 1] if i <= n else {}
            render_ctx[f"pelaksana_tugas_{i}"] = p.get('nama', '')
            render_ctx[f"jabatan_pelaksana_{i}"] = p.get('jabatan', '')
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        cleanup_surat_tugas_biasa(doc, 'dprd')
        doc.save(out_path)
    else:
        template_path = TEMPLATE_ST_DPRD_TABEL
        render_ctx["loop"] = {"index": ""}
        render_ctx["tabel"] = {"nama": "", "jabatan": ""}
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        rows_data = [[str(i + 1), p.get('nama', ''), p.get('jabatan', '')] for i, p in enumerate(selected_dprd)]
        _fill_table_rows_from_master(doc, ["No", "Nama", "Jabatan"], rows_data)
        doc.save(out_path)

def buat_surat_tugas_asn(ctx, selected_asn, out_path):
    n = len(selected_asn)
    render_ctx = ctx.copy()

    if n <= 3:
        template_path = TEMPLATE_ST_ASN_BIASA
        for i in range(1, 4):
            p = selected_asn[i - 1] if i <= n else {}
            render_ctx[f"nama_asn_{i}"] = p.get('nama', '')
            render_ctx[f"pangkat_asn_{i}"] = p.get('pangkat', '')
            render_ctx[f"nip_asn_{i}"] = p.get('nip', '')
            render_ctx[f"jabatan_asn_{i}"] = p.get('jabatan', '')
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        cleanup_surat_tugas_biasa(doc, 'asn')
        doc.save(out_path)
    else:
        template_path = TEMPLATE_ST_ASN_TABEL
        render_ctx["loop"] = {"index": ""}
        render_ctx["tabel"] = {"nama_asn": "", "jabatan_asn": ""}
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        rows_data = []
        for i, p in enumerate(selected_asn):
            nama_col = f"{p.get('nama', '')}\nNIP. {p.get('nip', '-')}"
            jabatan_col = f"{p.get('jabatan', '-')}\n{p.get('pangkat', '-')}"
            rows_data.append([str(i + 1), nama_col, jabatan_col])
        _fill_table_rows_from_master(doc, ["No", "Nama", "Jabatan"], rows_data)
        doc.save(out_path)

def _label_kategori_dprd(cat, jabatan_list):
    if cat == "Pimpinan DPRD": return "Pimpinan DPRD"
    # AKD Lainnya (Banggar/Banmus/Bapemperda) disimpan di database/Excel
    # sebagai singkatan supaya pencocokan kategori tidak berubah, tapi
    # dicetak di surat dengan nama lengkapnya (Badan Anggaran/Badan
    # Musyawarah/Badan Pembentukan Perda). "Badan Kehormatan" sudah nama
    # lengkap jadi tidak berubah.
    display_cat = AKD_LAINNYA_DISPLAY_NAMES.get(cat, cat)
    has_pimpinan = any(("ketua" in j.lower() or "sekretaris" in j.lower()) for j in jabatan_list)
    has_anggota = any(("anggota" in j.lower() and "ketua" not in j.lower()) for j in jabatan_list)
    if has_pimpinan and has_anggota: label = f"Pimpinan dan Anggota {display_cat}"
    elif has_pimpinan: label = f"Pimpinan {display_cat}"
    elif has_anggota: label = f"Anggota {display_cat}"
    else: label = display_cat
    if "DPRD" not in label: label += " DPRD"
    return label

def compute_pelaksana_dprd_summary(selected_dprd):
    by_cat = {}
    for p in selected_dprd:
        cat = str(p.get('kategori', '')).strip()
        by_cat.setdefault(cat, []).append(p.get('jabatan', ''))
    summary = []
    for cat in KATEGORI_DPRD_ORDER:
        jabatan_list = by_cat.get(cat)
        if not jabatan_list: continue
        label = _label_kategori_dprd(cat, jabatan_list)
        summary.append((label, len(jabatan_list)))
    for cat, jabatan_list in by_cat.items():
        if cat not in KATEGORI_DPRD_ORDER:
            label = _label_kategori_dprd(cat, jabatan_list)
            summary.append((label, len(jabatan_list)))
    return summary

def _remove_empty_pelaksana_lines(doc):
    for p in list(doc.paragraphs):
        txt = p.text
        normalized = re.sub(r'\s+', ' ', txt).strip()
        if "Pendamping ASN" in txt:
            m = re.search(r':\s*(\d+)\s+Orang', normalized)
            if m and int(m.group(1)) == 0:
                p._element.getparent().remove(p._element)
                continue
        if "Kota Bitung" in txt and "Orang" in txt:
            m = re.search(r':\s*(\d+)\s+Orang', normalized)
            if not m:
                p._element.getparent().remove(p._element)
                continue

def apply_pelaksana_dprd_summary_to_ctx(ctx, selected_dprd, max_slots=4):
    summary = compute_pelaksana_dprd_summary(selected_dprd)
    for i in range(1, max_slots + 1):
        if i <= len(summary):
            label, jumlah = summary[i - 1]
            ctx[f"pelaksana_tugas_{i}"] = label
            ctx[f"jlh_pelaksana_dprd{i}"] = jumlah
        else:
            ctx[f"pelaksana_tugas_{i}"] = ""
            ctx[f"jlh_pelaksana_dprd{i}"] = ""
    return ctx

def buat_surat_pemberitahuan_multi(template_path, ctx, selected_dprd, selected_asn, destinations, base_number, out_path, label_asn="Pendamping ASN"):
    """Satu halaman per tujuan/periode, digabung jadi satu dokumen akhir.

    BUG YANG DIPERBAIKI ("teks naik turun antar halaman"): versi lama
    menggabungkan tiap halaman dengan mencampur `master_doc.add_page_break()`
    (API tingkat tinggi python-docx yang menyisipkan paragraf barunya TEPAT
    SEBELUM elemen <w:sectPr>) dengan `body.append(element)` manual untuk
    elemen dari halaman berikutnya (yang selalu masuk ke posisi PALING
    AKHIR, yaitu SETELAH <w:sectPr>). Mencampur keduanya membuat <w:sectPr>
    -- yang menurut skema OOXML wajib jadi elemen TERAKHIR di <w:body> --
    malah nyangkut di TENGAH dokumen. Word "memperbaiki" struktur tidak
    valid ini sendiri saat dibuka, dan hasil perbaikan otomatis itu yang
    membuat urutan/posisi konten tiap halaman kadang bergeser naik/turun
    tidak terduga.

    Sekarang tiap periode dirender sbg file .docx sendiri-sendiri (sama
    seperti pola SPD/SPPD dan Undangan Rapat Biasa), lalu digabung lewat
    `_combine_word_pages` yang sudah menangani `sectPr` dengan benar
    (dilepas dulu, semua elemen digabung berurutan lewat `append()` biasa,
    baru dipasang lagi di posisi terakhir) sehingga setiap halaman selalu
    tetap persis pada halamannya sendiri.
    """
    periods = generate_periods(ctx.get("tanggal_mulai", ""), destinations)
    base_ctx = ctx.copy()
    apply_pelaksana_dprd_summary_to_ctx(base_ctx, selected_dprd)
    base_ctx["pelaksana_tugas_asn_info"] = label_asn
    base_ctx["jlh_pelaksana_asn"] = len(selected_asn)

    tmpdir = tempfile.mkdtemp()
    page_files = []
    try:
        for idx, period in enumerate(periods):
            nomor_surat = increment_nomor(base_number, idx)
            page_ctx = base_ctx.copy()
            page_ctx["nomor_surat_info"] = nomor_surat
            page_ctx["tujuan_surat_info"] = period["tujuan"]
            page_ctx["hari_info"] = period["hari"]
            page_ctx["tanggal_bertugas_info"] = period["tanggal"]

            tmp_docx = os.path.join(tmpdir, f"pemberitahuan_{idx}.docx")
            doc_tpl = DocxTemplate(template_path)
            doc_tpl.render(page_ctx)
            doc_tpl.save(tmp_docx)
            doc = Document(tmp_docx)
            _remove_empty_pelaksana_lines(doc)
            doc.save(tmp_docx)
            page_files.append(tmp_docx)

        if page_files:
            _combine_word_pages(page_files, out_path)
    finally:
        for f in page_files:
            try:
                os.unlink(f)
            except OSError:
                pass
        try:
            os.rmdir(tmpdir)
        except OSError:
            pass
