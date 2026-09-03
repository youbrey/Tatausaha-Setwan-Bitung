from __future__ import annotations

import shutil
import tempfile
from collections import OrderedDict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Iterable

from docx import Document
from docxtpl import DocxTemplate, RichText

from sekretariat_app.sips.constants import (
    PELAKSANA_RAPAT_CUSTOM,
    TUJUAN_SURAT_DPRD_MAP,
)
from sekretariat_app.sips.document_generators import (
    buat_surat_pemberitahuan_multi,
    buat_surat_tugas_asn,
    buat_surat_tugas_dprd,
)
from sekretariat_app.sips.docx_utils import (
    _combine_word_pages,
    _fill_table_rows_from_master,
    _force_daftar_hadir_page_break,
    _force_paripurna_page_break,
    build_halaman_tujuan_lain_template,
    build_tujuan_richtext,
    cleanup_skenario_paripurna,
    ensure_numbered_slots,
)
from sekretariat_app.sips.master_data import PersonnelMaster
from sekretariat_app.sips.meeting_attendance import generate_daftar_hadir_rapat
from sekretariat_app.sips.models import (
    InvitationFormData,
    TravelFormData,
    day_name_id,
    format_date_id,
    inclusive_days,
    terbilang_small,
)
from sekretariat_app.sips.official_note import generate_naskah_dinas
from sekretariat_app.sips.settings import (
    TEMPLATE_DAFTAR_HADIR,
    TEMPLATE_DAFTAR_HADIR_PARIPURNA,
    TEMPLATE_DAFTAR_HADIR_PIHAK_TERKAIT,
    TEMPLATE_DAFTAR_HADIR_RAPAT_MAP,
    TEMPLATE_DAFTAR_HADIR_SEKRETARIAT,
    TEMPLATE_DAFTAR_HADIR_TAF,
    TEMPLATE_NASKAH_DINAS_RAPAT,
    TEMPLATE_PARIPURNA,
    TEMPLATE_PEMBERITAHUAN,
    TEMPLATE_RAPAT_BIASA,
    TEMPLATE_SPD_BELAKANG,
    TEMPLATE_SPD_DEPAN,
)
from sekretariat_app.sips.sppd_generators import buat_sppd_asn, buat_sppd_dprd
from sekretariat_app.sips.text_utils import (
    detect_zona_waktu,
    extract_city_name,
    format_jabatan_penandatanganan,
    format_signature_position,
    generate_periods,
    increment_nomor_paripurna,
    is_in_sulawesi_utara,
    is_plain_region_name,
    slugify_filename,
    strip_jenis_perjalanan_prefix,
)


def _split_signer(value: str, default_position: str) -> tuple[str, str]:
    position, name = value.split(" - ", 1) if " - " in value else (default_position, value)
    return format_jabatan_penandatanganan(position.strip()), name.strip()


def _is_taf(value: str) -> bool:
    return " ".join(value.lower().split()).startswith("tenaga ahli fraksi")


@dataclass(frozen=True, slots=True)
class GenerationFailure:
    document: str
    message: str


@dataclass(slots=True)
class GenerationReport:
    """Hasil pembuatan batch; file yang berhasil tidak dibuang saat satu cabang gagal."""

    files: list[Path] = field(default_factory=list)
    failures: list[GenerationFailure] = field(default_factory=list)

    @property
    def error_message(self) -> str:
        return "\n".join(f"{item.document}: {item.message}" for item in self.failures)


class DocumentGenerationError(RuntimeError):
    def __init__(self, report: GenerationReport) -> None:
        self.report = report
        super().__init__(report.error_message or "Pembuatan dokumen gagal.")


class SIPSService:
    """Orkestrasi generator SIPS tanpa ketergantungan pada widget UI."""

    def __init__(self, master: PersonnelMaster | None = None) -> None:
        self.master = master or PersonnelMaster()

    def build_travel_context(self, data: TravelFormData, *, preview: bool = False) -> dict[str, Any]:
        data.validate_preview() if preview else data.validate()
        duration = inclusive_days(data.start_date, data.end_date)
        signer_dprd_position, signer_dprd_name = _split_signer(data.signer_dprd, "KETUA")
        signer_asn_position, signer_asn_name = _split_signer(data.signer_asn, "SEKRETARIS DPRD")
        signer_dprd_formatted = format_signature_position(signer_dprd_position)
        signer_asn_formatted = format_signature_position(signer_asn_position)
        destination_text = " / ".join(data.destinations)
        city_names = [extract_city_name(destination) for destination in data.destinations]
        transport = "Pesawat / Mobil / Kereta" if any(
            not is_in_sulawesi_utara(city) for city in city_names
        ) else "Mobil"
        numbers = data.document_numbers
        return {
            "nomor_surat": numbers.get("surat_tugas_dprd", ""),
            "nomor_surat_asn": numbers.get("surat_tugas_asn", ""),
            "nomor_pemberitahuan_dprd": numbers.get("pemberitahuan_dprd", ""),
            "nomor_pemberitahuan_asn": numbers.get("pemberitahuan_asn", ""),
            "nomor_spd_dprd": numbers.get("spd_dprd", ""),
            "nomor_spd_asn": numbers.get("spd_asn", ""),
            "nomor_spd_pelaksana": numbers.get("spd_pelaksana", numbers.get("spd_asn", "")),
            "nomor_spd_pendamping": numbers.get("spd_pendamping", numbers.get("spd_asn", "")),
            "nomor_spd": numbers.get("spd_dprd", numbers.get("spd_asn", "")),
            "tanggal_surat": format_date_id(data.letter_date),
            "tanggal_surat_asn": format_date_id(data.letter_date),
            "jenis_perjalanan": data.travel_type,
            "tujuan_bertugas": destination_text,
            "tujuan_bertugas_list": list(data.destinations),
            "dasar_surat_dprd": data.basis_dprd.strip(),
            "dasar_surat_asn": data.basis_asn.strip(),
            "materi_tugas": data.subject.strip(),
            "materi_tugas_asn": data.subject.strip(),
            "isi_surat_pemberitahuan": data.notice_subject.strip(),
            "isi_surat_izin": data.subject.strip(),
            "tanggal_mulai": format_date_id(data.start_date),
            "tanggal_akhir": format_date_id(data.end_date),
            "jumlah_angka": duration,
            "jumlah_teks": terbilang_small(duration),
            "jabatan_ttd": signer_dprd_formatted,
            "nama_ttd": signer_dprd_name,
            "jabatan_ttd_asn": signer_asn_formatted,
            "nama_ttd_asn": signer_asn_name,
            "transportasi_otomatis": transport,
            "tanggal_surat_info": format_date_id(data.letter_date),
            "tujuan_surat_info": destination_text,
            "pelaksana_tugas_dprd_info": "Pimpinan dan Anggota",
            "jenis_perjalanan_info": data.travel_type,
            "tujuan_bertugas_info": destination_text,
            "materi_tugas_info": data.notice_subject.strip(),
            "hari_info": "Sesuai Jadwal",
            "tanggal_bertugas_info": f"{format_date_id(data.start_date)} s/d {format_date_id(data.end_date)}",
            "pelaksana_tugas_info": "Anggota DPRD",
            "jlh_pelaksana_dprd": len(data.dprd),
            "pelaksana_tugas_asn_info": "Pendamping ASN" if data.mode == "dprd" else "Pelaksana ASN",
            "jlh_pelaksana_asn": len(data.asn if data.mode == "dprd" else data.executors),
            "jlh_pelaksana": len(data.executors),
            "jlh_pendamping": len(data.companions),
            "jabatan_ttd_info": signer_dprd_formatted,
            "nama_ttd_info": signer_dprd_name,
            "pelaksana_dprd": data.dprd,
            "pelaksana_asn": data.asn,
            "pelaksana_list": data.executors,
            "pendamping_list": data.companions,
        }

    @staticmethod
    def _run_generation(
        report: GenerationReport,
        label: str,
        expected_files: Iterable[Path],
        callback: Callable[[], None],
    ) -> None:
        expected = list(expected_files)
        before = {
            path: (path.stat().st_mtime_ns, path.stat().st_size)
            for path in expected
            if path.exists()
        }
        failed = False
        try:
            callback()
        except Exception as exc:
            failed = True
            report.failures.append(GenerationFailure(label, str(exc)))
        produced: list[Path] = []
        for path in expected:
            if not path.exists():
                continue
            after = (path.stat().st_mtime_ns, path.stat().st_size)
            if not failed or path not in before or before[path] != after:
                produced.append(path)
                if path not in report.files:
                    report.files.append(path)
        if not produced and not failed and not any(
            failure.document == label for failure in report.failures
        ):
            report.failures.append(GenerationFailure(label, "Generator tidak menghasilkan file."))

    def generate_travel_report(
        self,
        data: TravelFormData,
        output_directory: str | Path,
        *,
        preview: bool = False,
        preview_document: str | None = None,
    ) -> GenerationReport:
        context = self.build_travel_context(data, preview=preview)
        target = Path(output_directory)
        target.mkdir(parents=True, exist_ok=True)
        report = GenerationReport()

        def output(label: str) -> Path:
            return self._travel_output_path(data, target, label)

        def selected(*keys: str) -> bool:
            return preview_document is None or preview_document in keys

        if data.mode == "dprd":
            if data.dprd and selected("task-dprd"):
                path = output("surat-tugas")
                self._run_generation(
                    report, "Surat Tugas DPRD", [path],
                    lambda: buat_surat_tugas_dprd(context, data.dprd, path),
                )
            if data.asn and selected("task-asn"):
                path = output("surat-tugas-pendamping")
                self._run_generation(
                    report, "Surat Tugas Pendamping ASN", [path],
                    lambda: buat_surat_tugas_asn(context, data.asn, path),
                )
            if (data.dprd or data.asn) and selected("notice-dprd"):
                path = output("surat-pemberitahuan")
                self._run_generation(
                    report, "Surat Pemberitahuan DPRD", [path],
                    lambda: buat_surat_pemberitahuan_multi(
                        TEMPLATE_PEMBERITAHUAN, context, data.dprd, data.asn,
                        data.destinations, context["nomor_pemberitahuan_dprd"], path,
                        label_asn="Pendamping ASN",
                    ),
                )
            if data.dprd and selected("spd-dprd-front", "spd-dprd-back"):
                front, back = output("spd-depan"), output("spd-belakang")
                self._run_generation(
                    report, "SPD DPRD", [front, back],
                    lambda: buat_sppd_dprd(
                        TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, context,
                        data.dprd, data.destinations, front, back,
                    ),
                )
            if data.asn and selected("spd-asn-front", "spd-asn-back"):
                front, back = output("spd-pendamping-depan"), output("spd-pendamping-belakang")
                self._run_generation(
                    report, "SPD Pendamping ASN", [front, back],
                    lambda: buat_sppd_asn(
                        TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, context,
                        data.asn, data.destinations, front, back,
                    ),
                )
            # Perilaku SIPS lama: daftar hadir mode DPRD hanya berisi anggota
            # DPRD. Pendamping ASN sudah memiliki Surat Tugas/SPD sendiri.
            attendance_people = data.dprd
            attendance_mode = "dprd"
        else:
            for people, role, role_key, task_label, notice_label, spd_key in (
                (data.executors, "Pelaksana ASN", "executor", "surat-tugas-pelaksana", "surat-pemberitahuan-pelaksana", "nomor_spd_pelaksana"),
                (data.companions, "Pendamping ASN", "companion", "surat-tugas-pendamping", "surat-pemberitahuan-pendamping", "nomor_spd_pendamping"),
            ):
                if not people:
                    continue
                role_context = dict(context)
                role_context["pelaksana_tugas_asn_info"] = role
                role_context["jlh_pelaksana_asn"] = len(people)
                if selected(f"task-{role_key}"):
                    path = output(task_label)
                    self._run_generation(
                        report, f"Surat Tugas {role}", [path],
                        lambda p=path, c=role_context, people_for_role=people: buat_surat_tugas_asn(
                            c, people_for_role, p
                        ),
                    )
                if selected(f"notice-{role_key}"):
                    notice_path = output(notice_label)
                    self._run_generation(
                        report, f"Surat Pemberitahuan {role}", [notice_path],
                        lambda p=notice_path, c=role_context, people_for_role=people, role_name=role: buat_surat_pemberitahuan_multi(
                            TEMPLATE_PEMBERITAHUAN, c, [], people_for_role,
                            data.destinations, context["nomor_pemberitahuan_asn"], p,
                            label_asn=role_name,
                        ),
                    )
                spd_context = dict(role_context)
                spd_context["nomor_spd_asn"] = context[spd_key]
                if selected(f"spd-{role_key}-front", f"spd-{role_key}-back"):
                    front = output(f"spd-{slugify_filename(role)}-depan")
                    back = output(f"spd-{slugify_filename(role)}-belakang")
                    self._run_generation(
                        report, f"SPD {role}", [front, back],
                        lambda c=spd_context, people_for_role=people, f=front, b=back: buat_sppd_asn(
                            TEMPLATE_SPD_DEPAN, TEMPLATE_SPD_BELAKANG, c,
                            people_for_role, data.destinations, f, b,
                        ),
                    )
            attendance_people = data.executors + data.companions
            attendance_mode = "setwan"

        if attendance_people and selected("attendance"):
            path = output("daftar-hadir")
            self._run_generation(
                report, "Daftar Hadir Perjalanan Dinas", [path],
                lambda: self.generate_travel_attendance(
                    context, attendance_people, data.destinations, attendance_mode, path,
                ),
            )
        return report

    def generate_travel(
        self,
        data: TravelFormData,
        output_directory: str | Path,
        *,
        preview: bool = False,
        preview_document: str | None = None,
    ) -> list[Path]:
        report = self.generate_travel_report(
            data,
            output_directory,
            preview=preview,
            preview_document=preview_document,
        )
        if report.failures:
            raise DocumentGenerationError(report)
        return report.files

    @staticmethod
    def _travel_category_slug(data: TravelFormData) -> str:
        if data.mode == "setwan":
            labels = []
            if data.executors:
                labels.append("pelaksana-asn")
            if data.companions:
                labels.append("pendamping-asn")
            return "-".join(labels) or "asn"
        categories = list(OrderedDict.fromkeys(item.get("kategori", "") for item in data.dprd if item.get("kategori")))
        if data.asn:
            categories.append("Pendamping ASN")
        return "-".join(filter(None, (slugify_filename(value) for value in categories))) or "umum"

    def _travel_output_path(self, data: TravelFormData, target: Path, label: str) -> Path:
        mode = "dprd" if data.mode == "dprd" else "setwan"
        category = self._travel_category_slug(data)
        destination = slugify_filename(data.destinations[0]) or "tujuan"
        return target / f"{label}-{mode}-{category}-{destination}.docx"

    def travel_preview_documents(self, data: TravelFormData) -> list[tuple[str, str]]:
        """Daftar dokumen preview tanpa merender seluruh paket DOCX."""

        data.validate_preview()
        target = Path(".")
        options: list[tuple[str, str]] = []

        def add(key: str, label: str) -> None:
            options.append((key, self._travel_output_path(data, target, label).name))

        if data.mode == "dprd":
            if data.dprd:
                add("task-dprd", "surat-tugas")
            if data.asn:
                add("task-asn", "surat-tugas-pendamping")
            if data.dprd or data.asn:
                add("notice-dprd", "surat-pemberitahuan")
            if data.dprd:
                add("spd-dprd-front", "spd-depan")
                add("spd-dprd-back", "spd-belakang")
            if data.asn:
                add("spd-asn-front", "spd-pendamping-depan")
                add("spd-asn-back", "spd-pendamping-belakang")
            if data.dprd:
                add("attendance", "daftar-hadir")
        else:
            for people, role, role_key, task_label, notice_label in (
                (data.executors, "Pelaksana ASN", "executor", "surat-tugas-pelaksana", "surat-pemberitahuan-pelaksana"),
                (data.companions, "Pendamping ASN", "companion", "surat-tugas-pendamping", "surat-pemberitahuan-pendamping"),
            ):
                if not people:
                    continue
                add(f"task-{role_key}", task_label)
                add(f"notice-{role_key}", notice_label)
                add(f"spd-{role_key}-front", f"spd-{slugify_filename(role)}-depan")
                add(f"spd-{role_key}-back", f"spd-{slugify_filename(role)}-belakang")
            if data.executors or data.companions:
                add("attendance", "daftar-hadir")
        return options

    def generate_travel_attendance(
        self,
        context: dict[str, Any],
        people: list[dict[str, Any]],
        destinations: list[str],
        mode: str,
        output_path: str | Path,
    ) -> Path:
        periods = generate_periods(context.get("tanggal_mulai", ""), destinations)
        temporary_files: list[str] = []
        try:
            for period in periods:
                actor_text, institution = self._attendance_heading(people, period["tujuan"], mode)
                render_context = {
                    "pelaksana_tugas_dprd": actor_text,
                    "JENIS_PERJALANAN_DAFTAR_HADIR": context.get("jenis_perjalanan", "").upper(),
                    "TEMPAT_TUGAS_DPRD_DAFTAR_HADIR": institution,
                    "MATERI_TUGAS_DPRD_DAFTAR_HADIR": strip_jenis_perjalanan_prefix(
                        context.get("materi_tugas", ""), context.get("jenis_perjalanan", "")
                    ).upper(),
                    "HARI": period["hari"],
                    "TANGGAL_DAFTAR_HADIR": period["tanggal"],
                    "TEMPAT_DAFTAR_HADIR": period["tujuan"],
                    "zona": detect_zona_waktu(period["tujuan"]),
                    "loop": {"index": ""},
                    "tabel": {"NAMA_DAFTAR_HADIR": "", "jabatan_daftar_hadir": ""},
                }
                template = DocxTemplate(TEMPLATE_DAFTAR_HADIR)
                template.render(render_context)
                temporary = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                template.save(temporary)
                document = Document(temporary)
                rows = [[str(index + 1), item.get("nama", ""), item.get("jabatan", "")] for index, item in enumerate(people)]
                _fill_table_rows_from_master(document, ["no", "nama", "jabatan", "tanda tangan"], rows, max_tables=1)
                _force_daftar_hadir_page_break(document)
                document.save(temporary)
                temporary_files.append(temporary)
            if temporary_files:
                _combine_word_pages(temporary_files, output_path)
            return Path(output_path)
        finally:
            for path in temporary_files:
                Path(path).unlink(missing_ok=True)

    @staticmethod
    def _attendance_heading(people: list[dict[str, Any]], destination: str, mode: str) -> tuple[str, str]:
        if mode == "setwan":
            actor_text = "PEJABAT DAN PEGAWAI SEKRETARIAT DPRD KOTA BITUNG"
        else:
            groups: OrderedDict[str, list[dict[str, Any]]] = OrderedDict()
            for person in people:
                groups.setdefault(person.get("kategori", "Lainnya") or "Lainnya", []).append(person)
            labels: list[str] = []
            for category, members in groups.items():
                if category == "Pimpinan DPRD":
                    label = members[0].get("jabatan", "Pimpinan DPRD Kota Bitung") if len(members) == 1 else "Pimpinan DPRD Kota Bitung"
                else:
                    positions = [member.get("jabatan", "").lower() for member in members]
                    leaders = any("ketua" in position or "sekretaris" in position for position in positions)
                    members_present = any("anggota" in position for position in positions)
                    if leaders and members_present:
                        label = f"Pimpinan dan Anggota {category}"
                    elif leaders:
                        label = f"Pimpinan {category}"
                    elif members_present:
                        label = f"Anggota {category}"
                    else:
                        label = category
                    if "DPRD" not in label.upper():
                        label += " DPRD Kota Bitung"
                labels.append(label)
            actor_text = (
                " BERSAMA ".join(labels).upper()
                if len(labels) <= 3 else "PIMPINAN DAN ANGGOTA DPRD KOTA BITUNG"
            )
        clean_destination = destination.strip()
        institution = (
            clean_destination if "DPRD" in clean_destination.upper() else f"DPRD {clean_destination}"
        ) if is_plain_region_name(clean_destination) else clean_destination
        return actor_text, institution

    @staticmethod
    def _invitation_date_slug(data: InvitationFormData, *, letter_date: bool = False) -> str:
        selected_date = data.letter_date if letter_date else data.meeting_date
        return (
            f"{slugify_filename(day_name_id(selected_date))}-"
            f"{selected_date.day}-"
            f"{slugify_filename(format_date_id(selected_date).split()[1])}"
        )

    def generate_official_note(
        self,
        data: InvitationFormData,
        output_directory: str | Path,
    ) -> Path:
        """Buat Naskah Dinas secara mandiri, setara tombol pada SIPS lama."""
        if data.invitation_type not in {"paripurna", "biasa"}:
            raise ValueError("Jenis undangan tidak valid.")
        if not data.number.strip():
            raise ValueError("Nomor undangan wajib diisi untuk membuat Naskah Dinas.")
        if not data.agenda.strip():
            raise ValueError("Isi surat/agenda rapat wajib diisi untuk membuat Naskah Dinas.")
        target = Path(output_directory)
        target.mkdir(parents=True, exist_ok=True)
        path = target / (
            f"naskah-dinas-{'paripurna' if data.invitation_type == 'paripurna' else 'rapat'}-"
            f"{self._invitation_date_slug(data)}.docx"
        )
        generate_naskah_dinas(
            TEMPLATE_NASKAH_DINAS_RAPAT,
            path,
            data.number,
            format_date_id(data.letter_date),
            data.agenda,
        )
        return path

    def generate_meeting_attendance(
        self,
        data: InvitationFormData,
        output_directory: str | Path,
    ) -> Path:
        """Buat Daftar Hadir secara mandiri dengan lembar tambahan terkait."""
        data.validate_preview()
        target = Path(output_directory)
        target.mkdir(parents=True, exist_ok=True)
        label = "paripurna" if data.invitation_type == "paripurna" else slugify_filename(data.meeting_executor)
        path = target / f"daftar-hadir-{label}-{self._invitation_date_slug(data)}.docx"
        self._generate_meeting_attendance(data, path)
        return path

    def generate_invitation_report(
        self,
        data: InvitationFormData,
        output_directory: str | Path,
        *,
        preview: bool = False,
        preview_document: str | None = None,
    ) -> GenerationReport:
        data.validate_preview() if preview else data.validate()
        target = Path(output_directory)
        target.mkdir(parents=True, exist_ok=True)
        date_slug = self._invitation_date_slug(data)
        report = GenerationReport()

        def selected(key: str) -> bool:
            return preview_document is None or preview_document == key

        if data.invitation_type == "paripurna" and selected("invitation"):
            main_path = target / (
                f"undangan-paripurna-{self._invitation_date_slug(data, letter_date=True)}.docx"
            )
            self._run_generation(
                report, "Undangan Paripurna", [main_path],
                lambda: self._generate_plenary(data, main_path),
            )
        elif data.invitation_type == "biasa" and selected("invitation"):
            executor_slug = slugify_filename(data.meeting_executor) or "pelaksana"
            meeting_slug = slugify_filename(data.meeting_type) or "rapat"
            main_path = target / f"{meeting_slug}-{executor_slug}-{date_slug}.docx"
            self._run_generation(
                report, "Undangan Rapat Biasa", [main_path],
                lambda: self._generate_regular(data, main_path),
            )

        if data.include_official_note and selected("official-note"):
            path = target / f"naskah-dinas-{'paripurna' if data.invitation_type == 'paripurna' else 'rapat'}-{date_slug}.docx"
            self._run_generation(
                report, "Naskah Dinas", [path],
                lambda: generate_naskah_dinas(
                    TEMPLATE_NASKAH_DINAS_RAPAT, path, data.number,
                    format_date_id(data.letter_date), data.agenda,
                ),
            )
        if data.include_attendance and selected("attendance"):
            path = target / f"daftar-hadir-{'paripurna' if data.invitation_type == 'paripurna' else slugify_filename(data.meeting_executor)}-{date_slug}.docx"
            self._run_generation(
                report, "Daftar Hadir Rapat", [path],
                lambda: self._generate_meeting_attendance(data, path),
            )
        return report

    def generate_invitation(
        self,
        data: InvitationFormData,
        output_directory: str | Path,
        *,
        preview: bool = False,
        preview_document: str | None = None,
    ) -> list[Path]:
        report = self.generate_invitation_report(
            data,
            output_directory,
            preview=preview,
            preview_document=preview_document,
        )
        if report.failures:
            raise DocumentGenerationError(report)
        return report.files

    def invitation_preview_documents(self, data: InvitationFormData) -> list[tuple[str, str]]:
        """Daftar dokumen undangan yang tersedia untuk live preview lazy."""

        data.validate_preview()
        date_slug = self._invitation_date_slug(data)
        if data.invitation_type == "paripurna":
            main_name = (
                f"undangan-paripurna-{self._invitation_date_slug(data, letter_date=True)}.docx"
            )
        else:
            executor_slug = slugify_filename(data.meeting_executor) or "pelaksana"
            meeting_slug = slugify_filename(data.meeting_type) or "rapat"
            main_name = f"{meeting_slug}-{executor_slug}-{date_slug}.docx"
        options = [("invitation", main_name)]
        if data.include_official_note:
            kind = "paripurna" if data.invitation_type == "paripurna" else "rapat"
            options.append(("official-note", f"naskah-dinas-{kind}-{date_slug}.docx"))
        if data.include_attendance:
            kind = (
                "paripurna"
                if data.invitation_type == "paripurna"
                else slugify_filename(data.meeting_executor)
            )
            options.append(("attendance", f"daftar-hadir-{kind}-{date_slug}.docx"))
        return options

    @staticmethod
    def _generate_plenary(data: InvitationFormData, output_path: str | Path) -> None:
        position, name = _split_signer(data.signer, "KETUA")
        context = {
            "tgl_surat_paripurna": format_date_id(data.letter_date),
            "isi_undangan_paripurna": data.agenda.strip(),
            "hari_paripurna": day_name_id(data.meeting_date),
            "tanggal_paripurna": format_date_id(data.meeting_date),
            "jam_paripurna": data.time_text.strip(),
            "pakaian_paripurna": data.clothing,
            "jabatan_ttd_paripurna": position,
            "nama_ttd_paripurna": name,
        }
        for index in range(7):
            context[f"skenario_paripurna_{index + 1}"] = data.scenarios[index].strip() if index < len(data.scenarios) else ""
        for index in range(8):
            context[f"nomor_surat_paripurna_{index + 1}"] = increment_nomor_paripurna(data.number, index)
        template = DocxTemplate(TEMPLATE_PARIPURNA)
        template.render(context)
        template.save(output_path)
        document = Document(output_path)
        cleanup_skenario_paripurna(document)
        _force_paripurna_page_break(document)
        document.save(output_path)

    @staticmethod
    def _generate_regular(data: InvitationFormData, output_path: str | Path) -> None:
        parties = [value.strip() for value in data.related_parties if value.strip()]
        parties_mayor = [value for value in parties if not _is_taf(value)]
        position, name = _split_signer(data.signer, "KETUA")
        executor = "" if data.meeting_executor == PELAKSANA_RAPAT_CUSTOM else data.meeting_executor.strip()
        purpose_text = TUJUAN_SURAT_DPRD_MAP.get(executor, executor.upper())
        purpose = RichText()
        purpose.add(purpose_text, bold=True)
        pages = [[item.strip() for item in page if item.strip()] for page in data.other_destination_pages]
        pages = [page for page in pages if page]
        main_destinations = pages[0] if pages else []
        extra_destinations = pages[1:] if len(pages) > 1 else []
        context: dict[str, Any] = {
            "tgl_surat_rapat_biasa": format_date_id(data.letter_date),
            "pelaksana_rapat_skpd": executor,
            "jenis_rapat_biasa_skpd": data.meeting_type,
            "isi_surat_rapat_biasa": data.agenda.strip(),
            "isi_surat_rapat_biasa_skpd": data.agenda.strip(),
            "hari_rapat_biasa": day_name_id(data.meeting_date),
            "tanggal_rapat_biasa": format_date_id(data.meeting_date),
            "jam_rapat_biasa": data.time_text.strip(),
            "tujuan_surat_dprd": purpose,
            "tujuan_surat_lainnya": build_tujuan_richtext(main_destinations),
            "jabatan_ttd_rapat": position,
            "nama_ttd_rapat": name,
        }
        for index, value in enumerate(parties, start=1):
            context[f"pihak_terkait_{index}"] = value
        for index, value in enumerate(parties_mayor, start=1):
            context[f"pihak_terkait_wk_{index}"] = value
        # Undangan biasa memakai pola nomor berurutan yang sama dengan
        # Paripurna: segmen ketiga bertambah, segmen pertama dan kedua tetap.
        for index in range(4):
            context[f"nomor_surat_rapat_{index + 1}"] = increment_nomor_paripurna(data.number, index)

        main_temporary = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
        combined_files = [main_temporary]
        try:
            template = DocxTemplate(TEMPLATE_RAPAT_BIASA)
            ensure_numbered_slots(template.get_docx(), "pihak_terkait", len(parties))
            ensure_numbered_slots(template.get_docx(), "pihak_terkait_wk", len(parties_mayor))
            template.render(context)
            template.save(main_temporary)
            document = Document(main_temporary)
            cleanup_skenario_paripurna(document)
            document.save(main_temporary)
            if extra_destinations:
                cache = Path(tempfile.gettempdir()) / "sips_halaman_tujuan_lain_template.docx"
                extra_template = build_halaman_tujuan_lain_template(TEMPLATE_RAPAT_BIASA, cache)
                for index, destinations in enumerate(extra_destinations):
                    extra_context = dict(context)
                    extra_context["tujuan_surat_lainnya"] = build_tujuan_richtext(destinations)
                    extra_context["nomor_surat_rapat_4"] = increment_nomor_paripurna(data.number, 4 + index)
                    temporary = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    extra = DocxTemplate(extra_template)
                    extra.render(extra_context)
                    extra.save(temporary)
                    combined_files.append(temporary)
            if len(combined_files) > 1:
                _combine_word_pages(combined_files, output_path)
            else:
                shutil.copy(main_temporary, output_path)
        finally:
            for path in combined_files:
                Path(path).unlink(missing_ok=True)

    @staticmethod
    def _generate_meeting_attendance(data: InvitationFormData, output_path: str | Path) -> None:
        if data.invitation_type == "paripurna":
            main_template = TEMPLATE_DAFTAR_HADIR_PARIPURNA
            additional: list[str] = []
        else:
            main_template = TEMPLATE_DAFTAR_HADIR_RAPAT_MAP.get(data.meeting_executor)
            if not main_template:
                raise ValueError(
                    f"Template daftar hadir untuk '{data.meeting_executor}' belum tersedia."
                )
            additional = []
            if data.include_related_attendance:
                additional.append(TEMPLATE_DAFTAR_HADIR_PIHAK_TERKAIT)
            if data.include_secretariat_attendance:
                additional.append(TEMPLATE_DAFTAR_HADIR_SEKRETARIAT)
            if any(_is_taf(value) for value in data.related_parties):
                additional.append(TEMPLATE_DAFTAR_HADIR_TAF)
        generate_daftar_hadir_rapat(
            main_template,
            output_path,
            hari=day_name_id(data.meeting_date),
            tanggal=format_date_id(data.meeting_date),
            jam=data.time_text,
            isi_perihal=data.agenda,
            is_paripurna=data.invitation_type == "paripurna",
            lembar_tambahan=additional,
        )
