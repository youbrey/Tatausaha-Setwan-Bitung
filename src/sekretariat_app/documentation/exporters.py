from __future__ import annotations

import json
import tempfile
import zipfile
from pathlib import Path

from PySide6.QtCore import QMarginsF, QRectF, QSizeF
from PySide6.QtGui import QPageLayout, QPageSize, QPainter, QPdfWriter
from PySide6.QtPrintSupport import QPrintDialog, QPrinter, QPrinterInfo
from PySide6.QtWidgets import QWidget

from sekretariat_app.documentation.models import DocumentProject
from sekretariat_app.documentation.scene import DocumentScene


class DocumentationExporter:
    def _scene_for_page(self, project: DocumentProject, page_index: int) -> DocumentScene:
        width, height = project.page_size_mm
        margins = (
            project.margins.top,
            project.margins.right,
            project.margins.bottom,
            project.margins.left,
        )
        scene = DocumentScene()
        letterhead = project.letterhead if page_index == 0 else {**project.letterhead, "enabled": False}
        scene.load_page(project.pages[page_index], width, height, margins, letterhead)
        return scene

    def export_png_pages(self, project: DocumentProject, directory: str | Path, dpi: int = 300) -> list[Path]:
        target_dir = Path(directory)
        target_dir.mkdir(parents=True, exist_ok=True)
        files: list[Path] = []
        stem = self._safe_name(project.title)
        for index in range(len(project.pages)):
            scene = self._scene_for_page(project, index)
            image = scene.render_image(dpi)
            target = target_dir / f"{stem}-halaman-{index + 1:02d}.png"
            if not image.save(str(target), "PNG"):
                raise OSError(f"Gagal menulis {target}")
            scene.clear()
            del image
            files.append(target)
        return files

    def export_pdf(self, project: DocumentProject, path: str | Path, dpi: int = 300) -> Path:
        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        width_mm, height_mm = project.page_size_mm
        page_size = QPageSize(QSizeF(width_mm, height_mm), QPageSize.Unit.Millimeter, project.paper_size)
        writer = QPdfWriter(str(target))
        writer.setResolution(dpi)
        writer.setPageLayout(QPageLayout(page_size, QPageLayout.Orientation.Portrait, QMarginsF(0, 0, 0, 0)))
        painter = QPainter(writer)
        for index in range(len(project.pages)):
            if index:
                writer.newPage()
            scene = self._scene_for_page(project, index)
            scene.render_to_painter(painter, QRectF(0, 0, writer.width(), writer.height()))
            scene.clear()
        painter.end()
        return target

    def export_docx(self, project: DocumentProject, path: str | Path, dpi: int = 300) -> Path:
        from docx import Document
        from docx.enum.text import WD_BREAK
        from docx.shared import Mm

        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        width_mm, height_mm = project.page_size_mm
        with tempfile.TemporaryDirectory(prefix="dokufoto-docx-") as temp_dir:
            images = self.export_png_pages(project, temp_dir, dpi)
            document = Document()
            section = document.sections[0]
            section.page_width = Mm(width_mm)
            section.page_height = Mm(height_mm)
            section.top_margin = Mm(0)
            section.right_margin = Mm(0)
            section.bottom_margin = Mm(0)
            section.left_margin = Mm(0)
            section.header_distance = Mm(0)
            section.footer_distance = Mm(0)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            paragraph.paragraph_format.line_spacing = 1
            for index, image_path in enumerate(images):
                if index:
                    paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
                paragraph.add_run().add_picture(str(image_path), width=Mm(width_mm), height=Mm(height_mm))
                if index < len(images) - 1:
                    paragraph.add_run().add_break(WD_BREAK.PAGE)
            document.save(target)
        return target

    @staticmethod
    def available_printer_names() -> list[str]:
        """Baca printer aktif langsung dari spooler/driver lokal Windows."""

        return [
            info.printerName()
            for info in QPrinterInfo.availablePrinters()
            if info.printerName()
        ]

    def print_project(
        self,
        project: DocumentProject,
        parent: QWidget,
        printer_name: str = "",
    ) -> bool:
        available = QPrinterInfo.availablePrinters()
        if not available:
            raise RuntimeError(
                "Printer tidak ditemukan. Pastikan printer terpasang, aktif, dan "
                "terlihat pada Settings > Bluetooth & devices > Printers & scanners."
            )
        selected_info = next(
            (info for info in available if info.printerName() == printer_name),
            None,
        )
        if printer_name and selected_info is None:
            raise RuntimeError(
                f"Printer {printer_name} tidak lagi ditemukan. Deteksi ulang printer lalu coba kembali."
            )
        if selected_info is None:
            selected_info = QPrinterInfo.defaultPrinter()
        if selected_info.isNull():
            selected_info = available[0]
        printer = QPrinter(selected_info, QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.NativeFormat)
        printer.setDocName(project.title)
        printer.setColorMode(QPrinter.ColorMode.Color)
        printer.setFullPage(True)
        width_mm, height_mm = project.page_size_mm
        page_size = QPageSize(QSizeF(width_mm, height_mm), QPageSize.Unit.Millimeter, project.paper_size)
        page_layout = QPageLayout(
            page_size,
            QPageLayout.Orientation.Portrait,
            QMarginsF(0, 0, 0, 0),
        )
        if not printer.setPageLayout(page_layout):
            raise RuntimeError(
                f"Printer {selected_info.printerName()} tidak menerima ukuran kertas "
                f"{width_mm:g} × {height_mm:g} mm."
            )
        dialog = QPrintDialog(printer, parent)
        dialog.setWindowTitle("Cetak Dokumentasi Foto")
        if dialog.exec() != QPrintDialog.DialogCode.Accepted:
            return False
        if not printer.isValid() or not printer.printerName():
            raise RuntimeError("Printer yang dipilih tidak lagi tersedia atau drivernya tidak aktif.")

        # Dialog driver dapat mengganti ukuran halaman. Terapkan kembali ukuran
        # proyek supaya 1 mm pada canvas tetap 1 mm pada hasil cetak.
        printer.setFullPage(True)
        printer.setColorMode(QPrinter.ColorMode.Color)
        if not printer.setPageLayout(page_layout):
            raise RuntimeError(
                f"Printer {printer.printerName()} tidak mendukung ukuran halaman proyek."
            )
        actual_size = printer.pageLayout().pageSize().size(QPageSize.Unit.Millimeter)
        if (
            abs(actual_size.width() - width_mm) > 0.75
            or abs(actual_size.height() - height_mm) > 0.75
        ):
            raise RuntimeError(
                "Driver printer mengubah ukuran kertas menjadi "
                f"{actual_size.width():g} × {actual_size.height():g} mm. "
                f"Pilih atau buat ukuran {width_mm:g} × {height_mm:g} mm pada driver printer."
            )
        painter = QPainter(printer)
        if not painter.isActive():
            raise RuntimeError(f"Gagal membuka printer {printer.printerName()} untuk mencetak.")
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
        painter.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform, True)
        try:
            for index in range(len(project.pages)):
                if index and not printer.newPage():
                    raise RuntimeError(f"Printer gagal memulai halaman {index + 1}.")
                scene = self._scene_for_page(project, index)
                try:
                    target_rect = QRectF(
                        printer.pageLayout().fullRectPixels(printer.resolution())
                    )
                    scene.render_to_painter(painter, target_rect)
                finally:
                    scene.clear()
        finally:
            painter.end()
        return True

    def export_archive(self, project: DocumentProject, path: str | Path) -> Path:
        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        payload = project.to_dict()
        media_map: dict[str, str] = {}
        with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for index, media_path in enumerate(project.media):
                source = Path(media_path)
                if not source.exists() or not source.is_file():
                    continue
                archive_name = f"media/{index:04d}-{source.name}"
                archive.write(source, archive_name)
                media_map[str(source)] = archive_name
            for page in payload.get("pages", []):
                for element in page.get("elements", []):
                    if element.get("kind") == "photo" and element.get("photo_path") in media_map:
                        element["photo_path"] = media_map[element["photo_path"]]
            payload["media"] = list(media_map.values())
            payload["marked_media"] = [
                media_map[path]
                for path in project.marked_media
                if path in media_map
            ]
            payload["archive_format"] = "dokufoto-python-v1"
            archive.writestr("project.dokufoto.json", json.dumps(payload, ensure_ascii=False, indent=2))
        return target

    def import_archive(self, path: str | Path, destination: str | Path) -> DocumentProject:
        source = Path(path)
        target_dir = Path(destination) / source.stem
        target_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(source, "r") as archive:
            names = archive.namelist()
            project_name = next((name for name in names if name == "project.dokufoto.json" or name.endswith("/project.dokufoto.json")), None)
            if not project_name:
                raise ValueError("Arsip tidak memiliki project.dokufoto.json.")
            for name in names:
                candidate = (target_dir / name).resolve()
                if not str(candidate).startswith(str(target_dir.resolve())):
                    raise ValueError("Arsip berisi path yang tidak aman.")
            archive.extractall(target_dir)
        project_path = target_dir / project_name
        payload = json.loads(project_path.read_text(encoding="utf-8"))
        if payload.get("archive_format") != "dokufoto-python-v1":
            project = DocumentProject.load(project_path)
            for candidate in target_dir.rglob("*"):
                if candidate.is_file() and candidate.suffix.lower() in {".jpg", ".jpeg", ".png", ".webp", ".bmp"}:
                    value = str(candidate.resolve())
                    if value not in project.media:
                        project.media.append(value)
                        project.marked_media.append(value)
            return project
        media = [str((target_dir / name).resolve()) for name in payload.get("media", [])]
        mapping = dict(zip(payload.get("media", []), media))
        payload["media"] = media
        if "marked_media" in payload:
            payload["marked_media"] = [
                mapping[name]
                for name in payload.get("marked_media") or []
                if name in mapping
            ]
        for page in payload.get("pages", []):
            for element in page.get("elements", []):
                if element.get("kind") == "photo" and element.get("photo_path") in mapping:
                    element["photo_path"] = mapping[element["photo_path"]]
        return DocumentProject.from_dict(payload)

    @staticmethod
    def _safe_name(name: str) -> str:
        cleaned = "".join(character if character.isalnum() or character in "-_" else "_" for character in name)
        return cleaned.strip("_") or "Dokumentasi_Foto"
